from docx import Document
from docx.shared import Pt
from cyrtranslit import to_latin

import json
import time
import csv
import difflib


def cirilica_u_latinicu(text):
    latinica_text = to_latin(text)
    return latinica_text

def clear_serbian_chars(text: str = None) -> str:
    if text is None:
        return None
    
    replace_table = [
        ["ć", "c"],
        ["č", "c"],
        ["š", "s"],
        ["ž", "z"],
        ["đ", "dj"],
        ["Ć", "c"],
        ["Č", "c"],
        ["Š", "S"],
        ["Ž", "Z"],
        ["Đ", "Dj"]
    ]
    for i in replace_table:
        text = text.replace(i[0], i[1])
    return text

total_t_start = time.perf_counter()
dictionary = {}
dictionary["serbian"] = {}
dictionary["serbian"]["dicts"] = {}

END_OF_TOPIC = ",.!:?[]()\"';"
END_OF_TOPIC_WITH_SPACE = ",.!:?[]()\"'; "

txt_pirot = """aor.	aorist
ar	arapski
arh.	arhaična reč
augm.	augmenitativ
bot.	botanički
br.	broj
bug.	bugarizam; bugarski
v.	vidi
vok.	vokativ
vulg.	vulgario
g	glagol
gen.	genitiv
geogr.	geografski termin
dat.	dativ
dem.	deminutiv
deč.	dečji govor
euf.	eufemizam
ž.	imenica ženskog roda
zam.	zamenica
zb.	zbirna imenica
zoo.	zoološki
izr.	izraz
imper.	imperativ
iron.	ironično
ital.	italijanski
j.	jednina
2 lmn. aor.	drugo lice množine aorista 
lmn. pre.	treće lice množine prezenta 
m.	imenica muškog roda
mn.	množina
nesvr.	nesvršeni glagol
pej	pejorativ
perf.	perfekat
poet.	poetski
prid.	pridev
pril.	prilog
ret.	retka reč
s.	imenica srednjeg roda
svr.	svršeni glagol
tur.	turski
uzv.	uzvik
uskl.	usklik
fig.	figurativno
franc.	francuski
hip.	hipokoristik
cig.	romski
crk.	crkveni termin
šalj.	šaljivo
špans.	španski"""

vlaski_skracenice = """(amer.) = američki-engleski 
(arami.) = aramejski
(arap.) = arapski
(balkkan.) = balkanski
(gr.) = grčki
(egip.) = egipatski
(engl.) = engleski
(esper .) = esperanto 
(indijan.) = indijanski 
(island.) = islandski 
(ital.) = italijanski 
(jap.) = japanski
(kelt.) = keltski
(kines.) = kineski
(lat.) = latinski
(mađ.) = mađarski
(malaj.) = malajski
(meks.) = meksikanski
(nem.) = nemački
(novogr.) = novogrčki
(novolat.) = novolatinski
(norv.) = norveški
(pers.) = persijanski
(peruan.) = peruanski
(port.) = portugalski
(rum.) = rumunski
(rus.) = ruski
(slav.) = slovenski
(sanskr.) =sanskritski 
(staroslov.) = staroslovenski 
(starogerm.) = starogermanski 
(tur.) = turski
(turkm.) = turkmenski
(fins.) = finski
(fr.) = francuski
(hebr.) = hebrejski
(holan.) = holandski
(crkv-slov.) = crkvenoslovenski
(češ.) = češki
(špan.) = španski"""






with open("spisak_srpskih_reci.txt", "r", encoding="utf-8") as file:
    text = file.read()

with open("srpske_reci2.txt", "r", encoding="utf-8") as file:
    text2 = file.read()


dictionary["serbian"]["latin"] = {}
dictionary["serbian"]["latin"]["abc+"] = []
dictionary["serbian"]["latin"]["abc"] = []

with_spec_set = set([x.strip() for x in text.split("\n") if x.strip() != ""])
text2 = text2.replace("\n", " ")
text2_list = [x for x in text2.split(" ") if x != ""]
for i in text2_list:
    with_spec_set.add(i)


with_spec_list = list(with_spec_set)

with_spec_list.sort()
with_spec_list.pop(0)
no_spec_list = [clear_serbian_chars(x) for x in with_spec_list]

print (f"Begin. with_spec_list={len(with_spec_list)}   no_spec_list={len(no_spec_list)}")
if len(with_spec_list) != len(no_spec_list):
    print (f"Error.  len do not match. with_spec_list={len(with_spec_list)}   no_spec_list={len(no_spec_list)}")

for i in range(len(with_spec_list)):
    dictionary["serbian"]["latin"]["abc+"].append(with_spec_list[i])
    dictionary["serbian"]["latin"]["abc"].append(no_spec_list[i])

print ("Serbian Recapitulation:")
print (f"Raw data : with_spec_list={len(with_spec_list)}   no_spec_list={len(no_spec_list)}")
print (f'Processed data : abc+={len(dictionary["serbian"]["latin"]["abc+"])} abc={len(dictionary["serbian"]["latin"]["abc"])}')
print ()




def topic_find(bold_text: str, paragraph: str) -> str:
    bold_text = bold_text.replace("Nj", "NJ")
    bold_text = bold_text.replace("Lj", "LJ")
    paragraph = paragraph.replace("Nj", "NJ")
    paragraph = paragraph.replace("Lj", "LJ")
    
    bold_list = [x.strip() for x in bold_text.split(",") if len(x) > 1]
    new_topic = ""
    new_bold = ""
    for i in bold_list:
        if i != i.upper():
            break
        new_topic += i + " "
    new_topic = new_topic.strip()
    bold_text = ",".join([x for x in bold_list if x not in new_topic.split(" ")])
    bold_text = f"{new_topic},{bold_text}"
    
    topic = ""
    new_topic = ""
    for i in bold_text:
        new_topic += i
        if i in END_OF_TOPIC:
            break
        if new_topic != new_topic.upper():
            return None
        if new_topic != paragraph[:len(new_topic)]:
            break
        topic += i
    if topic:
        return topic, bold_text
    else:
        return None, bold_text

def vidi_find(bold_text: str) -> list:
    bold_text = bold_text.replace("Nj", "NJ")
    bold_text = bold_text.replace("Lj", "LJ")
    topic = ""
    new_topic = ""
    for i in bold_text:
        new_topic += i
        if i in END_OF_TOPIC:
            break
        if new_topic != new_topic.upper():
            break
        topic += i

    if topic:
        bold_text = bold_text[len(topic):]
    
    do = True
    while do:
        do = False
        if bold_text:
            if bold_text[:1] in END_OF_TOPIC_WITH_SPACE:
                bold_text = bold_text[1:]
                do = True
            if bold_text[-1:] in END_OF_TOPIC_WITH_SPACE:
                bold_text = bold_text[:-1]
                do = True
    
    bold_list = [x.strip("()., ") for x in bold_text.split(",") if x.strip() != ""]

    return bold_list
            
def add_links(original_list: list, links_list: list):
    for i in links_list:
        original_list.append(i)


# kulisicsrpski_mitol_recnik.docx

doc = Document("kulisicsrpski_mitol_recnik.docx")

dictionary["serbian"]["dicts"]["mit"] = {}

txt = ""
start = False
bold_text = ""
topic_name = ""
new_topic = ""
for paragraph in doc.paragraphs:
    par_text = cirilica_u_latinicu(paragraph.text.strip().replace("\t", "")) + "\n"
    if par_text.find("RUSA(L)NA") >= 0:
        par_text = par_text.replace("RUSA(L)NA", "RUSALNA")
    if par_text.find("SV.") >= 0:
        par_text = par_text.replace("SV.", "SVETI")
    if par_text.find("DUKLjAN(IN)") >= 0:
        par_text = par_text.replace("DUKLjAN(IN)", "DUKLJAN")

    if par_text == "A\n":
        start = True
    
    if par_text.find("VAŽNIJA LITERATURA") >= 0:
        start = False

    if not start:
        continue

    bold_text = ""
    if paragraph.alignment == 1:
        continue

    for run in paragraph.runs:
        if run.bold:
            bold_text += run.text + ","
            bold_text = cirilica_u_latinicu(bold_text.strip())

    bold_text = bold_text.replace("—", "")
    
    if bold_text.find("RUSA(L)NA") >= 0:
        bold_text = bold_text.replace("RUSA(L)NA", "RUSALNA")
    if bold_text.find("SV.") >= 0:
        bold_text = bold_text.replace("SV.", "SVETI")
    if bold_text.find("DUKLjAN(IN)") >= 0:
        bold_text = bold_text.replace("DUKLjAN(IN)", "DUKLJAN")
        print ("Found DUKLJAN")
    if bold_text.find(" ili ") >= 0:
        bold_text = bold_text.replace(" ili ", ",")

    bold_text = bold_text.replace("(", ",")
    bold_text = bold_text.replace(")", ",")
    bold_text = bold_text.replace("III", "")

    if bold_text.find("GLUHO") >= 0:
        print ("FOUND GLUHO")

    if bold_text:
        new_topic, bold_text = topic_find(bold_text, par_text)
        if new_topic:
            new_topic = new_topic.strip()
    
    if not topic_name and not new_topic:
        continue

    if bold_text != "":
        links = vidi_find(bold_text)
    else:
        links = []

    if new_topic:
        if new_topic == "KRALJICE" and topic_name == "KRALJICE":
            new_topic = "KRALJICE I KRALJEVI DUBOČKI"
            links.pop(0)
        if new_topic in dictionary["serbian"]["dicts"]["mit"]:
            print ("ERROR:  Duplicate key: ", new_topic)
            print (13/0)
        topic_name = new_topic
        dictionary["serbian"]["dicts"]["mit"][topic_name] = {}
        dictionary["serbian"]["dicts"]["mit"][topic_name]["text"] = par_text
        dictionary["serbian"]["dicts"]["mit"][topic_name]["links"] = []
        add_links(dictionary["serbian"]["dicts"]["mit"][topic_name]["links"], links)
        new_topic = ""
        continue

    dictionary["serbian"]["dicts"]["mit"][topic_name]["text"] += par_text
    add_links(dictionary["serbian"]["dicts"]["mit"][topic_name]["links"], links)

replace_items = [
]
add_items = [
    ["„TAMO ONA", "TAMO ONA"],
    ["„TAMO ONA", "TAMO ON"]
]

for item in dictionary["serbian"]["dicts"]["mit"]:
    make_new_item = []
    pop_items = []
    count = 0

    if item == "ISTOK" or item == "KRALJICE":
        print ("FOUND")

    if item in [x[0] for x in replace_items]:
        continue

    for i in dictionary["serbian"]["dicts"]["mit"][item]["links"]:
        if i == i.upper():
            pop_items.append(count)
            if len(i) > 1:
                make_new_item.append(i)
        count += 1
    pop_items.sort(reverse=True)
    for i in pop_items:
        dictionary["serbian"]["dicts"]["mit"][item]["links"].pop(i)

    for i in make_new_item:
        print (f'ITEM: {item} ++ {i}    LIST: {make_new_item}', end="")
        if item == i.strip(".— "):
            print ("  REJECTED")
            continue
        if i not in dictionary["serbian"]["dicts"]["mit"]:
            print ("  .........................................................................................CREATED")
        else:
            print ("  .........................................................................................DISCARDED")
        if i not in dictionary["serbian"]["dicts"]["mit"]:
            add_items.append([item, i])

delete_items = [
    "14. X",
    "XI",
    "13. XI",
    "„TAMO ONA",
    "„TAMO ON",
    "CRNI",
    "CRNI 'TORNIK"
]

for i in replace_items:
    if i[1] not in dictionary["serbian"]["dicts"]["mit"]:
        dictionary["serbian"]["dicts"]["mit"][i[1]] = {}
        dictionary["serbian"]["dicts"]["mit"][i[1]]["text"] = dictionary["serbian"]["dicts"]["mit"][i[0]]["text"]
        dictionary["serbian"]["dicts"]["mit"][i[1]]["links"] = list(dictionary["serbian"]["dicts"]["mit"][i[0]]["links"])
        dictionary["serbian"]["dicts"]["mit"].pop(i[0])

for i in add_items:
    if i[1] not in dictionary["serbian"]["dicts"]["mit"]:
        dictionary["serbian"]["dicts"]["mit"][i[1]] = {}
        dictionary["serbian"]["dicts"]["mit"][i[1]]["text"] = dictionary["serbian"]["dicts"]["mit"][i[0]]["text"]
        dictionary["serbian"]["dicts"]["mit"][i[1]]["links"] = list(dictionary["serbian"]["dicts"]["mit"][i[0]]["links"])

for i in delete_items:
    dictionary["serbian"]["dicts"]["mit"].pop(i)

for i in dictionary["serbian"]["dicts"]["mit"]:
    dictionary["serbian"]["dicts"]["mit"][i]["links"] = []
    for j in dictionary["serbian"]["dicts"]["mit"]:
        txt = dictionary["serbian"]["dicts"]["mit"][i]["text"].lower()
        for k in END_OF_TOPIC:
            txt = txt.replace(k, " ")
        txt = txt + " "

        if txt.find(f' {j.lower()} ') >= 0:
            if j not in dictionary["serbian"]["dicts"]["mit"][i]["links"] and j != i:
                dictionary["serbian"]["dicts"]["mit"][i]["links"].append(j)



# vujaklija_recnik_stranih_reci.docx

doc = Document("vujaklija_recnik_stranih_reci.docx")

dictionary["serbian"]["dicts"]["vujaklija"] = {}

txt = ""
start = False
bold_text = ""
topic_name = ""
new_topic = ""
for paragraph in doc.paragraphs:
    par_text: str = paragraph.text.strip().replace("\t", "") + "\n"

    if par_text.strip() == "A":
        start = True
    
    if par_text.find("MILAN VUJAKLIJA") >= 0:
        start = False

    if not start:
        continue

    if par_text.lower().startswith("aviomehani"):
        print ()

    if len(par_text.strip()) < 3:
        continue
    if par_text.strip().find(" ") == -1:
        cancel_cont = False
        for i in END_OF_TOPIC:
            if par_text.find(i) >= 0:
                cancel_cont = True
        if not cancel_cont:
            continue

    # fix paragraph
    par_text = par_text.replace("ñ", "đ")
    correct_lj = par_text.find(">")
    if correct_lj > 0:
        if par_text[correct_lj-1:correct_lj].lower() == "l":
            par_text = par_text[:correct_lj] + "j" + par_text[correct_lj+1:]

    # find bold text

    bold_text = ""
    bold_text_alt = ""
    for run in paragraph.runs:
        bold_text_alt += run.text
        if run.bold:
            bold_text = bold_text_alt
    
    while True:
        if bold_text.find("  ") >= 0:
            bold_text = bold_text.replace("  ", " ")
        else:
            break

    bold_text = bold_text.strip()
    bold_text = bold_text.replace("—", "")

    # find item if not bolded
    if not bold_text:
        pos_bracket = par_text.find("(")
        if pos_bracket > 0:
            if len(par_text[:pos_bracket].strip().split(" ")) < 3:
                if par_text[:pos_bracket].find("\\") == -1:
                    bold_text = par_text[:pos_bracket].strip()

    if len(par_text) > 5:
        if par_text[:5] == "šućur":
            bold_text = "šućur"
    
    if bold_text:
        topic_name = ""
        if bold_text in dictionary["serbian"]["dicts"]["vujaklija"]:
            print (f"Duplicate entry : {bold_text}")
            dictionary["serbian"]["dicts"]["vujaklija"][bold_text]["text"] += "\n\n"
        else:
            dictionary["serbian"]["dicts"]["vujaklija"][bold_text] = {}
            dictionary["serbian"]["dicts"]["vujaklija"][bold_text]["text"] = ""
            dictionary["serbian"]["dicts"]["vujaklija"][bold_text]["links"] = []
        topic_name = bold_text
    
    if not topic_name:
        continue

    par_text = par_text.replace("-", "")

    dictionary["serbian"]["dicts"]["vujaklija"][topic_name]["text"] += par_text

print ("Searching for links...")
# a = len(dictionary["serbian"]["dicts"]["vujaklija"])
# count_step = int(a/1000)
# lnks = 0
# for item in dictionary["serbian"]["dicts"]["vujaklija"]:
#     if count % count_step == 0:
#         print (f"{count/a*100: .2f} %")
#     count += 1
#     for i in dictionary["serbian"]["dicts"]["vujaklija"]:
#         if (dictionary["serbian"]["dicts"]["vujaklija"][item]["text"].find(f" {i} ") > 0 or dictionary["serbian"]["dicts"]["vujaklija"][item]["text"].find(f" {i}.") > 0 or dictionary["serbian"]["dicts"]["vujaklija"][item]["text"].find(f" {i})") > 0) and item != i and len(i) > 2:
#             dictionary["serbian"]["dicts"]["vujaklija"][item]["links"].append(i)
#             lnks += 1
    
# print (f"Total {lnks} added.")




# Sanovnik TXT

dictionary["serbian"]["dicts"]["san"] = {}

def san_format_txt(txt: str) -> str:
    txt = txt.replace("–", "-")
    txt = txt.strip()
    while txt.find("  ") >= 0:
        txt = txt.replace("  ", " ")
    return txt

def san_get_title(txt) -> str:
    title, _ = san_split_title_body(txt)
    return title

def san_get_body(txt) -> str:
    _, body = san_split_title_body(txt)
    return body

def san_split_title_body(txt: str) -> tuple:
    if txt.find("-") == -1:
        return ("", san_format_body(txt))

    title = txt[:txt.find("-")].strip()
    body = txt[txt.find("-")+1:].strip()

    pos_space = title.find(" ")
    if pos_space >= 0:
        if title.find(" ", pos_space + 1) == -1:
            pos_space2 = len(title)
        else:
            pos_space2 = title.find(" ", pos_space + 1)
        if title[pos_space:pos_space2].strip().lower() == "ili":
            title = title.upper()

    up_title = title
    while up_title.find("(") >= 0:
        if up_title.find(")") == -1:
            print (f"Split TITLE-BODY error: title: {title}")
            print (f"TEXT: {txt}")
            up_title = up_title[:up_title.find("(")].strip()
            break
        up_remain = ""
        if up_title.find(")") + 1 < len(up_title):
            up_remain = up_title[up_title.find(")") + 1:]

        up_title = up_title[:up_title.find("(")] + ", " + up_title[up_title.find("(")+1:up_title.find(")")].upper() + up_remain

    if up_title.find(" ILI ") >= 0:
        up_title = up_title.replace(" ILI ", ", ")

    if up_title == up_title.upper():
        return (up_title, san_format_body(body))
    else:
        return ("", san_format_body(txt))

def san_format_body(txt):
    txt_list = [x.strip() for x in txt.split("; -") if x.strip() != ""]
    txt = "\n".join(txt_list)
    return txt

def san_get_links(txt: str, current_item: str) -> list:
    links = []
    txt = txt.replace(current_item, "").lower()
    for i in END_OF_TOPIC:
        txt = txt.replace(i, " ")
    for item in dictionary["serbian"]["dicts"]["san"]:
        if txt.find(f" {item.lower()} ") >= 0 and item != current_item:
            links.append(item)
    return links


with open("Sanovnik.txt", "r", encoding="utf-8") as file:
    text = file.read()

text_list = [x.strip() for x in text.split("\n") if x.strip() != ""]

title = ""
error_flag = False
for line in text_list:
    line = san_format_txt(line)

    new_title = san_get_title(line)
    body = san_get_body(line)

    links = []

    if new_title:
        if new_title in dictionary["serbian"]["dicts"]["san"]:
            print (f'Error - duplicate entry: {new_title}')
            error_flag = True
            continue
        error_flag = False
        dictionary["serbian"]["dicts"]["san"][new_title] = {}
        dictionary["serbian"]["dicts"]["san"][new_title]["text"] = f'{new_title}\n{body}'
        dictionary["serbian"]["dicts"]["san"][new_title]["links"] = []
        for link in links:
            dictionary["serbian"]["dicts"]["san"][new_title]["links"].append(link)
        title = new_title
        continue
    
    if error_flag:
        continue
    dictionary["serbian"]["dicts"]["san"][title]["text"] += f' {body}'
    for link in links:
        dictionary["serbian"]["dicts"]["san"][title]["links"].append(link)


for item in dictionary["serbian"]["dicts"]["san"]:
    txt = dictionary["serbian"]["dicts"]["san"][item]["text"]
    dictionary["serbian"]["dicts"]["san"][item]["links"] = san_get_links(txt, item)


print (f'Sanovnik 1 completed. Records: {len(dictionary["serbian"]["dicts"]["san"])}')



# Sanovnik DOCX

def san2_get_title(txt: str) -> str:
    count = 0
    while txt.find("(") >= 0:
        count += 1
        if count > 10:
            print ("Endless loop")
            break

        pos1 = txt.find("(")
        pos2 = txt.find(")")
        if pos2 < 0:
            print ("Error. ", txt)
            break
        if pos2 + 1 < len(txt):
            txt = txt[:pos1].strip() + ", " + txt[pos1+1:pos2].upper() + txt[pos2+1:]
        else:
            txt = txt[:pos1].strip() + ", " + txt[pos1+1:pos2].upper()

    if txt == txt.upper():
        return txt
    else:
        return ""

def san2_get_links(txt: str, current_item: str) -> list:
    links = []
    txt = txt.replace(current_item, "").lower()
    for i in END_OF_TOPIC:
        txt = txt.replace(i, " ")
    for item in dictionary["serbian"]["dicts"]["san"]:
        if txt.find(f" {item.lower()} ") >= 0 and item != current_item:
            links.append(item)
    return links

doc = Document("sanovnik-recnik-snova-ii-deopdf.docx")

dictionary["serbian"]["dicts"]["sanovnik2"] = {}

txt = ""
start = False
concatenate = ""
topic_name = ""
new_topic = ""
for paragraph in doc.paragraphs:
    par_text = paragraph.text.strip().replace("\t", "") + "\n"

    if par_text == "ABECEDA\n":
        start = True
    
    if par_text == "ŽUTA BOJA\n":
        print ("At end")

    if not start:
        continue

    if concatenate:
        par_text = concatenate + par_text
        concatenate = ""
    else:
        find_space = par_text.find(" ")
        if find_space > 0:
            if par_text[:find_space] == par_text[:find_space].upper() and par_text.strip()[-1] == "-":
                concatenate = par_text.strip()[:-1]
                continue

    new_topic = ""
    body = par_text.strip()
    is_valid = True
    validate_text = par_text.strip()
    if validate_text.find("SANOVNIK") >= 0:
        validate_text = validate_text.replace("SANOVNIK", "")
        try:
            test = int(validate_text)
            is_valid = False
        except ValueError or TypeError:
            is_valid = True
            

    if not is_valid or len(par_text) < 3:
        continue

    new_topic = san2_get_title(par_text.strip())

    links = []

    if new_topic:
        if new_topic in dictionary["serbian"]["dicts"]["sanovnik2"]:
            print ("ERROR:  Duplicate key: ", new_topic)
            continue
        topic_name = new_topic
        dictionary["serbian"]["dicts"]["sanovnik2"][topic_name] = {}
        dictionary["serbian"]["dicts"]["sanovnik2"][topic_name]["text"] = f"{new_topic}\n"
        dictionary["serbian"]["dicts"]["sanovnik2"][topic_name]["links"] = []
        for i in links:
            dictionary["serbian"]["dicts"]["sanovnik2"][topic_name]["links"].append(i)
        new_topic = ""
        continue

    if dictionary["serbian"]["dicts"]["sanovnik2"][topic_name]["text"][-1] == "-":
        dictionary["serbian"]["dicts"]["sanovnik2"][topic_name]["text"] = dictionary["serbian"]["dicts"]["sanovnik2"][topic_name]["text"][:-1]

    dictionary["serbian"]["dicts"]["sanovnik2"][topic_name]["text"] += f"{body} "
    for i in links:
        dictionary["serbian"]["dicts"]["sanovnik2"][topic_name]["links"].append(i)

for item in dictionary["serbian"]["dicts"]["sanovnik2"]:
    txt = dictionary["serbian"]["dicts"]["sanovnik2"][item]["text"]
    pos = 0
    while True:
        pos = txt.find("- ", pos)
        if pos == -1:
            break
        if pos > 0:
            if txt[pos-1:pos] != " ":
                txt = txt[:pos] + txt[pos+2:]
        pos += 1

    txt_list = [x.strip() for x in txt.split(".") if x.strip() != ""]
    txt = "\n".join(txt_list)
    dictionary["serbian"]["dicts"]["sanovnik2"][item]["text"] = txt

for item in dictionary["serbian"]["dicts"]["sanovnik2"]:
    if item in dictionary["serbian"]["dicts"]["san"]:
        dictionary["serbian"]["dicts"]["san"][item]["text"] += "\n\nSanovnik (2):\n"
        dictionary["serbian"]["dicts"]["san"][item]["text"] += dictionary["serbian"]["dicts"]["sanovnik2"][item]["text"]
    else:
        dictionary["serbian"]["dicts"]["san"][item] = {}
        dictionary["serbian"]["dicts"]["san"][item]["text"] = dictionary["serbian"]["dicts"]["sanovnik2"][item]["text"]
        dictionary["serbian"]["dicts"]["san"][item]["links"] = []
for item in dictionary["serbian"]["dicts"]["san"]:
    txt = dictionary["serbian"]["dicts"]["san"][item]["text"]
    dictionary["serbian"]["dicts"]["san"][item]["links"] = san2_get_links(txt, item)

dictionary["serbian"]["dicts"].pop("sanovnik2")

print (f'Sanovnik 2 concatenated to san.  LEN = {len(dictionary["serbian"]["dicts"]["san"])}')



# Recnik Zargona
# 9=ć  3=č  +=Č  &=Ć  :=Đ  )(=đ  ∗=ignore  –=Separator

def zargon_get_title(text: str) -> str:
    spliter_pos = text.find("–")
    if spliter_pos == -1:
        return ""
    
    txt = text[:spliter_pos].strip()

    while txt.find("(") >= 0:
        pos1 = txt.find("(")
        pos2 = txt.find(")")
        if pos2 < 0:
            print (f"Error. {txt}   .... returning: {_zargon_format_line(txt)}")
            return _zargon_format_line(txt)
        if pos2 + 1 < len(txt):
            txt = txt[:pos1].strip() + ", " + txt[pos1+1:pos2] + txt[pos2+1:]
        else:
            txt = txt[:pos1].strip() + ", " + txt[pos1+1:pos2]

    return _zargon_format_line(txt)

def zargon_get_body(text: str) -> str:
    if text:
        if text[-1] == "-":
            text = text[:-1]
    
    spliter_pos = text.find("–")
    if spliter_pos == -1:
        return _zargon_format_line(text.strip())
    
    text = text[spliter_pos + 1:].strip()
    return _zargon_format_line(text)

def _zargon_format_line(txt: str) -> str:
    txt = txt.replace("9", "ć")
    txt = txt.replace("3", "č")
    txt = txt.replace("+", "Č")
    txt = txt.replace("&", "ć")
    txt = txt.replace(":", "Đ")
    txt = txt.replace(")", "đ")
    txt = txt.replace("(", "đ")
    return txt

def zargon_get_links(txt: str) -> list:
    links = []
    text_sections = [x for x in txt.split(",") if x != ""]
    for txt in text_sections:
        txt_list = [x for x in txt.split(" ") if x != ""]
        count = 0
        search = ""
        while count < len(txt_list):
            search += txt_list[count]
            if search in dictionary["serbian"]["dicts"]["zargon"]:
                links.append(search)
            search += " "
            count += 1
    return links


doc = Document("zargon.docx")

dictionary["serbian"]["dicts"]["zargon"] = {}

txt = ""
start = False
concatenate = ""
topic_name = ""
new_topic = ""
for paragraph in doc.paragraphs:
    par_text = paragraph.text.strip().replace("\t", "")
    par_text = par_text.replace("- ", "")

    if par_text.find("Žargonizmi Holandija, Jevrej i Šiptar") >= 0:
        continue

    if par_text == "RE+NIK":
        start = True
        continue
    
    if par_text.find("Sociolingvistika. Nov") >= 0:
        print (f"At end:  {par_text}")
        start = False

    if not start:
        continue
    
    if par_text.find("čam(uga)") >=0:
        par_text = par_text.replace("čam(uga)", "čamuga")
    
    if par_text.strip() == "":
        topic_name = ""

    segments = []
    normal_done = False
    segment_text = ""
    for run in paragraph.runs:
        if run.bold:
            if segment_text and normal_done:
                segments.append(segment_text)
                segment_text = ""
            segment_text += run.text
            normal_done = False
        else:
            segment_text += run.text
            normal_done = True
    if segment_text and normal_done:
        segments.append(segment_text)

    for par_text in segments:
        new_topic = ""
        body = ""
        is_valid = True
        validate_text = par_text.strip()
        if validate_text.find("OBELE@JA NOVOSADSKOG OMLADINSKOG @ARGONA SA RE^NIKOM") >= 0:
            is_valid = False

        if not is_valid or len(par_text) < 3:
            continue

        new_topic = zargon_get_title(par_text)
        body = zargon_get_body(par_text)

        body += " "
        if body.find(" đ") >= 0 and (body.find("đ ") >= 0 or body.find("đ,") >= 0):
            body = body.replace(" đ", " (").replace("đ ", ") ").replace("đ,", "),")
        
        body = body.strip()
        if body.find("- "):
            body = body.replace("- ", "")

        links = []

        if new_topic:
            if new_topic in dictionary["serbian"]["dicts"]["zargon"]:
                print ("ERROR:  Duplicate key: ", new_topic)
                continue
            topic_name = new_topic
            dictionary["serbian"]["dicts"]["zargon"][topic_name] = {}
            dictionary["serbian"]["dicts"]["zargon"][topic_name]["text"] = f"{new_topic}\n{body}"
            dictionary["serbian"]["dicts"]["zargon"][topic_name]["links"] = []
            for i in links:
                dictionary["serbian"]["dicts"]["zargon"][topic_name]["links"].append(i)
            new_topic = ""
            continue

        if topic_name:
            if dictionary["serbian"]["dicts"]["zargon"][topic_name]["text"][-1] == "-":
                dictionary["serbian"]["dicts"]["zargon"][topic_name]["text"] = dictionary["serbian"]["dicts"]["sanovnik2"][topic_name]["text"][:-1]
            else:
                dictionary["serbian"]["dicts"]["zargon"][topic_name]["text"] += " "

            dictionary["serbian"]["dicts"]["zargon"][topic_name]["text"] += f"{body} "
            for i in links:
                dictionary["serbian"]["dicts"]["zargon"][topic_name]["links"].append(i)


for item in dictionary["serbian"]["dicts"]["zargon"]:
    txt = dictionary["serbian"]["dicts"]["zargon"][item]["text"]
    txt_list = [x.strip() for x in txt.split(" v.") if x.strip() != ""]
    
    if len(txt_list) < 2:
        continue

    for i in range(len(txt_list)):
        if i == 0:
            continue
        links = zargon_get_links(txt_list[i])
        for link in links:
            if link not in dictionary["serbian"]["dicts"]["zargon"][item]["links"]:
                dictionary["serbian"]["dicts"]["zargon"][item]["links"].append(link)
        
    

print (f'Zargon finnished.  LEN = {len(dictionary["serbian"]["dicts"]["zargon"])}')



# Bosanski recnik

def bos_get_title_and_body(txt: str) -> tuple:
    pos = txt.find("=")
    if pos == -1:
        print (f"Error, separator not found:  {txt}")
        return ("", "")
    
    title = txt[:pos].strip()
    body = txt[pos+1:].strip()
    if body.find("=") >= 0:
        print (f"Error, multi separator detected: {txt}")
    
    return (title, body)


dictionary["serbian"]["dicts"]["bos"] = {}

with open("bosanski_recnik.txt", "r", encoding="utf-8") as file:
    txt = file.read()

txt_lines = [x.strip() for x in txt.split("\n") if x.strip() != ""]

for line in txt_lines:
    title, body = bos_get_title_and_body(line)

    if title in dictionary["serbian"]["dicts"]["bos"]:
        print (f"Duplicate entry: {line}")
        continue

    dictionary["serbian"]["dicts"]["bos"][title] = {}
    dictionary["serbian"]["dicts"]["bos"][title]["text"] = body
    dictionary["serbian"]["dicts"]["bos"][title]["links"] = []

lnks = 0
for item in dictionary["serbian"]["dicts"]["bos"]:
    txt = dictionary["serbian"]["dicts"]["bos"][item]["text"] + " "
    for i in END_OF_TOPIC:
        txt = txt.replace(i, " ")

    for i in dictionary["serbian"]["dicts"]["bos"]:
        if txt.lower().find(f" {i.lower()} ") >= 0 and item != i:
            dictionary["serbian"]["dicts"]["bos"][item]["links"].append(i)
            lnks += 1

print (f'Bosanski recnik:  Total {len(dictionary["serbian"]["dicts"]["bos"])} records.  Links: {lnks}')


# Englesko-Srpski recnik sa izgovorom

def en_sr_get_data(txt: str) -> tuple:
    txt_list = [x.strip() for x in txt.split("\t")]
    if len(txt_list) != 3:
        print (f"Error. LEN = {len(txt_list)},   {txt}")
        return ("", "", "")
    title = txt_list[0]
    izgovor = txt_list[1]
    body = txt_list[2]
    return (title, body, izgovor)

dictionary["serbian"]["dicts"]["en-sr"] = {}

with open("englesko_srpski_recnik_sa_izgovorom.txt", "r", encoding="utf-8") as file:
    txt = file.read()

txt_lines = [x.strip() for x in txt.split("\n") if x.strip() != ""]

for line in txt_lines:
    title, body, izgovor = en_sr_get_data(line)

    if not title:
        continue

    if title in dictionary["serbian"]["dicts"]["en-sr"]:
        print (f"Duplicate entry: {line}")
        continue

    dictionary["serbian"]["dicts"]["en-sr"][title] = {}
    dictionary["serbian"]["dicts"]["en-sr"][title]["text"] = body
    dictionary["serbian"]["dicts"]["en-sr"][title]["links"] = []
    dictionary["serbian"]["dicts"]["en-sr"][title]["izgovor"] = izgovor


print (f'en-sr :  Total {len(dictionary["serbian"]["dicts"]["en-sr"])} records.')



# Psiholoski recnik

def psiho_get_links(txt: str, curr_item: str) -> list:
    txt += " "
    for i in END_OF_TOPIC:
        txt = txt.replace(i, " ")

    links = []
    for i in dictionary["serbian"]["dicts"]["psiho"]:
        if i == curr_item:
            continue
        search = f" {i} "
        if txt.find(search) >= 0:
            if i not in links:
                links.append(i)
    return links


doc = Document("recnik_psiholoskih_pojmova.docx")

dictionary["serbian"]["dicts"]["psiho"] = {}

start = False
bold_text = ""
body = ""
topic_name = ""
new_topic = ""
for paragraph in doc.paragraphs:
    par_text = paragraph.text.strip()

    if par_text.strip() == "a":
        start = True
        continue
    
    if not start:
        continue

    if len(par_text.strip()) < 3:
        topic_name = ""
        continue
    try:
        _ = int(par_text)
        continue
    except ValueError:
        _ = "Ok"

    new_topic = ""
    body = ""
    for run in paragraph.runs:
        if run.bold:
            new_topic += run.text
        else:
            body += run.text

    new_topic = new_topic.strip()

    links = []

    if new_topic:
        topic_name = new_topic
        new_topic = ""
        if new_topic in dictionary["serbian"]["dicts"]["psiho"]:
            dictionary["serbian"]["dicts"]["psiho"][topic_name]["text"] += "\n"
            print ("ERROR:  Duplicate key: ", topic_name)
        else:
            dictionary["serbian"]["dicts"]["psiho"][topic_name] = {}
            dictionary["serbian"]["dicts"]["psiho"][topic_name]["text"] = par_text + " "
            dictionary["serbian"]["dicts"]["psiho"][topic_name]["links"] = []
            new_topic = ""
            continue

    if topic_name:
        dictionary["serbian"]["dicts"]["psiho"][topic_name]["text"] += f"{par_text} "


with open("psiholoski_recnik2.txt", "r", encoding="utf-8") as file:
    txt = file.read()

txt_list = [x.strip() for x in txt.split("\n") if x.strip() != ""]

for line in txt_list:
    line = line.replace("–", "-")
    pos = line.find("-")

    if pos == -1:
        print (f"Error, no separator: {line}")
        continue

    title = line[:pos].strip().upper()

    if title in dictionary["serbian"]["dicts"]["psiho"]:
        dictionary["serbian"]["dicts"]["psiho"][title]["text"] += "\n\n"
        dictionary["serbian"]["dicts"]["psiho"][title]["text"] += line
    else:
        dictionary["serbian"]["dicts"]["psiho"][title] = {}
        dictionary["serbian"]["dicts"]["psiho"][title]["text"] = line
        dictionary["serbian"]["dicts"]["psiho"][title]["links"] = []

print (f"Added from psiholoski_recnik2 {len(txt_list)} items.")

lnks = 0
for item in dictionary["serbian"]["dicts"]["psiho"]:
    txt = dictionary["serbian"]["dicts"]["psiho"][item]["text"]
    links = psiho_get_links(txt, item)
    dictionary["serbian"]["dicts"]["psiho"][item]["links"] = links
    lnks += len(links)
        
    

print (f'psiho finnished.  LEN = {len(dictionary["serbian"]["dicts"]["psiho"])},  Links: {lnks}')



# Recnik starih izraza

dictionary["serbian"]["dicts"]["stari_izrazi"] = {}

with open("stari_izrazi.txt", "r", encoding="utf-8") as file:
    txt = file.read()

txt_list = [x.strip() for x in txt.split("\n") if len(x.strip()) > 2]

for line in txt_list:
    line = line.replace("–", "-")
    pos = line.find("-")

    if pos == -1:
        print (f"Error, no separator: {line}")
        continue

    new_title = line[:pos].strip()
    body = line[pos:].strip()
    if not new_title:
        body = line.strip()

    if new_title != new_title.upper():
        print (f"Error, title incorect case: {line}")
        title = title.upper()

    if new_title:
        title = new_title
        if title in dictionary["serbian"]["dicts"]["stari_izrazi"]:
            dictionary["serbian"]["dicts"]["stari_izrazi"][title]["text"] += "\n\n"
            dictionary["serbian"]["dicts"]["stari_izrazi"][title]["text"] += body
        else:
            dictionary["serbian"]["dicts"]["stari_izrazi"][title] = {}
            dictionary["serbian"]["dicts"]["stari_izrazi"][title]["text"] = body
            dictionary["serbian"]["dicts"]["stari_izrazi"][title]["links"] = []
        continue    

    dictionary["serbian"]["dicts"]["stari_izrazi"][title]["text"] += f"\n{body}"


print (f"Added from stari_izrazi {len(txt_list)} items.")

print (f'stari_izrazi finnished.  LEN = {len(dictionary["serbian"]["dicts"]["stari_izrazi"])}')



# Filozofski recnik

def filozof_get_links(txt: str, curr_item: str) -> list:
    txt += " "
    for i in END_OF_TOPIC:
        txt = txt.replace(i, " ")

    links = []
    for i in dictionary["serbian"]["dicts"]["filoz"]:
        if i == curr_item or len(i) < 2:
            continue
        search = f" {i} "
        if txt.find(search) >= 0:
            if i not in links:
                links.append(i)
    return links


dictionary["serbian"]["dicts"]["filoz"] = {}

with open("filozofski_recnik.txt", "r", encoding="utf-8") as file:
    txt = file.read()

txt_list = [x.strip() for x in txt.split("\n")]


new_topic = ""
topic = ""
topic_ended = 2
for line in txt_list:
    line = line.replace("–", "-").strip()
    pos = line.find("-")

    if not line:
        topic_ended += 1
        continue

    
    new_topic = ""
    if pos >= 0:
        if topic_ended >= 2:
            new_topic = line[:pos].strip()
            body = line[pos+1:].strip()
        else:
            if line[pos+1:].strip()[0] == "(":
                new_topic = line[:pos].strip()
                body = line[pos+1:].strip()
            else:
                body = " " + line
    else:
        body = " " + line
    
    topic_ended = 0
    
    if new_topic:
        topic = new_topic
        if topic in dictionary["serbian"]["dicts"]["filoz"]:
            dictionary["serbian"]["dicts"]["filoz"][topic]["text"] += "\n\n"
            dictionary["serbian"]["dicts"]["filoz"][topic]["text"] += body
        else:
            dictionary["serbian"]["dicts"]["filoz"][topic] = {}
            dictionary["serbian"]["dicts"]["filoz"][topic]["text"] = body.strip()
            dictionary["serbian"]["dicts"]["filoz"][topic]["links"] = []
        continue    

    dictionary["serbian"]["dicts"]["filoz"][topic]["text"] += body

lnks = 0
for item in dictionary["serbian"]["dicts"]["filoz"]:
    txt = dictionary["serbian"]["dicts"]["filoz"][item]["text"]

    links = filozof_get_links(txt, item)
    dictionary["serbian"]["dicts"]["filoz"][item]["links"] = links
    lnks += len(links)

print (f'filoz finnished.  LEN = {len(dictionary["serbian"]["dicts"]["filoz"])},  Links: {lnks}')





# Recnik emocija

def emo_get_links(txt: str, curr_item: str) -> list:
    txt += " "
    for i in END_OF_TOPIC:
        txt = txt.replace(i, " ")

    links = []
    for i in dictionary["serbian"]["dicts"]["emo"]:
        if i == curr_item or len(i) < 2:
            continue
        search = f" {i} "
        if txt.find(search) >= 0:
            if i not in links:
                links.append(i)
    return links


dictionary["serbian"]["dicts"]["emo"] = {}

with open("recnik_emocija.txt", "r", encoding="utf-8") as file:
    txt = file.read()

txt_list = [x.strip() for x in txt.split("\n")]


new_topic = ""
topic = ""
topic_ended = 1
for line in txt_list:
    line = line.replace("–", "-").strip()

    if not line:
        topic_ended += 1
        continue

    
    new_topic = ""
    if topic_ended >= 1:
        new_topic = line
        body = ""
    else:
        body = " " + line
    
    topic_ended = 0
    
    if new_topic:
        topic = new_topic
        if topic in dictionary["serbian"]["dicts"]["emo"]:
            dictionary["serbian"]["dicts"]["emo"][topic]["text"] += "\n\n"
            dictionary["serbian"]["dicts"]["emo"][topic]["text"] += body.strip()
        else:
            dictionary["serbian"]["dicts"]["emo"][topic] = {}
            dictionary["serbian"]["dicts"]["emo"][topic]["text"] = body.strip()
            dictionary["serbian"]["dicts"]["emo"][topic]["links"] = []
        continue    

    if dictionary["serbian"]["dicts"]["emo"][topic]["text"]:
        dictionary["serbian"]["dicts"]["emo"][topic]["text"] += f" {body}"
    else:
        dictionary["serbian"]["dicts"]["emo"][topic]["text"] += body

lnks = 0
for item in dictionary["serbian"]["dicts"]["emo"]:
    txt = dictionary["serbian"]["dicts"]["emo"][item]["text"]

    links = emo_get_links(txt, item)
    dictionary["serbian"]["dicts"]["emo"][item]["links"] = links
    lnks += len(links)

print (f'emo finnished.  LEN = {len(dictionary["serbian"]["dicts"]["emo"])},  Links: {lnks}')



# Narodna verovanja o biljkama

def bilj_get_links(txt: str, curr_item: str) -> list:
    txt += " "
    for i in END_OF_TOPIC:
        txt = txt.replace(i, " ")

    links = []
    for i in dictionary["serbian"]["dicts"]["biljke"]:
        if i == curr_item or len(i) < 2:
            continue
        search = f" {i} "
        if txt.find(search) >= 0:
            if i not in links:
                links.append(i)
    return links

def bilj_fix_text(txt: str, item: str) -> str:
    original_txt = txt

    txt = txt + " "
    item_clean = item
    for i in END_OF_TOPIC:
        if i != ".":
            txt = txt.replace(i, " ")
        item_clean = item_clean.replace(i, " ")
    
    item_list = [x.strip() for x in item_clean.split(" ") if x.strip() != ""]
    repl_list = []
    for idx, i in enumerate(item_list):
        repl_list.append([f" {i[0].lower()}. ", i.lower(), f"|{idx*2}"])
        repl_list.append([f" {i[0].upper()}. ", i.capitalize(), f"|{idx*2+1}"])
        if idx == 4:
            break
    
    for i in repl_list:
        pos = txt.find(i[0])
        while pos >= 0:
            original_txt = original_txt[:pos+1] + i[2] + original_txt[pos+3:]
            pos += 1
            pos = txt.find(i[0], pos)

    for i in repl_list:
        original_txt = original_txt.replace(i[2], i[1])

    return original_txt


dictionary["serbian"]["dicts"]["biljke"] = {}

with open("biljke_lat.txt", "r", encoding="utf-8") as file:
    txt = file.read()


txt_list = [x.strip() for x in txt.split("\n")]


new_topic = ""
topic = ""
for line in txt_list:
    line = line.replace("–", "-").strip()

    if not line:
        continue

    if line.find("|") >= 0:
        line = line.replace("|", "").replace(".", "").strip().upper()
    
    new_topic = ""
    if line == line.upper():
        new_topic = line
        body = ""
    else:
        body = " " + line
    
    if new_topic:
        topic = new_topic
        if topic in dictionary["serbian"]["dicts"]["biljke"]:
            dictionary["serbian"]["dicts"]["biljke"][topic]["text"] += "\n\n"
            dictionary["serbian"]["dicts"]["biljke"][topic]["text"] += body.strip()
        else:
            dictionary["serbian"]["dicts"]["biljke"][topic] = {}
            dictionary["serbian"]["dicts"]["biljke"][topic]["text"] = body.strip()
            dictionary["serbian"]["dicts"]["biljke"][topic]["links"] = []
        continue    

    if dictionary["serbian"]["dicts"]["biljke"][topic]["text"]:
        dictionary["serbian"]["dicts"]["biljke"][topic]["text"] += f" {body}"
    else:
        dictionary["serbian"]["dicts"]["biljke"][topic]["text"] += body

lnks = 0
for item in dictionary["serbian"]["dicts"]["biljke"]:
    txt = dictionary["serbian"]["dicts"]["biljke"][item]["text"]

    links = bilj_get_links(txt, item)
    dictionary["serbian"]["dicts"]["biljke"][item]["links"] = links
    lnks += len(links)
    txt = bilj_fix_text(txt, item)
    dictionary["serbian"]["dicts"]["biljke"][item]["text"] = txt


print (f'biljke finnished.  LEN = {len(dictionary["serbian"]["dicts"]["biljke"])},  Links: {lnks}')




# IT Recnik

def it_get_links(txt: str, curr_item: str) -> list:
    txt += " "
    for i in END_OF_TOPIC:
        txt = txt.replace(i, " ")

    links = []
    for i in dictionary["serbian"]["dicts"]["it"]:
        if i == curr_item or len(i) < 2:
            continue
        search = f" {i} "
        if txt.find(search) >= 0:
            if i not in links:
                links.append(i)
    return links


dictionary["serbian"]["dicts"]["it"] = {}

with open("it_recnik.txt", "r", encoding="utf-8") as file:
    txt = file.read()


txt_list = [x.strip() for x in txt.split("\n")]


new_topic = ""
topic = ""

for line in txt_list:
    line = line.replace("–", "-").strip()

    if not line:
        topic = ""
        continue

    if not topic:
        new_topic = line
        body = ""
    else:
        body = " " + line


    if new_topic:
        topic = new_topic
        if topic in dictionary["serbian"]["dicts"]["it"]:
            dictionary["serbian"]["dicts"]["it"][topic]["text"] += "\n\n"
            dictionary["serbian"]["dicts"]["it"][topic]["text"] += body.strip()
        else:
            dictionary["serbian"]["dicts"]["it"][topic] = {}
            dictionary["serbian"]["dicts"]["it"][topic]["text"] = body.strip()
            dictionary["serbian"]["dicts"]["it"][topic]["links"] = []
        continue    

    if dictionary["serbian"]["dicts"]["it"][topic]["text"]:
        dictionary["serbian"]["dicts"]["it"][topic]["text"] += f" {body}"
    else:
        dictionary["serbian"]["dicts"]["it"][topic]["text"] += body

lnks = 0
for item in dictionary["serbian"]["dicts"]["it"]:
    txt = dictionary["serbian"]["dicts"]["it"][item]["text"]

    links = it_get_links(txt, item)
    dictionary["serbian"]["dicts"]["it"][item]["links"] = links
    lnks += len(links)


print (f'it finnished.  LEN = {len(dictionary["serbian"]["dicts"]["it"])},  Links: {lnks}')




# Bokeljski Recnik

dictionary["serbian"]["dicts"]["bokeljski"] = {}

with open("bokeljski_recnik.txt", "r", encoding="utf-8") as file:
    txt = file.read()

txt_list = [x.strip() for x in txt.split("\n") if x.strip() != ""]

new_topic = ""
topic = ""
for line in txt_list:
    line = line.replace("–", "-")

    if line.find("-") >= 0:
        new_topic = line[:line.find("-")].strip()
        body = line[line.find("-") + 1:].strip()
    else:
        body = line
    
    new_topic = new_topic.lower()
    body = body.lower()
    
    if new_topic:
        topic = new_topic
        new_topic = ""
        dictionary["serbian"]["dicts"]["bokeljski"][topic] = {}
        dictionary["serbian"]["dicts"]["bokeljski"][topic]["text"] = body
        dictionary["serbian"]["dicts"]["bokeljski"][topic]["links"] = []
        continue

    dictionary["serbian"]["dicts"]["bokeljski"][topic]["text"] += f' {body}'

print (f'Bokeljski recnik finnished: {len(dictionary["serbian"]["dicts"]["bokeljski"])} items.')



# Engelsko - Hrvatski Recnik Bankarstva, Osiguranja i ostalih finansijskih usluga

doc = Document("ENGLESKI_RECNIK_bankarstva_osiguranja_i.docx")

dictionary["serbian"]["dicts"]["bank"] = {}

start = False
bold_text = ""
body = ""
topic_name = ""
new_topic = ""
for paragraph in doc.paragraphs:
    par_text = paragraph.text.strip()

    if par_text.find("aaa") >= 0:
        start = True
    
    if par_text.find("162\tMinistarstvo vanjskih poslova i europskih integracija") >= 0:
        start = False

    if not start:
        continue

    if len(par_text) < 3:
        continue

    new_topic = ""
    body = ""
    next_topic = False
    topic_list = []
    abbreviation_mode = False
    abbreviation_pos = 0

    for run in paragraph.runs:
        if run.text.strip() == "➔":
            abbreviation_mode = True
            abbreviation_pos = 0
            next_topic = False
            if not new_topic:
                new_topic = body
            body = ""
            new_topic = new_topic.strip()
            continue

        if abbreviation_mode:
            body += run.text
            if run.text.strip() == "":
                continue

            for i in range(abbreviation_pos, len(new_topic)):
                if new_topic[i].lower() == run.text[0].strip().lower():
                    abbreviation_pos = i + 1
                    break
            else:
                abbreviation_pos = 1000000
            
            if abbreviation_pos >= len(new_topic):
                topic_list.append([new_topic, body.strip()])
                new_topic = ""
                body = ""
                abbreviation_mode = False
                continue
            else:
                continue
        
        if run.bold:
            if run.text == "Bilj.:":
                body = run.text
                new_topic = ""
            else:
                if next_topic:
                    if new_topic.strip() or body.strip():
                        topic_list.append([new_topic.strip(), body.strip()])
                    new_topic = ""
                    body = ""
                    next_topic = False
                if new_topic and run.text:
                    if new_topic[-1] != " " and run.text[0] != " ":
                        new_topic += " "
                new_topic += run.text
        else:
            body += run.text
            if run.text.strip():
                next_topic = True

    if new_topic.strip() or body.strip():
        topic_list.append([new_topic.strip(), body.strip()])

    for i in topic_list:
        new_topic = i[0]
        if new_topic.strip() == ",":
            new_topic = ""
        body = i[1]

        if new_topic:
            topic_name = new_topic
            new_topic = ""
            if topic_name in dictionary["serbian"]["dicts"]["bank"]:

                dictionary["serbian"]["dicts"]["bank"][topic_name]["text"] += "\n\n"
                print (".", end="")
            else:
                dictionary["serbian"]["dicts"]["bank"][topic_name] = {}
                dictionary["serbian"]["dicts"]["bank"][topic_name]["text"] = ""
                dictionary["serbian"]["dicts"]["bank"][topic_name]["links"] = []
                new_topic = ""
                # continue

        if topic_name:
            if dictionary["serbian"]["dicts"]["bank"][topic_name]["text"]:
                if body.strip() != dictionary["serbian"]["dicts"]["bank"][topic_name]["text"].strip():
                    dictionary["serbian"]["dicts"]["bank"][topic_name]["text"] += f"\n{body}"
            else:
                dictionary["serbian"]["dicts"]["bank"][topic_name]["text"] += body


print ()
print (f'GLOSAR Bankarstva: {len(dictionary["serbian"]["dicts"]["bank"])} items.')



# Recnik Google i Microsoft Termina

dictionary["serbian"]["dicts"]["google&ms"] = {}

with open("recnik_google_i_microsoft_termina.txt", "r", encoding="utf-8") as file:
    txt = file.read()

txt_list = [x.strip() for x in txt.split("\n") if x.strip() != ""]

new_topic = ""
topic = ""
has_new_topic = False
for line in txt_list:
    line = line.strip()

    if len(line) < 3:
        continue
    
    if line.find("plug") >= 0:
        print ("Found")

    if line.find(" – ") >= 0:
        has_new_topic = True
    else:
        has_new_topic = False

    line = line.replace("–", "-")
    
    if has_new_topic:
        new_topic = line[:line.find(" - ")].strip()
        body = line[line.find(" - ") + 3:].strip()
    else:
        body = line
    
    new_topic = new_topic.strip()
    body = body.strip()
    
    if new_topic:
        topic = new_topic
        new_topic = ""
        dictionary["serbian"]["dicts"]["google&ms"][topic] = {}
        dictionary["serbian"]["dicts"]["google&ms"][topic]["text"] = body
        dictionary["serbian"]["dicts"]["google&ms"][topic]["links"] = []
        continue

    dictionary["serbian"]["dicts"]["google&ms"][topic]["text"] += f' {body}'

print (f'Google&MS recnik finnished: {len(dictionary["serbian"]["dicts"]["google&ms"])} items.')




# Frazeoloski recnik Djordje Otasevic

dictionary["serbian"]["dicts"]["fraze"] = {}

with open("frazeoloski_recnik_djordje_otasevic.txt", "r", encoding="utf-8") as file:
    txt = file.read()

txt_list = [x.strip() for x in txt.split("\n") if x.strip() != ""]

new_topic = ""
topic = ""
has_new_topic = False
for line in txt_list:
    line = line.strip()

    if not line:
        continue

    line = cirilica_u_latinicu(line)

    line_up = line.replace("Nj", "NJ")
    line_up = line_up.replace("Lj", "LJ")
    
    if line_up == line_up.upper():
        has_new_topic = True
    else:
        has_new_topic = False

    line = line.replace("–", "-")
    
    if has_new_topic:
        new_topic = line
        body = ""
    else:
        body = line
    
    new_topic = new_topic.strip()
    body = body.strip()
    
    if new_topic:
        topic = new_topic
        new_topic = ""
        dictionary["serbian"]["dicts"]["fraze"][topic] = {}
        dictionary["serbian"]["dicts"]["fraze"][topic]["text"] = body
        dictionary["serbian"]["dicts"]["fraze"][topic]["links"] = []
        continue

    dictionary["serbian"]["dicts"]["fraze"][topic]["text"] += f' {body}'

print (f'Frazeoloski recnik finnished: {len(dictionary["serbian"]["dicts"]["fraze"])} items.')





# Ekonomski recnik Momir Jaksic

dictionary["serbian"]["dicts"]["ekonom"] = {}

with open("ekonomski_recnik_momir_jaksic.txt", "r", encoding="utf-8") as file:
    txt = file.read()

txt_list = [x.strip() for x in txt.split("\n") if x.strip() != ""]

new_topic = ""
topic = ""
has_new_topic = False
for line in txt_list:
    line = line.strip()

    if not line:
        continue

    line = line.replace("–", "-")

    if line.find("-") >= 0:
        has_new_topic = True
    else:
        has_new_topic = False
    
    if has_new_topic:
        new_topic = line[:line.find("-")]
        body = line[line.find("-") + 1:]
    else:
        body = line
    
    new_topic = new_topic.strip()
    body = body.strip()
    
    if new_topic:
        topic = new_topic
        new_topic = ""
        dictionary["serbian"]["dicts"]["ekonom"][topic] = {}
        dictionary["serbian"]["dicts"]["ekonom"][topic]["text"] = body
        dictionary["serbian"]["dicts"]["ekonom"][topic]["links"] = []
        continue

    dictionary["serbian"]["dicts"]["ekonom"][topic]["text"] += f' {body}'

print (f'Ekonomski recnik Momir Jaksic finnished: {len(dictionary["serbian"]["dicts"]["ekonom"])} items.')




# Recnik Ekonomskih Pojmova 

def ekonom2_fix_topic_name(topic_name: str) -> str:
    find_num = "0123456789"
    pos = 0
    has_num = False
    while True:
        pos = topic_name.find("(", pos)
        if pos == -1:
            break
        pos2 = topic_name.find(")", pos)
        if pos2 == -1:
            break

        inside_b = topic_name[pos:pos2]
        has_num = False
        for i in inside_b:
            if i in find_num:
                has_num = True
                break
        
        if has_num:
            part1 = topic_name[:pos]
            if pos2 != len(topic_name) - 1:
                part2 = topic_name[pos2 + 1:]
            else:
                part2 = ""
            
            topic_name = part1.strip() + " " + part2.strip()
        
        pos += 1
    
    return topic_name


dictionary["serbian"]["dicts"]["ekonom2"] = {}

with open("recnik_ekonomskih_pojmova.txt", "r", encoding="utf-8") as file:
    txt = file.read()

txt_list = [x.strip() for x in txt.split("\n") if x.strip() != ""]

new_topic = ""
topic = ""
has_new_topic = False
for line in txt_list:
    line = line.strip()

    if not line:
        continue

    line = line.replace("–", "-")

    if line.find(":") >= 0:
        has_new_topic = True
    else:
        has_new_topic = False
    
    if has_new_topic:
        new_topic = line[:line.find(":")]
        body = line[line.find(":") + 1:]
    else:
        body = line
    
    new_topic = new_topic.strip()
    new_topic = ekonom2_fix_topic_name(new_topic)
    body = body.replace("- ", "")    
    body = body.strip()
    
    if new_topic:
        topic = new_topic
        new_topic = ""
        dictionary["serbian"]["dicts"]["ekonom2"][topic] = {}
        dictionary["serbian"]["dicts"]["ekonom2"][topic]["text"] = body
        dictionary["serbian"]["dicts"]["ekonom2"][topic]["links"] = []
        continue

    dictionary["serbian"]["dicts"]["ekonom2"][topic]["text"] += f' {body}'

print (f'Recnik Ekonomskih Pojmova finnished: {len(dictionary["serbian"]["dicts"]["ekonom2"])} items.')



# GLOSAR PROCESNIH POJMOVA


dictionary["serbian"]["dicts"]["proces"] = {}

with open("GLOSAR_PROCESNIH_POJMOVA.txt", "r", encoding="utf-8") as file:
    txt = file.read()

txt_list = [x.strip() for x in txt.split("\n") if x.strip() != ""]

new_topic = ""
topic = ""
has_new_topic = False
for line in txt_list:
    line = line.strip()

    if not line:
        continue

    line = line.replace("–", "-")

    if line.find("•") >= 0:
        has_new_topic = True
    else:
        has_new_topic = False
    
    line = line.replace("•", "").strip()

    if has_new_topic:
        new_topic = line[:line.find("-")]
        body = line[line.find("-") + 1:]
    else:
        body = line
    
    new_topic = new_topic.strip()
    body = body.strip()
    
    if new_topic:
        topic = new_topic
        new_topic = ""
        dictionary["serbian"]["dicts"]["proces"][topic] = {}
        dictionary["serbian"]["dicts"]["proces"][topic]["text"] = body
        dictionary["serbian"]["dicts"]["proces"][topic]["links"] = []
        continue

    dictionary["serbian"]["dicts"]["proces"][topic]["text"] += f' {body}'

print (f'GLOSAR PROCESNIH POJMOVA finnished: {len(dictionary["serbian"]["dicts"]["ekonom2"])} items.')



# Recnik Ekonomskih Pojmova DODATAK - LEKSIKON POJMOVA za pripremu prijemnog ispita za MASTER STUDIJE 

dictionary["serbian"]["dicts"]["ekonom2"] = {}

with open("LEKSIKONPOJMOVAzapripremuprijemnogispitazaMASTERSTUDIJE.txt", "r", encoding="utf-8") as file:
    txt = file.read()

txt_list = [x.strip() for x in txt.split("\n") if x.strip() != ""]

new_topic = ""
topic = ""
has_new_topic = False
for line in txt_list:
    line = line.strip()

    if not line:
        continue

    line = line.replace("–", "-")

    if line == line.upper():
        has_new_topic = True
    else:
        has_new_topic = False
    
    if has_new_topic:
        new_topic = line
        body = ""
    else:
        body = line
    
    new_topic = new_topic.strip()
    body = body.strip()
    
    if new_topic:
        topic = new_topic
        new_topic = ""
        if topic in dictionary["serbian"]["dicts"]["ekonom2"]:
           continue 
        dictionary["serbian"]["dicts"]["ekonom2"][topic] = {}
        dictionary["serbian"]["dicts"]["ekonom2"][topic]["text"] = body
        dictionary["serbian"]["dicts"]["ekonom2"][topic]["links"] = []
        continue

    dictionary["serbian"]["dicts"]["ekonom2"][topic]["text"] += f' {body}'

print (f'Recnik Ekonomskih Pojmova DODATAK !!! finnished: {len(dictionary["serbian"]["dicts"]["ekonom2"])} items.')




# Filozofijski Rjecnik Vladimira Filipovica

doc = Document("kupdf.net_filozofski-recnik.docx")

dictionary["serbian"]["dicts"]["filoz2"] = {}

start = False
body = ""
topic_name = ""
header = False
new_topic = ""
remove_from_topic_name = ["-","(lat","(grč"]
for paragraph in doc.paragraphs:
    par_text = paragraph.text.strip()

    if par_text.find("a = a") >= 0:
        print ("FOUND")
    
    if par_text == "A":
        start = True
    
    if par_text.find("Grupa autora u redakciji Vladimira Fil") >= 0 and start:
        start = False

    if not start or len(par_text) < 4:
        continue

    if len(par_text) < 3:
        continue

    if par_text.count("\t") > 1 and len(par_text) > len(topic_name):
        if topic_name.lower() == par_text[:len(topic_name)].lower():
            header = True
            continue
        if header:
            header = False
            continue

    new_topic = ""
    body = ""

    for run in paragraph.runs:
        if run.bold:
            new_topic += run.text
        else:
            body += run.text

    if new_topic != par_text[:len(new_topic)]:
        body = par_text
        new_topic = ""
    if new_topic == "sufi	320	sumaniya" or new_topic == "sumaniya	321	supstancija":
        print ("PASSED topic : ", new_topic)
        new_topic = ""
        continue
    
    new_topic = new_topic.replace("\xad ", "")
    body = body.replace("\xad ", "")
    
    if body.find("\t") >= 0:
        body = body[:body.rfind("\t")]

    new_topic = new_topic.replace("–", "-")
    new_topic = new_topic.replace("\t", "")
    new_topic = new_topic.strip()
    if new_topic:
        for remove_item in remove_from_topic_name:
            if len(new_topic) >= len(remove_item):
                if new_topic[-len(remove_item)] == remove_item:
                    new_topic = new_topic[:-len(remove_item)]
        while True:
            if new_topic.find("  ") >= 0:
                new_topic = new_topic.replace("  ", " ")
            else:
                break
    
    body = body.strip()

    while True:
        if not new_topic:
            break
        if new_topic[-1] in ",. :=":
            new_topic = new_topic[:-1]
        else:
            break

    if new_topic:
        topic_name = new_topic
        new_topic = ""
        if topic_name in dictionary["serbian"]["dicts"]["filoz2"]:
            dictionary["serbian"]["dicts"]["filoz2"][topic_name]["text"] += "\n\n\n"
            print (f"Duplicate Entry : {topic_name}")
        else:
            dictionary["serbian"]["dicts"]["filoz2"][topic_name] = {}
            dictionary["serbian"]["dicts"]["filoz2"][topic_name]["text"] = ""
            dictionary["serbian"]["dicts"]["filoz2"][topic_name]["links"] = []
            new_topic = ""
            # continue

    if dictionary["serbian"]["dicts"]["filoz2"][topic_name]["text"]:
        dictionary["serbian"]["dicts"]["filoz2"][topic_name]["text"] += f" {body}"
    else:
        dictionary["serbian"]["dicts"]["filoz2"][topic_name]["text"] += body


print ()
print (f'Filozofijski Rjecnik Vladimira Filipovica: {len(dictionary["serbian"]["dicts"]["filoz2"])} items.')




# 1999_leksikon_srpskog_srednjeg_veka_prir_sima_cirkovic_i_rade_mihaljcic

def srp_srednji_vek_is_valid_title(title: str) -> bool:
    a = "QWERTYUIOPLKJHGFDSAZXCVBNMŠĐŽĆČ"
    illegal = ["1A", "2A", "3A", "4A", "5A", "6A", "7A", "8A", "9A",
               "1B", "2B", "3B", "4B", "5B", "6B", "7B", "8B", "9B",
               "1C", "2C", "3C", "4C", "5C", "6C", "7C", "8C", "9C",
               "1D", "2D", "3D", "4D", "5D", "6D", "7D", "8D", "9D",
               "1E", "2E", "3E", "4E", "5E", "6E", "7E", "8E", "9E",
               "1F", "2F", "3F", "4F", "5F", "6F", "7F", "8F", "9F",
               "V1"]
    for i in illegal:
        if i in title:
            print (f"Izbacen {title}")
            return False
    pos = 0
    pos2 = 0
    while True:
        pos = title.find("(")
        if pos >=0:
            pos2 = title.find(")")
            if pos2 > pos:
                title = title[:pos] + title[pos2+1:]
            else:
                break
        else:
            break
    if not title.strip():
        return True
    
    for i in title:
        if i in a:
            return True
    return False

def srp_srednji_vek_is_final_check(title: str) -> bool:
    pos = 0
    pos2 = 0
    while True:
        pos = title.find("(")
        if pos >=0:
            pos2 = title.find(")")
            if pos2 > pos:
                title = title[:pos] + title[pos2+1:]
            else:
                break
        else:
            break

    if title.find("^") >= 0:
        return False
    
    if title != title.upper():
        return False
    return True

def srp_srednji_vek_is_upper(text: str, replace_v: bool = False) -> tuple:
    if replace_v:
        text = text.replace(" в. ", " - в. ")
        text = text.replace(" v. ", " - v. ")

    a = "QWERTYUIOPLKJHGFDSAZXCVBNMŠĐŽĆČ ,.()"

    pos = 0
    pos2 = 0
    while True:
        pos = text.find("(")
        if pos >=0:
            pos2 = text.find(")")
            if pos2 > pos:
                text = text[:pos].strip() + " " + text[pos2+1:]
            else:
                break
        else:
            break

    text = text.replace("–", "-")
    pos = text.find("-")
    if pos == -1:
        return None
    
    text = cirilica_u_latinicu(text)
    bd = text[text.find("-")+1:].strip()

    text = text.replace("Nj", "NJ").replace("Lj", "LJ")

    text = text[:text.find("-")].strip()

    is_valid = True
    for i in text:
        if i not in a:
            is_valid = False
            break
    
    if not is_valid:
        return None
    
    return (text, bd)


doc = Document("1999_leksikon_srpskog_srednjeg_veka_prir_sima_cirkovic_i_rade_mihaljcic.docx")

dictionary["serbian"]["dicts"]["srp_srednji_vek"] = {}

start = False
body = ""
topic_name = ""
header = False
new_topic = ""
for paragraph in doc.paragraphs:
    par_text = paragraph.text.strip()

    if par_text.find("Сима Ћирковић") >= 0:
        print ("Found")

    if par_text == "I ОПШТА ДЕЛА":
        start = False

    par_text = cirilica_u_latinicu(par_text)

    if par_text == "A":
        start = True
    
    if not start:
        continue

    if len(par_text) < 3:
        continue

    if par_text == par_text.upper():
        continue

    new_topic = ""
    body = ""

    if cirilica_u_latinicu(par_text).find("Baštinska vojska") >= 0 or cirilica_u_latinicu(par_text).find("AMIRA") >= 0:
        print ("Found CRKVE")

    for run in paragraph.runs:
        if run.bold or run.text == "ЧЕЉАДИН ":
            if not body.strip():
                if srp_srednji_vek_is_valid_title(cirilica_u_latinicu(run.text)):
                    if run.text.rfind("-") > 0:
                        if srp_srednji_vek_is_valid_title(run.text[run.text.rfind("-")+1:]):
                            new_topic += run.text.strip() + " "
                        else:
                            new_topic += run.text[:run.text.rfind("-")].strip() + " "
                            body += run.text[run.text.rfind("-")+1:]
                    else:
                        new_topic += run.text.strip() + " "
                else:
                    body += run.text
            else:
                body += run.text
        else:
            if len(run.text) > 3:
                if body.strip() == "" and new_topic.strip() == "":
                    replace_v = False
                    if not srp_srednji_vek_is_upper(par_text):
                        replace_v = True
                    result = srp_srednji_vek_is_upper(run.text, replace_v=replace_v)
                    if result:
                        new_topic = result[0] + " "
                        body += result[1]
                    else:
                        if srp_srednji_vek_is_upper(par_text, replace_v=replace_v):
                            if len(run.text.strip()) > 2 and run.text == run.text.upper():
                                new_topic = run.text.strip() + " "
                            else:
                                body += run.text
                        else:
                            body += run.text
                else:
                    body += run.text
            else:
                body += run.text

    new_topic = new_topic.replace("–", "-")
    new_topic = new_topic.replace("\t", "")

    pos = 0
    while True:
        if new_topic.find("(", pos) >= 0:
            pos = new_topic.find("(", pos)
        else:
            break
        if new_topic.find(")", pos) >= 0:
            new_topic = new_topic[:pos] + new_topic[new_topic.find(")", pos) + 1:]
        else:
            new_topic = new_topic[:pos]
            break

    new_topic = cirilica_u_latinicu(new_topic)
    body = cirilica_u_latinicu(body)

    if new_topic in ["DŠVATI veliki Lmirl Pdkizitk;", "NIKA VB DOMB SVOI?, ILI "]:
        new_topic = ""

    if new_topic.find("ČIZME") >= 0:
        print ("FOUND")

    new_topic = new_topic.replace("Nj", "NJ").replace("Lj", "LJ")
    if not srp_srednji_vek_is_final_check(new_topic):
        new_topic = ""
        body = cirilica_u_latinicu(par_text)

    new_topic = new_topic.replace("\t", " ")
    while True:
        if new_topic.find("  ") >= 0:
            new_topic = new_topic.replace("  ", " ")
        else:
            break

    new_topic = new_topic.strip("- ")
    pos = 0
    while True:
        pos = body.find("-", pos)
        if pos == -1:
            break

        replace_v = True
        if pos > 0:
            if body[pos-1] in "0123456789":
                replace_v = False
        if pos < len(body) - 1:
            if body[pos+1] in "0123456789":
                replace_v = False
        
        if replace_v:
            body = body[:pos] + body[pos+1:]
        else:
            pos += 1

    body = body.replace("LITERATURA:", "\n\nLITERATURA:").replace("IZVORI:", "\n\nIZVORI:").strip()

    if new_topic:
        topic_name = new_topic
        new_topic = ""
        if topic_name in dictionary["serbian"]["dicts"]["srp_srednji_vek"]:
            dictionary["serbian"]["dicts"]["srp_srednji_vek"][topic_name]["text"] += "\n\n\n"
            print (f"Duplicate Entry : {topic_name}")
        else:
            dictionary["serbian"]["dicts"]["srp_srednji_vek"][topic_name] = {}
            dictionary["serbian"]["dicts"]["srp_srednji_vek"][topic_name]["text"] = ""
            dictionary["serbian"]["dicts"]["srp_srednji_vek"][topic_name]["links"] = []
            new_topic = ""
            # continue

    if dictionary["serbian"]["dicts"]["srp_srednji_vek"][topic_name]["text"]:
        dictionary["serbian"]["dicts"]["srp_srednji_vek"][topic_name]["text"] += f" {body}"
    else:
        dictionary["serbian"]["dicts"]["srp_srednji_vek"][topic_name]["text"] += body

dictionary["serbian"]["dicts"]["srp_srednji_vek"]["KRIPTA"] = dictionary["serbian"]["dicts"]["srp_srednji_vek"].pop("KRIPTA )")

remove_items = ["PRO", ",", "U", "IH", "RI", "S.", "BACJ", "STIŠL"]
for i in remove_items:
    dictionary["serbian"]["dicts"]["srp_srednji_vek"].pop(i)

print ()
print (f'1999_leksikon_srpskog_srednjeg_veka_prir_sima_cirkovic_i_rade_mihaljcic: {len(dictionary["serbian"]["dicts"]["srp_srednji_vek"])} items.')




# Sindikalni Leksikon

def sind_get_title_body(text: str) -> tuple:
    pos = text.find("-")
    if pos == -1:
        return ("", text)
    
    title = text[:pos].strip()
    body = text[pos+1:].strip()

    if title != title.upper():
        return ("", text)

    return (title, body)


dictionary["serbian"]["dicts"]["sind"] = {}

with open("24sindikalni_leksikon.txt", "r", encoding="utf-8") as file:
    text = file.read()

text_lines = [x.strip() for x in text.split("\n") if x.strip() != ""]

text = ""
for line in text_lines:
    if len(line) < 4 and line.find("-") == -1:
        continue

    if line.find("slidepdf.com") >= 0:
        continue
    if line.find(" LEKSIKON") >=0:
        continue
    if line.find("TERMINOLOŠKI REČNIK") >=0:
        continue

    text += line + " "

text = text.replace("", "")
text = text.replace("Nj", "NJ")
text = text.replace("Lj", "LJ")
text = text.replace("Dž", "DŽ")
if text[-1] != ".":
    text += "."

pos = -1
text_lines = []
topic = ""
new_topic = False
while True:
    pos2 = text.find(".", pos + 1)
    if pos2 == -1:
        topic += text[pos+1:]
        text_lines.append(topic)
        break

    sentence = text[pos+1:pos2+1]
    pos = pos2
    sep_pos = sentence.find("-")
    if sep_pos == -1:
        if sentence.find("/") >= 0:
            sep_pos = sentence.find("/")
            sentence = sentence[:sep_pos] + " - " + sentence[sep_pos:]
            sep_pos = sentence.find("-")

    else:
        if sentence.find("/") < sep_pos and sentence.find("/") > 0:
            sep_pos = sentence.find("/")
            sentence = sentence[:sep_pos] + " - " + sentence[sep_pos:]
            sep_pos = sentence.find("-")

    if sep_pos == -1:
        topic += sentence
        continue

    pref = sentence[:sep_pos].strip()
    if len(pref) < 3:
        topic += sentence
        continue
    if pref != pref.upper():
        topic += sentence
        continue

    if topic:
        text_lines.append(topic)
    
    topic = sentence

for topic in text_lines:
    pos = topic.find("-")
    if pos == -1:
        print ("Error.  ", topic)
        continue

    new_topic = topic[:pos].strip()
    body = topic[pos+1:].strip()

    new_topic = new_topic.replace("-   ", "")
    new_topic = new_topic.replace("-  ", "")
    new_topic = new_topic.replace("- ", "")
    body = body.replace("-   ", "")
    body = body.replace("-  ", "")
    body = body.replace("- ", "")

    if new_topic:
        topic_name = new_topic
        new_topic = ""
        if topic_name in dictionary["serbian"]["dicts"]["sind"]:
            dictionary["serbian"]["dicts"]["sind"][topic_name]["text"] += "\n\n\n"
            print (f"Duplicate Entry : {topic_name}")
        else:
            dictionary["serbian"]["dicts"]["sind"][topic_name] = {}
            dictionary["serbian"]["dicts"]["sind"][topic_name]["text"] = body
            dictionary["serbian"]["dicts"]["sind"][topic_name]["links"] = []
            new_topic = ""

dictionary["serbian"]["dicts"]["sind"].pop("000")

print ()
print (f'Sindikalni Leksikon: {len(dictionary["serbian"]["dicts"]["sind"])} items.')





# Leksikon Temeljnih Religijskih Pojmova

padding = [
    "1. kršćanski",
    "2. kršćanski",
    "3. kršćanski",
    "kršćanski",
    "1. židovski",
    "2. židovski",
    "3. židovski",
    "židovski",
    "1. islamski",
    "2. islamski",
    "3. islamski",
    "islamski"
]

def religije_valid_title(text: str, paragraph) -> tuple:
    excl = [
        "L. HAGE'MANN 3. islamski",
        "kršćanski",
        "islamski",
        "židovski"
    ]

    for i in excl:
        if text.find(i) >= 0:
            print (f"Rejected ........ EXCL List .......... {text}")
            return ("", text)
    
    if paragraph.alignment == 2:
        print (f"Rejected ........ Align=2 .......... {text}")
        return ("", text)
    
    if text.find(":") >= 0:
        print (f"Rejected ........ Found (:) .......... {text}")
        return ("", text)
    
    if len(text.strip()) < 3:
        print (f"Rejected ........ LEN < 2 .......... {text}")
        return ("", text)
    
    text = text.replace("»", ">")
    text = text.replace("«", "<")
    text = text.replace("–", "-")
    text = text.replace("►", ">")

    if text.find(">") >= 0:
        text_split = text.split(">")
        if len(text_split) < 2:
            print (f"Rejected ........ SPLIT LEN < 2 .......... {text}")
            return ("", text)
        if text_split[0] != text_split[0].upper():
            print (f"Rejected ........ LEFT <> LEFT.UP .......... {text}")
            return ("", text)
        text = " > ".join(text_split)
        print (f"Accepted ..... Topic ({text_split[0].strip(' -')}) - Body :  -> {text}")
        return (text_split[0].strip(" -"), text)

    if paragraph.text[:len(text)] != text:
        print (f"Rejected ........ NOT BEGINNING .......... {text}")
        return ("", text)
    
    print (f"New Topic: {text}")
    return (text, "")

doc = Document("a_khoury_leksikon_temeljnih_religijskih_pojmova_zidovstvo_hriscanstvo_islampdf.docx")

dictionary["serbian"]["dicts"]["religije"] = {}

start = False
body = ""
topic_name = ""
header = False
new_topic = ""
counter = 0
for paragraph in doc.paragraphs:
    par_text = paragraph.text.strip()

    if par_text == "ABRAHAM":
        start = True

    if par_text == "SlNOPTIČKA KRONOLOŠKA TABLICA":
        start = False

    if not start:
        continue

    if len(par_text) < 3:
        continue

    new_topic = ""
    body = ""
    for run in paragraph.runs:
        if run.bold:
            if not new_topic:
                result = religije_valid_title(run.text, paragraph)
                new_topic = result[0]
                body += result[1]
            else:
                print (f"Already has topic ! ! !  : {run.text}")
                body += run.text
        else:
            body += run.text

    new_topic = new_topic.strip()
    body = body.strip()
    counter += 1
    if counter % 10 == 0:
        print ("---------------------------------------------------------------------------")

    for i in padding:
        if par_text[:len(i)] == i and not new_topic:
            body = "\n\n" + body

    if new_topic:
        topic_name = new_topic
        new_topic = ""
        if topic_name in dictionary["serbian"]["dicts"]["religije"]:
            dictionary["serbian"]["dicts"]["religije"][topic_name]["text"] += "\n\n\n"
            print (f"Duplicate Entry : {topic_name}")
        else:
            print (topic_name)
            dictionary["serbian"]["dicts"]["religije"][topic_name] = {}
            dictionary["serbian"]["dicts"]["religije"][topic_name]["text"] = ""
            dictionary["serbian"]["dicts"]["religije"][topic_name]["links"] = []
            new_topic = ""
            # continue

    if dictionary["serbian"]["dicts"]["religije"][topic_name]["text"]:
        dictionary["serbian"]["dicts"]["religije"][topic_name]["text"] += f" {body}"
    else:
        dictionary["serbian"]["dicts"]["religije"][topic_name]["text"] += body

remove_items = [
    "DUAKONUA LJUBAV PREMA BLIŽNJEMU",
    "VJEROISPOVIJEST"
]

for i in dictionary["serbian"]["dicts"]["religije"]:
    if not dictionary["serbian"]["dicts"]["religije"][i]["text"].strip():
        if i not in remove_items:
            remove_items.append(i)
    
for i in remove_items:
    dictionary["serbian"]["dicts"]["religije"].pop(i)


add_items = [
    ["DOGMA", " VJEROISPOVIJEST"]
]

print ()
print (f'Leksikon religijskih pojmova: {len(dictionary["serbian"]["dicts"]["religije"])} items.')




# Recnik svetske mitologije Artur Koterel


doc = Document("artur_koterel_recnik_svetske_mitologije.docx")

dictionary["serbian"]["dicts"]["svet_mit"] = {}

start = False
body = ""
topic_name = ""
header = False
new_topic = ""
counter = 0
for paragraph in doc.paragraphs:
    par_text = paragraph.text.strip()

    if par_text in ["Adapa", "Agni", "Adžisit", "Afrodita", "Agve", "Abasi", "Ahoeitu"]:
        start = True

    if par_text in ["JUŽNA I SREDNJA AZIJA", "Sibir Mongolija Kina Japan\njugoistočna Azija", "EVROPA", "AMERIKA", "AFRIKA", "OKEANIJA", "LITERATURA"]:
        start = False

    if not start:
        continue

    if len(par_text) < 2:
        continue

    new_topic = ""
    body = ""


    if par_text == "Bal":
        print ()

    if par_text.count(" ") > 3:
        body = par_text
    else:
        new_topic = par_text

    header = False
    # for run in paragraph.runs:
    #     if run.font.size is None:
    #         header = False
        # if run.bold:
        #     new_topic = run.text
        # else:
        #     body += run.text

    pos = par_text.find(" ")
    if pos >= 0:
        try:
            _ = int(par_text.split(" ")[0])
            if par_text.find("laićanski bog mora.") == -1 and par_text.find("judi od davnina čuvaju predmete za koje veruju da su zbog svojih ranijih veza ispunjeni vr") == -1:
                header = True
        except:
            _ = 0

    if par_text in ["Rcčmksvelske mitologije 1 13"]:
        header = True

    if header:
        print (f"Rejected: {par_text}")
        continue

    new_topic = new_topic.strip()
    body = body.strip()

    if new_topic:
        topic_name = new_topic
        new_topic = ""
        if topic_name in dictionary["serbian"]["dicts"]["svet_mit"]:
            dictionary["serbian"]["dicts"]["svet_mit"][topic_name]["text"] += "\n\n\n"
            print (f"Duplicate Entry : {topic_name}")
        else:
            print (topic_name)
            dictionary["serbian"]["dicts"]["svet_mit"][topic_name] = {}
            dictionary["serbian"]["dicts"]["svet_mit"][topic_name]["text"] = ""
            dictionary["serbian"]["dicts"]["svet_mit"][topic_name]["links"] = []
            new_topic = ""
            # continue

    if dictionary["serbian"]["dicts"]["svet_mit"][topic_name]["text"]:
        dictionary["serbian"]["dicts"]["svet_mit"][topic_name]["text"] += f" {body}"
    else:
        dictionary["serbian"]["dicts"]["svet_mit"][topic_name]["text"] += body

remove_items = [
    "o ’"
]


print ()
print (f'Recnik svetske mitologije: {len(dictionary["serbian"]["dicts"]["svet_mit"])} items.')




# Arvacki Recnik kupdf_net_arvacki-recnik

def aracki_is_valid(text: str) -> bool:
    for i in " ,.()-;'\"<>/?[]:~":
        text = text.replace(i, "")
    if text.strip():
        return True
    return False

def aracki_fix_space(text: str) -> str:
    text = text.replace("\t", " ")
    while True:
        text = text.replace("  ", " ")
        if text.find("  ") == -1:
            break
    return text

doc = Document("kupdf_net_arvacki-recnik.docx")

dictionary["serbian"]["dicts"]["arvacki"] = {}

start = False
body = ""
topic_name = ""
header = False
new_topic = ""
counter = 0
for paragraph in doc.paragraphs:
    par_text = paragraph.text.strip()

    start = True
    if par_text in ["Adapa"]:
        start = True

    if par_text in ["JUŽNA I SREDNJA AZIJA"]:
        start = False

    if not start:
        continue

    if len(par_text) < 3:
        continue

    new_topic = ""
    body = ""

    topic_list = []

    for run in paragraph.runs:
        if run.bold:
            if aracki_is_valid(new_topic) and aracki_is_valid(body):
                topic_list.append([aracki_fix_space(new_topic), aracki_fix_space(body)])
                new_topic = ""
                body = ""
            new_topic += run.text + " "
        else:
            body += run.text + " "

    if  new_topic and body:
        topic_list.append([aracki_fix_space(new_topic), aracki_fix_space(body)])

    # counter += 1
    # if counter % 10 == 0:
    #     print(f"{counter} items processed")


    for i in topic_list:
        new_topic = i[0]
        body = i[1]
        new_topic = new_topic.strip().replace("\t", " ")
        body = body.strip().replace("\t", " ")

        if new_topic:
            topic_name = new_topic
            new_topic = ""
            if topic_name in dictionary["serbian"]["dicts"]["arvacki"]:
                dictionary["serbian"]["dicts"]["arvacki"][topic_name]["text"] += "\n\n\n"
                print (f"Duplicate Entry : {topic_name}")
            else:
                # print (topic_name)
                dictionary["serbian"]["dicts"]["arvacki"][topic_name] = {}
                dictionary["serbian"]["dicts"]["arvacki"][topic_name]["text"] = ""
                dictionary["serbian"]["dicts"]["arvacki"][topic_name]["links"] = []
                new_topic = ""

        if dictionary["serbian"]["dicts"]["arvacki"][topic_name]["text"]:
            dictionary["serbian"]["dicts"]["arvacki"][topic_name]["text"] += f" {body}"
        else:
            dictionary["serbian"]["dicts"]["arvacki"][topic_name]["text"] += body


print ()
print (f'Arvacki recnik: {len(dictionary["serbian"]["dicts"]["arvacki"])} items.')





# Beogradski frajerski recnik

def frajer_get_fixed_text(txt: str) -> str:
    delete_items = [
        "BEOGRAD",
        "PETRIT",
        "IMAMI",
        " I",
        "I "
    ]

    for item in delete_items:
        while True:
            pos = txt.find(item)
            if pos == -1:
                break
            start = txt[:pos].rfind("\n", pos)
            if start < 0:
                start = 0
            end = txt.find("\n", pos)
            if end < 0:
                end = len(txt)
            txt = txt[:start] + txt[end:]
            pos = start

    txt = txt.lower()
    spec_chars = "~!@#$%^&*(_+=0987654321[]){}';:\"|<>?,./"
    for i in spec_chars:
        txt = txt.replace(i, " ")

    while True:
        txt = txt.replace("  ", " ")
        if txt.find("  ") == -1:
            break
    

    return txt

def frajer_topics_list(txt: str) -> list:
    txt = frajer_get_fixed_text(txt)

    txt_list = [x.strip() for x in txt.split("\n") if x.strip() != ""]
    topics = []
    for item in txt_list:
        t = [x.strip() for x in item.split(" ") if x.strip() != ""]
        for i in t:
            topics.append(i)
    
    char_set = set()
    for i in topics:
        char_set.add(i[0])
    
    char_list = ""
    for i in topics:
        char_list += i[0]

    topic_char = None
    for i in char_set:
        if int(char_list.count(i) / len(char_list) * 100) > 45:
            topic_char = i
            break
    
    if not topic_char:
        return None
    
    result = []
    topic = ""
    for i in topics:
        if i[0] == topic_char:
            if topic:
                result.append(topic.strip())
            topic = i + " "
        else:
            topic += i + " "
    if topic:
        result.append(topic.strip())
    
    return result

def frajer_themes_list(txt: str, theme_list: list) -> str:
    delete_items = [
        "BEOGRAD",
        "PETRIT",
        "IMAMI",
        " I",
        "I "
    ]
    delimiter = "          "
    txt_list = [x.strip() for x in txt.split("\n") if x.strip() != ""]
    
    remaining = ""
    for item in txt_list:
        is_valid = True
        for i in delete_items:
            if item.find(i) >= 0:
                is_valid = False
        if not is_valid:
            continue

        if item.find(delimiter) == -1:
            item = item.replace("=", delimiter, 1)
        item = item.replace(delimiter, "=", 1)
        while True:
            item = item.replace("  ", " ")
            if item.find("  ") == -1:
                break
        
        pos = item.find("=")
        if pos == -1:
            remaining += item + " "
            continue

        item = item.replace(":", ",")
        theme_text = "\n".join([x.strip() for x in item[pos+1:].strip().split(",") if x.strip() != ""])

        theme_list.append([item[:pos].strip(), theme_text])

    return remaining.strip()

def frajer_examine_topic(txt: str, topic_list: list, theme_list: list):
    txt = frajer_themes_list(txt, theme_list)

    result = frajer_topics_list(txt)
    if result:
        for i in result:
            topic_list.append(i)


# doc = Document("beogradski_frajerski_recnik.docx")

dictionary["serbian"]["dicts"]["frajer"] = {}

with open("beogradski_frajerski_recnik_topics.txt", "r", encoding="utf-8") as file:
    topics = file.read()



topic_list = []
theme_list = []

pos = 0
start = None
while True:
    pos = topics.find("\"", pos)
    if pos == -1:
        break

    if start is None:
        start = pos
        pos += 1
        continue

    text = topics[start+1:pos]

    frajer_examine_topic(text, topic_list, theme_list)

    pos += 1
    start = None

def frajer_replace(txt: str, in_quote: bool = False, only_recur: bool = False) -> str:
    txt = txt.replace("\n", ",")
    txt = txt.replace("////", " ili ")
    txt = txt.replace("///", " ili ")
    txt = txt.replace("//", " ili ")
    txt = f' {txt.replace(":", ",")} '
    for i in range(1, 31):
        if txt.find(str(i)) >= 0:
            txt = txt.replace(str(i), f"{str(i)}.")
        else:
            break
    txt = txt.replace("..", ".")
    
    if txt.find(",jebali su mu majku!") >= 0:
        print ()

    # Replace , between items
    count = 1
    pos_start = 0
    while True:
        pos_start = txt.find(f"{str(count)}.", pos_start)
        if pos_start >= 0:
            if count == 1:
                first_1 = pos_start
            pos_end = txt.find(f"{str(count+1)}.", pos_start)
            if pos_end >= 0 and txt[pos_start+2:pos_end].find("1.") >= 0:
                pos_start = first_1 +1
                count = 1
                continue
            # print (txt[pos_start:pos_start+20])
            if pos_end >= 0:
                occur = txt[pos_start:pos_end].count(",") - 1
                if pos_start < txt.find("="):
                    if occur > 0:
                        txt = txt[:pos_start] + txt[pos_start:pos_end].replace(",", ";", occur) + txt[pos_end:]
                else:
                    txt = txt[:pos_start] + txt[pos_start:pos_end].replace(",", ";") + txt[pos_end:]
                pos_start = pos_end
                count += 1
            else:
                pos_start = first_1 + 1
                count = 1
        else:

            break

    while True:
        if txt.find("//") == -1:
            break
        txt = txt.replace("//", "/")

    if not only_recur:
        # if in_quote:
        #     txt = txt.replace(",", " ")

        repl_first = [
            ["—", "="],
            [";", "; "],
            ["=", " = "]
        ]
        
        repl = [
            ["/r.", "/v."],
            [" r.", " v."],
            ["oA", "od"],
            ["7 ", "/ "],
            [" če ", " će "],
            ["jr.", "fr."],
            ["(nm.)", "(mn.)"],
            ["(nm)", "(mn)"],
            ["/ćr.", "/izv."],
            [" Jig.", " Fig."],
            ["!v. ", "/v. "],
            [" lv. ", " /v. "],
            ["iasoc.", "/asoc."],
            ["(atigni.)", "(augm.)"],
            [" tinver.", " /inver."],
            ["//cr.", "//izv."],
            [":,", ":"],
            ["tinver.", "/inver."],
            [",,", ","],
            ["(augni.)", "(augm.)"],
            ["oA", "od"],
            ["(bipok.)", "(hipok.)"],
            ["lamer.", "/amer."],
            ["lamer;", "/amer."],
            ["lizv.", "/izv."],
            ["/f.", "/v."],
            ["!eng.", "/eng."],
            ["(bipok", "(hipok."],
            ["/Jr.", "/fr."],
            ["inem.", "/nem."],
            ["(um.)", "(mn.)"],
            ["Uzv.", "/izv."],
            ['"!', "/"],
            ["(l)ipok", "(hipok"],
            ['"I', "/"],
            [" J", " l"],
            ["I", "l"],
            ['"I', "/"],
            ["/ćr", "/izv"],
            ["/avoc", "/asoc"],
            ["(jnn.)", "(izv.)"],
            ["\"\"'", "/"],
            ['""1', "/"],
            [" lamer.", " /amer."],
            [" lamer ", " /amer."],
            ["/fcr.", "/izv."],
            ["Zfeu ", "/izv. "],
            ["/bipok.", "/hipok."],
            ["(nm/)", "(mn.)"],
            ["(se)", "se"],
            ["\\", " "],
            ["/skrač.", "/skrać."],
            ["/ćv.", "/izv."],
            ["(fcu)", "(izv.)"],
            ["ezz/", "euf."],
            [" si.", " sl."],
            ["/fer.", "/izv."],
            ["»", " "],
            [" tv.", " /v."],
            ["/skrač.", "/skrać."],
            ["draškan", "druškan"],
            ["(nm,)", "(mn.)"],
            ["(jnn.)", "(izv.)"],
            ["{", "("],
            ["}", ")"],
            ["[", "("],
            ["]", ")"],
            ["(/wz.)", "(mn.)"],
            ["je 1’", "jel "],
            [" 1’", "l "],
            ["je F", "jel "],
            ["je I", "jel "],
            ["/£r.", "/izv."],
            ["/ću ", "/izv. "],
            ["jasoc", "/asoc"],
            ["'", ""],
            [" Izv.", " /izv."],
            ["(inver. ", "/inver. "],
            ["(ft ", "/fr. "],
            [" 'izv.", " /izv."],
            ["", ""],
            ["", ""],
            ["", ""],
            ["", ""],
            ["", ""],
            ["", ""],

            ["’", " "],
            ['"', " "],
            ["„", " "],
            ["/", " / "],
            ["”", " "]
        ]
        repl.sort(key=lambda x: len(x[0]), reverse=True)

        repl_after = [
            [" miga!", "mi ga!"],
            ["(skrač)", "(skrać)"],
            ["lugoslav", "Jugoslav"],
            ["", ""],
            ["", ""],
            ["", ""],
            ["", ""],
            ["", ""],
            ["", ""],
            ["", ""],
            ["", ""],
            ["", ""]

        ]

        for i in repl_first:
            if not i[0]:
                continue
            if txt.find(i[0]) >= 0:
                txt = txt.replace(i[0], i[1])

        for i in repl:
            if not i[0]:
                continue
            if txt.find(i[0]) >= 0:
                txt = txt.replace(i[0], i[1])
        
    repl_recur = [
        ["  ", " "],
        [". .", ".."],
        ["..", "."],
        [", ,", ",,"],
        [",,", ","],
        ["/ /", "//"],
        ["//", "/"]
    ]

    for i in repl_recur:
        while True:
            if txt.find(i[0]) == -1:
                break
            txt = txt.replace(i[0], i[1])
    
    return txt.strip()

def frajer_get_fixed_body(txt: str) -> str:
    delim = "@@"
    txt = f" {frajer_replace(txt.replace(':', ','))} "
    # txt = txt.replace("/", " / ")
    
    pos_start = 0
    pos_end = None
    pos = 0
    fixed_topic_body = ""
    while True:
        pos = txt.find("/", pos_start)

        if pos == -1:
            fixed_topic_body += " " + txt[pos_start:].replace(",", delim)
            break

        pos_end = txt.find("/", pos + 1)

        if pos_end == -1:
            fixed_topic_body += txt[pos_start:pos].replace(",", delim)
            fixed_topic_body += txt[pos:] + " /"
            break

        fixed_topic_body += txt[pos_start:pos].replace(",", delim)
        
        fixed_topic_body += " " + txt[pos:pos_end+1]
        if txt[pos_end+1:].strip():
            if txt[pos_end+1:].strip()[0] != ",":
                fixed_topic_body += delim

        pos_start = pos_end + 1
    
    return fixed_topic_body.strip()
        

with open ("beogradski_frajerski_recnik_final_theme.txt", "r", encoding="utf-8") as file:
    themes = file.read()

themes_list = [x.strip() for x in themes.split("\n") if x.strip()]

topic_list = []
for theme in themes_list:
    topic_name = theme[:theme.find("=")].strip()
    topic_name = "@@@@" + topic_name
    topic_body = theme[theme.find("=") + 1:].strip()

    if topic_name == "blamotorno":
        print (f"Found : {topic_name}")

    # Find new topic name if already exist in base
    if topic_name in [x[0] for x in topic_list]:
        count = 1
        new_topic_name = topic_name
        while new_topic_name in topic_list:
            new_topic_name = f"{topic_name} ({count})"
            count += 1
        topic_name = new_topic_name
    
    topic_body = frajer_get_fixed_body(topic_body)
    
    topic_list.append([topic_name, topic_body])

topic_list.append(["@@@@@", ""])
for theme in theme_list:
    topic_name = theme[0].strip()
    topic_name = "@@@@" + topic_name
    topic_body = theme[1].strip()
    if topic_name == "frka":
        print ()

    # Find new topic name if already exist in base
    if topic_name in [x[0] for x in topic_list]:
        print ("Topic exist: ", topic_name)
        continue
        # count = 1
        # new_topic_name = topic_name
        # while new_topic_name in topic_list:
        #     new_topic_name = f"{topic_name} ({count})"
        #     count += 1
        # topic_name = new_topic_name
    
    topic_body = frajer_get_fixed_body(topic_body)

    topic_list.append([topic_name, topic_body])
    
txt = ""
for i in topic_list:
    if i[0] == "@@@@@":
        txt += "\n" * 10
        continue

    topic_name = i[0]
    body_list = [x.strip() for x in i[1].split("@@")]
    topic_body = "\n".join(body_list)
    
    txt += f'{topic_name}\n{topic_body}\n\n'


with open ("beogradski_frajerski_recnik_output.txt", "w", encoding="utf-8") as file:
    file.write(txt)


start = False
body = ""
topic_name = ""
header = False
new_topic = ""
counter = 0
topic_list = []
body_list = []

for line in txt.split("\n"):
    
    if not line:
        continue
    
    if line.find("@@@@") >= 0:
        new_topic = line.replace("@@@@", "")
        body = ""
    else:
        if line:
            body = line + "\n"
        else:
            continue

    if new_topic:
        topic_name = new_topic
        new_topic = ""
        if topic_name in dictionary["serbian"]["dicts"]["frajer"]:
            name_count = 1
            while True:
                try_new_name = topic_name + f"({name_count})"
                if try_new_name in dictionary["serbian"]["dicts"]["frajer"]:
                    name_count += 1
                else:
                    break
            topic_name = try_new_name

            dictionary["serbian"]["dicts"]["frajer"][topic_name] = {}
            dictionary["serbian"]["dicts"]["frajer"][topic_name]["text"] = ""
            dictionary["serbian"]["dicts"]["frajer"][topic_name]["links"] = []
            print (f"Duplicate Entry : {topic_name}")
        else:
            dictionary["serbian"]["dicts"]["frajer"][topic_name] = {}
            dictionary["serbian"]["dicts"]["frajer"][topic_name]["text"] = ""
            dictionary["serbian"]["dicts"]["frajer"][topic_name]["links"] = []
            new_topic = ""

    if dictionary["serbian"]["dicts"]["frajer"][topic_name]["text"]:
        dictionary["serbian"]["dicts"]["frajer"][topic_name]["text"] += f" {body}"
    else:
        dictionary["serbian"]["dicts"]["frajer"][topic_name]["text"] += body

with open ("beogradski_frajerski_recnik_tematski_registar.txt", "r", encoding="utf-8") as file:
    txt = file.read()

txt_list = txt.split("\n")

new_topic = ""
start_new_topic = True
for line in txt_list:
    if not line:
        start_new_topic = True
        continue

    if start_new_topic:
        new_topic = f"TEMA: {line}"
        body = ""
        start_new_topic = False
    else:
        new_topic = ""

    if not new_topic:
        body = line

    if new_topic:
        topic_name = new_topic
        new_topic = ""
        if topic_name in dictionary["serbian"]["dicts"]["frajer"]:
            name_count = 1
            while True:
                try_new_name = topic_name + f"({name_count})"
                if try_new_name in dictionary["serbian"]["dicts"]["frajer"]:
                    name_count += 1
                else:
                    break
            topic_name = try_new_name

            dictionary["serbian"]["dicts"]["frajer"][topic_name] = {}
            dictionary["serbian"]["dicts"]["frajer"][topic_name]["text"] = ""
            dictionary["serbian"]["dicts"]["frajer"][topic_name]["links"] = []
            print (f"Duplicate Entry : {topic_name}")
        else:
            dictionary["serbian"]["dicts"]["frajer"][topic_name] = {}
            dictionary["serbian"]["dicts"]["frajer"][topic_name]["text"] = ""
            dictionary["serbian"]["dicts"]["frajer"][topic_name]["links"] = []
            new_topic = ""

    if dictionary["serbian"]["dicts"]["frajer"][topic_name]["text"]:
        dictionary["serbian"]["dicts"]["frajer"][topic_name]["text"] += f" {body}"
    else:
        dictionary["serbian"]["dicts"]["frajer"][topic_name]["text"] += body


print ()
print (f'Beogradski frajerski recnik: {len(dictionary["serbian"]["dicts"]["frajer"])} items.')
print ()





# astroloski_recnik

with open("astroloski_recnik.txt", "r", encoding="utf-8") as file:
    txt = file.read()

txt_list = txt.split("\n")

dictionary["serbian"]["dicts"]["astroloski"] = {}

start = False
body = ""
topic_name = ""
header = False
new_topic = ""
counter = 0
for line in txt_list:
    line = line.strip()

    if not line:
        continue

    delim = line.find(":")
    new_topic = line[:delim].strip()
    body = line[delim+1:].strip()

    if new_topic:
        topic_name = new_topic
        new_topic = ""
        if topic_name in dictionary["serbian"]["dicts"]["astroloski"]:
            dictionary["serbian"]["dicts"]["astroloski"][topic_name]["text"] += "\n\n\n"
            print (f"Duplicate Entry : {topic_name}")
        else:
            print (topic_name)
            dictionary["serbian"]["dicts"]["astroloski"][topic_name] = {}
            dictionary["serbian"]["dicts"]["astroloski"][topic_name]["text"] = ""
            dictionary["serbian"]["dicts"]["astroloski"][topic_name]["links"] = []
            new_topic = ""
            # continue

    if dictionary["serbian"]["dicts"]["astroloski"][topic_name]["text"]:
        dictionary["serbian"]["dicts"]["astroloski"][topic_name]["text"] += f" {body}"
    else:
        dictionary["serbian"]["dicts"]["astroloski"][topic_name]["text"] += body


print ()
print (f'astroloski_recnik: {len(dictionary["serbian"]["dicts"]["astroloski"])} items.')





# Biblija  Biblija Stari Zavet

def biblija_get_verbs(txt: str, bib_dict: dict, knjiga: str, notes: list) -> None:
    txt = txt + " "
    txt_tmp = txt
    nums = "0123456789"
    for i in nums:
        txt_tmp = txt_tmp.replace(i, "|")

    main = 0
    sub = 1000000
    count = 0
    pos = 0
    while True:
        pos = txt_tmp.find("|", pos)
        if pos == -1:
            break

        verb_num = int(float(txt[pos:txt.find(" ", pos)]))
        verb_len = len(txt[pos:txt.find(" ", pos)])
        
        if sub + 1 != verb_num:
            if main + 1 == verb_num:
                main = verb_num
                sub = 1
            else:
                sub = verb_num
        else:
            sub = verb_num
        
        verb_txt = txt[pos + verb_len:txt_tmp.find("|", pos + verb_len)].strip()
        verb_item = f'{knjiga} {main},{sub}'
        for i in notes:
            if i[0] == main and i[1] == sub:
                verb_txt += f'\n\n{i[2]}'
        
        bib_dict[verb_item] = {}
        bib_dict[verb_item]["text"] = verb_txt
        bib_dict[verb_item]["links"] = []

        pos += verb_len

def biblija_get_note(notes: list, bib_dict: dict, knjiga: str) -> None:
    nums = "0123456789"

    result = []
    note_map = []
    for note in notes:
        note = f" {note} "
        pos = 0
        while True:
            pos = note.find(",", pos)
            if pos == -1:
                break
            main_txt = note[note[:pos+1].rfind(" ") + 1:pos]
            sub_txt = note[pos+1:note.find(" ", pos)]
            main_num = None
            sub_num = None
            try:
                main_num = int(main_txt)
                sub_num = int(sub_txt)
            except ValueError:
                pos += 1
                continue
            
            note_map.append((pos - len(main_txt), main_num, sub_num))
            pos += 1

        for idx, i in enumerate(note_map):
            if idx >= len(note_map) - 1:
                end_pos = len(note)
            else:
                end_pos = note_map[idx+1][0]
            
            item = note[i[0]:end_pos]
            result.append([i[1], i[2], f"*{item.strip()}"])
        note_map = []
    return result


with open("biblija_stari_zavet.txt", "r", encoding="utf-8") as file:
    txt = file.read()

txt = txt.replace("\t", "  ")
txt_list = txt.split("\n")

dictionary["serbian"]["dicts"]["biblija_stari_zavet"] = {}

bibl_dict = {
    "stari": {},
    "novi": {}
}

header = False
header_txt = ""
zavet = "stari"
knjiga = None
note_mode = False
note_txt = ""
knjiga_text = ""
knjiga_note = []
for line in txt_list:
    line = cirilica_u_latinicu(line)
    line = line.replace("Th", "Ć")

    if line.startswith("@@@@"):
        if header:
            header = False
        else:
            header = True
    
    if header:
        if knjiga_text:
            biblija_get_verbs(knjiga_text.strip("@ "), dictionary["serbian"]["dicts"]["biblija_stari_zavet"], knjiga, biblija_get_note(knjiga_note, dictionary["serbian"]["dicts"]["biblija_stari_zavet"], knjiga))
            
        if line.strip(" @"):
            header_txt += line.strip() + "\n"
        knjiga_text = ""
        knjiga_note = []
        continue

    if header_txt:
        knjiga = header_txt[:header_txt.find("\n")].strip()
        header_txt = ""

    if line.startswith("*"):
        note_mode = True
    
    if note_mode and not line.strip():
        note_mode = False
        knjiga_note.append(note_txt.strip(" *"))
        note_txt = ""

    if note_mode:
        note_txt += line.strip(" ") + " "
        continue

    if line.replace("j", "J") == line.upper():
        continue

    knjiga_text += line.strip() + " "



print ()
print (f'Biblija: {len(dictionary["serbian"]["dicts"]["biblija_stari_zavet"])} items.')



# Biblija Novi Zavet


with open("biblija_novi_zavet.txt", "r", encoding="utf-8") as file:
    txt = file.read()
txt = txt.replace("\t", "  ")
txt_list = txt.split("\n")

dictionary["serbian"]["dicts"]["biblija_novi_zavet"] = {}

bibl_dict = {
    "stari": {},
    "novi": {}
}

header = False
header_txt = ""
zavet = "stari"
knjiga = None
note_mode = False
note_txt = ""
knjiga_text = ""
knjiga_note = []
for line in txt_list:
    line = cirilica_u_latinicu(line)
    line = line.replace("Th", "Ć")

    if line.startswith("@@@@"):
        if header:
            header = False
        else:
            header = True
    
    if header:
        if knjiga_text:
            biblija_get_verbs(knjiga_text.strip("@ "), dictionary["serbian"]["dicts"]["biblija_novi_zavet"], knjiga, biblija_get_note(knjiga_note, dictionary["serbian"]["dicts"]["biblija_novi_zavet"], knjiga))
            
        if line.strip(" @"):
            header_txt += line.strip() + "\n"
        knjiga_text = ""
        knjiga_note = []
        continue

    if header_txt:
        knjiga = header_txt[:header_txt.find("\n")].strip()
        header_txt = ""

    if line.startswith("*"):
        note_mode = True
    
    if note_mode and not line.strip():
        note_mode = False
        knjiga_note.append(note_txt.strip(" *"))
        note_txt = ""

    if note_mode:
        note_txt += line.strip(" ") + " "
        continue

    if line.replace("j", "J") == line.upper():
        continue

    knjiga_text += line.strip() + " "



print ()
print (f'Biblija: {len(dictionary["serbian"]["dicts"]["biblija_novi_zavet"])} items.')



# Biblijski Leksikon

def bibl_leks_check_new_topic(text: str, clean_txt: str, pos: int) -> tuple:
    if text[pos:].strip().startswith("NZ") or text[pos:].strip().startswith("SZ"):
        return None
    
    txt = text
    if pos > 0 and text[:pos].strip()[-1] == ">":
        return None
    
    topic = ""
    topic_clean = ""
    while True:
        end = txt.find(" ", pos + 1)
        if end == -1:
            break
        
        if clean_txt[pos + 1:end + 1] == clean_txt[pos + 1:end + 1].upper():
            topic += txt[pos + 1:end + 1]
            topic_clean += clean_txt[pos + 1:end + 1]
        else:
            break
        pos = end
    
    count = 0
    if len(topic_clean) >= 2:
        if topic_clean[0] not in "QWERTYUIOPLKJHGFDSAZXCVBNMŠĐŽĆČ" or topic_clean[1] not in "QWERTYUIOPLKJHGFDSAZXCVBNMŠĐŽĆČ":
            return None

    for i in topic_clean:
        if i in "QWERTYUIOPLKJHGFDSAZXCVBNMŠĐŽĆČ":
            count += 1
    
    if count > 1:
        return topic
    else:
        return None

dictionary["serbian"]["dicts"]["bibl_leksikon"] = {}

with open("biblijski_leksikon.txt", "r", encoding="utf-8") as file:
    txt = file.read()

txt_list1 = [x + "\n" for x in txt.split("\n")]

txt = ""
txt_list = []
skip_next = False
for i in txt_list1:
    if not i.strip():
        skip_next = True
        continue

    if skip_next:
        skip_next = False
        continue

    txt_list.append(i)

repl = [
    ["-   ", ""],
    ["-  ", ""],
    ["- ", ""],
    ["-", ""],
    ["\t", "  "]
]

for idx in range(len(txt_list)):
    for i in repl:
        txt = txt_list[idx]

        txt = txt.replace(i[0], i[1])
        while txt.find("  ") != -1:
            txt = txt.replace("  ", " ")

        txt_list[idx] = txt



print (f"Biblijski leksikon STARTED:")
print ("Replacing () ... ", end="")

new_txt_list = []
percent = 0
for idx, txt in enumerate(txt_list):
    if int(idx/len(txt_list)*100) != percent and int(idx/len(txt_list)*100) % 4 == 0:
        print (f"{percent}%, ", end="")
        percent = int(idx/len(txt_list)*100)

    pos = 0
    old_pos = 0
    new_txt = ""
    while True:
        pos_start = txt.find("(", pos)
        if pos_start == -1:
            new_txt += txt[pos:]
            break
        pos_end = txt.find(")", pos_start)
        if pos_end == -1:
            new_txt += txt[pos:]
            break
        new_txt += txt[old_pos:pos_start]
        new_txt += "1" * (pos_end - pos_start + 1)
        old_pos = pos_end + 1
        pos = pos_end + 1 

    new_txt = new_txt.replace(" i ", " 1 ")
    new_txt = new_txt.replace("Lj", "LJ")
    new_txt = new_txt.replace("Nj", "NJ")
    new_txt = new_txt.replace("Dž", "DŽ")

    new_txt_list.append(new_txt)

print ("Done.")

pos = 0
old_pos = 0
new_topic = ""
topic_name = ""
body = ""
end = None
for idx in range(len(txt_list)):
    result = bibl_leks_check_new_topic(txt_list[idx], new_txt_list[idx], -1)
    if result:
        new_topic = result.strip(" :,=")
        body = txt_list[idx][len(result):]
    else:
        body = txt_list[idx]

    if new_topic:
        topic_name = new_topic
        new_topic = ""
        if topic_name in dictionary["serbian"]["dicts"]["bibl_leksikon"]:
            dictionary["serbian"]["dicts"]["bibl_leksikon"][topic_name]["text"] += "\n\n\n"
            print (f"Duplicate Entry : {topic_name}")
        else:
            # print (topic_name)
            dictionary["serbian"]["dicts"]["bibl_leksikon"][topic_name] = {}
            dictionary["serbian"]["dicts"]["bibl_leksikon"][topic_name]["text"] = ""
            dictionary["serbian"]["dicts"]["bibl_leksikon"][topic_name]["links"] = []
            new_topic = ""
            # continue

    if dictionary["serbian"]["dicts"]["bibl_leksikon"][topic_name]["text"]:
        dictionary["serbian"]["dicts"]["bibl_leksikon"][topic_name]["text"] += f" {body}"
    else:
        dictionary["serbian"]["dicts"]["bibl_leksikon"][topic_name]["text"] += body




print ()
print (f'Biblijski leksikon: {len(dictionary["serbian"]["dicts"]["bibl_leksikon"])} items.')



# Tracanska, Ilirska i Slovenska Mitologija

with open("tracanska_ilirska_slovenska_mitologija.txt", "r", encoding="utf-8") as file:
    txt = file.read()

txt_list = txt.split("\n")

dictionary["serbian"]["dicts"]["tis_mit"] = {}


main_theme = ""
body = ""
topic_name = ""
new_topic = ""
counter = 0
for line in txt_list:
    line = cirilica_u_latinicu(line.strip())

    if line.startswith("@@@"):
        main_theme = line.replace("@@@", "")
        continue
    
    if not line.strip():
        topic_name = ""
        continue

    if not topic_name:
        new_topic = f"{main_theme}: {line}"
        body = ""

    if not new_topic:
        if len(line) < 100:
            print (f"WARNING: {line}")
        body = line

    if new_topic:
        topic_name = new_topic
        new_topic = ""
        if topic_name in dictionary["serbian"]["dicts"]["tis_mit"]:
            dictionary["serbian"]["dicts"]["tis_mit"][topic_name]["text"] += "\n\n\n"
            print (f"Duplicate Entry : {topic_name}")
        else:
            # print (topic_name)
            dictionary["serbian"]["dicts"]["tis_mit"][topic_name] = {}
            dictionary["serbian"]["dicts"]["tis_mit"][topic_name]["text"] = ""
            dictionary["serbian"]["dicts"]["tis_mit"][topic_name]["links"] = []
            new_topic = ""
            continue

    if dictionary["serbian"]["dicts"]["tis_mit"][topic_name]["text"]:
        dictionary["serbian"]["dicts"]["tis_mit"][topic_name]["text"] += f"{body}\n"
    else:
        dictionary["serbian"]["dicts"]["tis_mit"][topic_name]["text"] += f"{body}\n"


print ()
print (f'Tracanska, Ilirska i Slovenska mitologija: {len(dictionary["serbian"]["dicts"]["tis_mit"])} items.')




# Dzepni pravni recnik

doc = Document("dzepni_pravni_recnik_zeljko_kuvizic_gradjevinska_knjiga.docx")

dictionary["serbian"]["dicts"]["dz_pravni"] = {}

start = False
body = ""
topic_name = ""
header = False
new_topic = ""
counter = 0
for paragraph in doc.paragraphs:
    par_text: str = paragraph.text.strip()

    if par_text.find("abatement") >= 0 or par_text.strip() == "Institucije Evropske unije":
        start = True

    if par_text == "Dodatak I" or par_text.startswith("Tabela 1. Hronološki spisak osnivačkih ugovora"):
        start = False

    if not start:
        continue

    if len(par_text) < 2:
        continue

    new_topic = ""
    body = ""
    for run in paragraph.runs:
        if run.bold:
            if par_text.startswith(new_topic + run.text):
                new_topic += run.text
            else:
                body += run.text
        else:
            body += run.text

    new_topic = new_topic.strip()
    body = body.strip()

    if new_topic.replace("1.", "").strip() in ["emit", "hazard", "label", "leave", "line", "material (opšt.)", "offense", "ordinary course of business", "petition", "portfolio", "purview", "relief", "scalping", "stake", "usage"]:
        print (f"Corrected: {new_topic}  ->  ", end="")
        new_topic = new_topic.replace("1.", "").strip()
        print (new_topic)
        body = "1. " + body
    if new_topic == "raid I.":
        print (f"Corrected: {new_topic}  ->  ", end="")
        new_topic = new_topic.replace("I.", "").strip()
        print (new_topic)
        body = "1. " + body

    if new_topic:
        if topic_name:
            txt = dictionary["serbian"]["dicts"]["dz_pravni"][topic_name]["text"]
            txt = txt.replace("-@@", "")
            txt = txt.replace("@@", "")
            dictionary["serbian"]["dicts"]["dz_pravni"][topic_name]["text"] = txt

        topic_name = new_topic
        new_topic = ""
        if topic_name in dictionary["serbian"]["dicts"]["dz_pravni"]:
            dictionary["serbian"]["dicts"]["dz_pravni"][topic_name]["text"] += "\n\n\n"
            print (f"Duplicate Entry : {topic_name}")
        else:
            # print (topic_name)
            dictionary["serbian"]["dicts"]["dz_pravni"][topic_name] = {}
            dictionary["serbian"]["dicts"]["dz_pravni"][topic_name]["text"] = ""
            dictionary["serbian"]["dicts"]["dz_pravni"][topic_name]["links"] = []
            new_topic = ""
            # continue

    if dictionary["serbian"]["dicts"]["dz_pravni"][topic_name]["text"]:
        dictionary["serbian"]["dicts"]["dz_pravni"][topic_name]["text"] += f"@@{body}"
    else:
        dictionary["serbian"]["dicts"]["dz_pravni"][topic_name]["text"] += body


print ()
print (f'Recnik svetske mitologije: {len(dictionary["serbian"]["dicts"]["dz_pravni"])} items.')




# Eponimski recnik


doc = Document("eponimski_leksikon_dubravko_mri.docx")

dictionary["serbian"]["dicts"]["eponim"] = {}

start = False
body = ""
topic_name = ""
header = False
new_topic = ""
counter = 0
for paragraph in doc.paragraphs:
    par_text: str = paragraph.text.strip()

    if par_text == "Abasidi":
        start = True

    if par_text in ["JUŽNA I SREDNJA AZIJA", "Sibir Mongolija Kina Japan\njugoistočna Azija", "EVROPA", "AMERIKA", "AFRIKA", "OKEANIJA", "LITERATURA"]:
        start = False

    if not start:
        continue

    if len(par_text) < 2:
        continue

    if par_text.startswith("dejvi"):
        print ()
        # for run in paragraph.runs:
        #     print (run.bold, run.font.size, run.text)

    if paragraph.style.font.size == 139700:
        header = True
        body = ""
        for run in paragraph.runs:
            if run.font.size:
                header = False
            
            if header:
                new_topic += run.text
            else:
                body += run.text
        if not body:
            continue
    else:
        if body.endswith(" "):
            body = par_text
        else:
            if body:
                body = " " + par_text
            else:
                body = par_text
    


    new_topic = new_topic.strip()
    body = body.strip()

    if new_topic:
        new_topic = new_topic.replace("~", "")
        new_topic = new_topic.strip(" ,")
        if topic_name:
            txt = dictionary["serbian"]["dicts"]["eponim"][topic_name]["text"]
            txt = txt.replace("- ", "")
            txt = txt.replace("~", "\n")
            txt = txt.strip()
            dictionary["serbian"]["dicts"]["eponim"][topic_name]["text"] = txt
        topic_name = new_topic
        new_topic = ""
        if topic_name in dictionary["serbian"]["dicts"]["eponim"]:
            dictionary["serbian"]["dicts"]["eponim"][topic_name]["text"] += "\n\n\n"
            print (f"Duplicate Entry : {topic_name}")
        else:
            # print (topic_name)
            dictionary["serbian"]["dicts"]["eponim"][topic_name] = {}
            dictionary["serbian"]["dicts"]["eponim"][topic_name]["text"] = ""
            dictionary["serbian"]["dicts"]["eponim"][topic_name]["links"] = []
            new_topic = ""
            # continue

    if dictionary["serbian"]["dicts"]["eponim"][topic_name]["text"]:
        dictionary["serbian"]["dicts"]["eponim"][topic_name]["text"] += f"\n{body}"
    else:
        dictionary["serbian"]["dicts"]["eponim"][topic_name]["text"] += body


print ()
print (f'Eponimski recnik: {len(dictionary["serbian"]["dicts"]["eponim"])} items.')



# Leksikon osnovnih jungovskih pojmova

def jung_fix_text(txt: str) -> str:
    txt = txt.replace("- ", "")
    txt = txt.replace("~", "\n")
    while txt.find("  ") >= 0:
        txt = txt.replace("  ", " ")

    for i in "qwertyplkjhgfdsazxcvbnmčćžšđ":
        txt = txt.replace(f" {i} ", f"{i} ")
        
    txt = txt.replace("esc", "e se")
    # txt = txt.replace(" sc ", " se ")
    for i in "qwtydnšđčćžmbvczlkjhgfs":
        txt = txt.replace(f"{i}c", f"{i}e")

    txt = txt.replace("šio", "što")
    txt = txt.replace("šc", "še")
    txt = txt.replace("polp", "potp")
    
    txt = txt.replace("šl", "št")
    txt = txt.replace("lime", "time")
    txt = txt.replace("Ilark", "Hark")
    txt = txt.replace("Terson", "Person")

    txt = txt.replace("kei", "kci")
    txt = txt.replace("pei", "pci")
    txt = txt.replace("cr", "er")
    txt = txt.replace("sir", "str")
    txt = txt.replace("Nc", "Ne")

    txt = txt.strip()
    return txt

doc = Document("hark_leksikon_osnovnih_jungovskh_pojmova.docx")

dictionary["serbian"]["dicts"]["jung"] = {}

start = False
body = ""
topic_name = ""
header = False
new_topic = ""
counter = 0
skip_until = True
for paragraph in doc.paragraphs:
    par_text: str = paragraph.text.strip()

    if par_text == "DERETA, 1998":
        skip_until = False

    if skip_until:
        continue

    if par_text.find("PREDGOVOR") >= 0:
        start = True
        

    if par_text == "Helmut Hark":
        dictionary["serbian"]["dicts"]["jung"][topic_name]["text"] = jung_fix_text(dictionary["serbian"]["dicts"]["jung"][topic_name]["text"])
        start = False

    if not start:
        continue

    if not par_text.strip():
        continue


    if par_text.startswith("OSET"):
        print ()

    try:
        _ = int(par_text.strip())
        continue
    except ValueError:
        _ = 0

    if par_text.startswith("dejvi"):
        print ()
        # for run in paragraph.runs:
        #     print (run.bold, run.font.size, run.text)

    if paragraph.style.font.size == 127000:
        new_topic = par_text
        body = ""
    else:
        body = par_text

    new_topic = new_topic.strip()
    body = body.strip()

    if new_topic:
        new_topic = new_topic.replace("~", "")
        new_topic = new_topic.strip(" ,")
        if topic_name:
            dictionary["serbian"]["dicts"]["jung"][topic_name]["text"] = jung_fix_text(dictionary["serbian"]["dicts"]["jung"][topic_name]["text"])
        topic_name = new_topic
        new_topic = ""
        if topic_name in dictionary["serbian"]["dicts"]["jung"]:
            dictionary["serbian"]["dicts"]["jung"][topic_name]["text"] += "\n\n\n"
            print (f"Duplicate Entry : {topic_name}")
        else:
            # print (topic_name)
            dictionary["serbian"]["dicts"]["jung"][topic_name] = {}
            dictionary["serbian"]["dicts"]["jung"][topic_name]["text"] = ""
            dictionary["serbian"]["dicts"]["jung"][topic_name]["links"] = []
            new_topic = ""
            # continue

    if dictionary["serbian"]["dicts"]["jung"][topic_name]["text"]:
        dictionary["serbian"]["dicts"]["jung"][topic_name]["text"] += f" {body}"
    else:
        dictionary["serbian"]["dicts"]["jung"][topic_name]["text"] += body


print ()
print (f'Leksikon osnovnih jungovskih pojmova: {len(dictionary["serbian"]["dicts"]["jung"])} items.')




# HRT Leksikon

def hrt_repl(txt: str) -> str:
    repl = [
        ["\xad ", ""],
        ["\xad", ""],
        ["\uf034", "\n"]
    ]
    for i in repl:
        txt = txt.replace(i[0], i[1])
    
    return txt

doc = Document("hrt_leksikon.docx")

dictionary["serbian"]["dicts"]["hrt"] = {}

start = False
body = ""
topic_name = ""
header = False
new_topic = ""
counter = 0

for paragraph in doc.paragraphs:
    par_text: str = paragraph.text.strip()

    par_text = hrt_repl(par_text)

    page_test_par_text = par_text
    if page_test_par_text:
        if page_test_par_text[-1] == "I":
            page_test_par_text = page_test_par_text[:-1]
    for i in page_test_par_text:
        if i in "QWERTYUOPLKJHGFDSAZXCVBNM":
            page_test_par_text = page_test_par_text.replace(i, "")

    par_text_list = page_test_par_text.split("I")
    if par_text_list:
        header = False
        try:
            _ = int(par_text_list[0])
            header = True
        except ValueError:
            _ = 0
        try:
            _ = int(par_text_list[-1])
            header = True
        except ValueError:
            _ = 0
        
        if header:
            continue

    if par_text.startswith("ABC"):
        start = True

    if par_text == "PRILOZI":
        start = False

    if not start:
        continue

    if len(par_text) < 2:
        continue

    is_picture = False
    if paragraph.style.font.bold:
        if par_text.find("→") > 0:
            new_topic = par_text[:par_text.find("→")]
            body = par_text[par_text.find("→")+1:]
        else:
            continue
    else:
        new_topic = ""
        body = ""
        header = True
        for run in paragraph.runs:
            if run.font.name == "Gill Sans MT":
                is_picture = True
            if run.bold and header:
                if par_text.startswith(f"{new_topic}{run.text}"):
                    new_topic += run.text
                else:
                    body += run.text
                    header = False
            else:
                body += run.text

    if is_picture:
        continue

    new_topic = hrt_repl(new_topic)
    new_topic = new_topic.strip()
    body = hrt_repl(body)
    body = body.strip()

    if new_topic == "2.":
        body = new_topic + " " + body
        new_topic = ""

    if new_topic:
        topic_name = new_topic
        new_topic = ""
        if topic_name in dictionary["serbian"]["dicts"]["hrt"]:
            counter = 1 
            while True:
                if f"{topic_name} ({counter})" not in dictionary["serbian"]["dicts"]["hrt"]:
                    topic_name = f"{topic_name} ({counter})"
                    dictionary["serbian"]["dicts"]["hrt"][topic_name] = {}
                    dictionary["serbian"]["dicts"]["hrt"][topic_name]["text"] = ""
                    dictionary["serbian"]["dicts"]["hrt"][topic_name]["links"] = []
                    new_topic = ""
                    break
                else:
                    counter += 1
            print (f"Duplicate Entry Corrected: {topic_name}")
        else:
            # print (topic_name)
            dictionary["serbian"]["dicts"]["hrt"][topic_name] = {}
            dictionary["serbian"]["dicts"]["hrt"][topic_name]["text"] = ""
            dictionary["serbian"]["dicts"]["hrt"][topic_name]["links"] = []
            new_topic = ""
            # continue

    if dictionary["serbian"]["dicts"]["hrt"][topic_name]["text"]:
        dictionary["serbian"]["dicts"]["hrt"][topic_name]["text"] += f" {body}"
    else:
        dictionary["serbian"]["dicts"]["hrt"][topic_name]["text"] += body


print ()
print (f'HRT Leksikon: {len(dictionary["serbian"]["dicts"]["hrt"])} items.')



# Znacenje vlastitih imena

dictionary["serbian"]["dicts"]["imena"] = {}

with open("spisak_licnih_imena.json", "r", encoding="utf-8") as file:
    names_dict = json.load(file)

new_topic = ""
body = ""
for name in names_dict:
    if name == "Muški":
        continue

    new_topic = name
    body = names_dict[name]["desc"] + "\n\n" + names_dict[name]["num"]

    if new_topic:
        topic_name = new_topic
        new_topic = ""
        # print (topic_name)
        dictionary["serbian"]["dicts"]["imena"][topic_name] = {}
        dictionary["serbian"]["dicts"]["imena"][topic_name]["text"] = ""
        dictionary["serbian"]["dicts"]["imena"][topic_name]["links"] = []
        new_topic = ""
        # continue

    if dictionary["serbian"]["dicts"]["imena"][topic_name]["text"]:
        dictionary["serbian"]["dicts"]["imena"][topic_name]["text"] += f"\n{body}"
    else:
        dictionary["serbian"]["dicts"]["imena"][topic_name]["text"] += body


print ()
print (f'Znacenje Vlastitih Imena: {len(dictionary["serbian"]["dicts"]["imena"])} items.')




# Recnik kosarke

with open("kosarka_recnik.txt", "r", encoding="utf-8") as file:
    txt = file.read()

txt_list = txt.split("\n")

dictionary["serbian"]["dicts"]["kosarka"] = {}

start = False
body = ""
topic_name = ""
header = False
new_topic = ""
counter = 0
for line in txt_list:
    line = line.strip()

    if not line:
        continue

    pos = line.find(" - ")
    new_topic = line[:pos].strip()
    body = line[pos + 3:].strip()

    if new_topic:
        topic_name = new_topic
        new_topic = ""
        if topic_name in dictionary["serbian"]["dicts"]["kosarka"]:
            dictionary["serbian"]["dicts"]["kosarka"][topic_name]["text"] += "\n\n\n"
            print (f"Duplicate Entry : {topic_name}")
        else:
            dictionary["serbian"]["dicts"]["kosarka"][topic_name] = {}
            dictionary["serbian"]["dicts"]["kosarka"][topic_name]["text"] = ""
            dictionary["serbian"]["dicts"]["kosarka"][topic_name]["links"] = []
            new_topic = ""
            # continue

    if dictionary["serbian"]["dicts"]["kosarka"][topic_name]["text"]:
        dictionary["serbian"]["dicts"]["kosarka"][topic_name]["text"] += f" {body}"
    else:
        dictionary["serbian"]["dicts"]["kosarka"][topic_name]["text"] += body


print ()
print (f'Recnik kosarke: {len(dictionary["serbian"]["dicts"]["kosarka"])} items.')





# Recnik jezickih nedoumica


doc = Document("Ivan_Klajn_Recnik_jezickih_nedoumica.docx")

dictionary["serbian"]["dicts"]["jez_nedoum"] = {}

start = False
body = ""
topic_name = ""
header = False
new_topic = ""
counter = 0
repl = [
    ["ušlo vom", "uslovom"],
    ["ađresant", "adresant"],
    ["-  ", ""],
    ["- ", ""],
    [" si.", " sl."],
    [" P ", " Pravopis "],
    [" PR 60 ", " Pravopis "],
    [" PR ", " Pravopis "],
    ["„", '"'],
    ['”', '"'],
    ["", ""],
    ["", ""]
]

for paragraph in doc.paragraphs:
    par_text: str = paragraph.text.strip()
    par_text = par_text.replace("\t", "")

    if par_text == "ž. - ženski":
        start = True
        continue

    if par_text.startswith("3a OHe"):
        start = False

    if not start:
        continue

    if len(par_text) < 2:
        continue

    new_topic = ""
    body = ""
    topic_list = []
    for run in paragraph.runs:
        if run.bold:
            if par_text.startswith(new_topic + run.text):
                new_topic += run.text
            else:
                topic_list.append([new_topic, body])
                new_topic = run.text
                body = ""
        else:
            body += run.text

    if new_topic or body:
        topic_list.append([new_topic, body])
    
    for item in topic_list:
        new_topic: str = item[0]
        body: str = item[1]

        new_topic = new_topic.strip(",.:; ")
        body = body.strip()
        if new_topic.endswith("(="):
            new_topic = new_topic.replace("(=", "").strip()
            body = "(=" + body
        if new_topic.endswith("(ne"):
            new_topic = new_topic.replace("(ne", "").strip()
            body = "(ne " + body

        for i in repl:
            new_topic = new_topic.replace(i[0], i[1])
            body = body.replace(i[0], i[1])

        if new_topic:
            topic_name = new_topic
            new_topic = ""
            if topic_name in dictionary["serbian"]["dicts"]["jez_nedoum"]:
                dictionary["serbian"]["dicts"]["jez_nedoum"][topic_name]["text"] += "\n\n\n"
                print (f"Duplicate Entry : {topic_name}")
            else:
                # print (topic_name)
                dictionary["serbian"]["dicts"]["jez_nedoum"][topic_name] = {}
                dictionary["serbian"]["dicts"]["jez_nedoum"][topic_name]["text"] = ""
                dictionary["serbian"]["dicts"]["jez_nedoum"][topic_name]["links"] = []
                new_topic = ""
                # continue

        if dictionary["serbian"]["dicts"]["jez_nedoum"][topic_name]["text"]:
            dictionary["serbian"]["dicts"]["jez_nedoum"][topic_name]["text"] += f" {body}"
        else:
            dictionary["serbian"]["dicts"]["jez_nedoum"][topic_name]["text"] += body



print ()
print (f'Recnik jezickih nedoumica: {len(dictionary["serbian"]["dicts"]["jez_nedoum"])} items.')



# Leksikon bibliotekarstva


doc = Document("kupdf_net_leksikon_bibliotekarstva.docx")

dictionary["serbian"]["dicts"]["bibliotek"] = {}

start = False
body = ""
topic_name = ""
header = False
new_topic = ""
counter = 0
keep_new_topic = False
repl = [
    ["ušlo vom", "uslovom"],
    ["ađresant", "adresant"],
    [" si.", " sl."],
    ["„", '"'],
    ['”', '"'],
    [" 1 ", " i "],
    ["\xad ", ""],
    ["»", ">"],
    ["«", "<"]
]
for paragraph in doc.paragraphs:
    par_text: str = paragraph.text.strip()

    if par_text == "K. Grubačić":
        start = True
        continue

    if par_text == "BIBLIOTEČKE 1 NAJVAŽNIJE OPŠTE SKRAĆENICE (NAŠE I STRANE)":
        start = False

    if not start:
        continue

    if par_text.startswith("ACME COLOUR SEPARATOR"):
        print ("---------------------------------------")

    if len(par_text) < 4:
        continue

    if par_text.strip().split(" ")[-1] == "Leksikon":
        # print (" .......... ........... Leksikon FOUND ! .......... ................ ")
        continue

    if paragraph.style.name == "List Paragraph":
        # print ()
        # print ("HEAD (Style): ", paragraph.text)
        # print ()
        continue
    
    if not keep_new_topic:
        new_topic = ""
    else:
        if new_topic.endswith("-"):
            new_topic = new_topic.strip("-")
        else:
            new_topic += " "
    body = ""
    for run in paragraph.runs:
        if run.bold:
            if par_text.startswith(new_topic + run.text):
                new_topic += run.text
            else:
                if keep_new_topic:
                    new_topic += run.text
                else:
                    body += run.text
        else:
            body += run.text
    keep_new_topic = False

    par_text_head = body.split("—")
    if len(par_text_head) == 2 and len(par_text_head[0]) < 50 and len(par_text_head[1]) < 50 and not new_topic:
        # print ()
        # print ("HEAD (Format - 1): ", paragraph.text)
        # print ()
        continue

    par_text_head = body.split("-")
    if len(par_text_head) == 2 and par_text_head[-1].find("\t") != -1 and not new_topic:
        try:
            _ = int(par_text_head[-1].split("\t")[-1])
            # print ()
            # print ("HEAD (Format - 2): ", paragraph.text)
            # print ()
            continue
        except Exception as e:
            _ = 0
            # print ("False alarm for heading; ", e)

    if body.endswith("\xad"):
        body = body[:body.rfind("\xad")] + "-"
    for i in repl:
        new_topic = new_topic.replace(i[0], i[1])
        body = body.replace(i[0], i[1])
    
    new_topic = new_topic.strip(" ,.—:=")
    body = body.strip()

    new_topic =  new_topic.replace("\t", " ")
    body = body.replace("\t", " ")
    while True:
        new_topic = new_topic.replace("  ", " ")
        body = body.replace("  ", " ")
        if new_topic.find("  ") == -1 and body.find("  ") == -1:
            break

    if new_topic and not body:
        keep_new_topic = True
        # print (f"Keep new topic activated   ({new_topic}) ......")
        continue

    if new_topic:
        topic_name = new_topic
        new_topic = ""
        if topic_name in dictionary["serbian"]["dicts"]["bibliotek"]:
            dictionary["serbian"]["dicts"]["bibliotek"][topic_name]["text"] += "\n\n\n"
            print (f"Duplicate Entry : {topic_name}")
        else:
            # print (topic_name)
            dictionary["serbian"]["dicts"]["bibliotek"][topic_name] = {}
            dictionary["serbian"]["dicts"]["bibliotek"][topic_name]["text"] = ""
            dictionary["serbian"]["dicts"]["bibliotek"][topic_name]["links"] = []
            new_topic = ""
            # continue

    if dictionary["serbian"]["dicts"]["bibliotek"][topic_name]["text"]:
        dictionary["serbian"]["dicts"]["bibliotek"][topic_name]["text"] += f" {body}"
    else:
        dictionary["serbian"]["dicts"]["bibliotek"][topic_name]["text"] += body



# Skracenice
with open("leksikon_bibliotekarstva_skracenice.txt", "r", encoding="utf-8") as file:
    txt = file.read()

txt_list = txt.split("\n")


start = False
body = ""
topic_name = ""
header = False
new_topic = ""
counter = 0
for line in txt_list:
    line = line.strip()

    if not line:
        continue

    new_topic = ""
    body = ""
    
    pos = line.find("=")
    if pos > 0:
        new_topic = line[:pos].strip()
        body = line[pos+1:].strip()
    else:
        pos = line.find("-")
        if pos > 0:
            new_topic = line[:pos].strip()
            body = line[pos+1:].strip()
        else:
            body = line

    if new_topic:
        topic_name = new_topic
        new_topic = ""
        if topic_name in dictionary["serbian"]["dicts"]["bibliotek"]:
            dictionary["serbian"]["dicts"]["bibliotek"][topic_name]["text"] += "\n\n\n"
            print (f"Duplicate Entry : {topic_name}")
        else:
            # print (topic_name)
            dictionary["serbian"]["dicts"]["bibliotek"][topic_name] = {}
            dictionary["serbian"]["dicts"]["bibliotek"][topic_name]["text"] = ""
            dictionary["serbian"]["dicts"]["bibliotek"][topic_name]["links"] = []
            new_topic = ""
            # continue

    if dictionary["serbian"]["dicts"]["bibliotek"][topic_name]["text"]:
        dictionary["serbian"]["dicts"]["bibliotek"][topic_name]["text"] += f" {body}"
    else:
        dictionary["serbian"]["dicts"]["bibliotek"][topic_name]["text"] += body


# Stari nazivi mesta
with open("leksikon_bibliotekarstva_latinski_nazivi_gradova.txt", "r", encoding="utf-8") as file:
    txt = file.read()

txt_list = txt.split("\n")


start = False
body = ""
topic_name = ""
header = False
new_topic = ""
counter = 0
for line in txt_list:
    line = line.strip()

    if not line:
        continue

    new_topic = ""
    body = ""
    
    pos = line.find("=")
    if pos > 0:
        new_topic = line[:pos].strip()
        body = line[pos+1:].strip()
    else:
        pos = line.find("-")
        if pos > 0:
            new_topic = line[:pos].strip()
            body = line[pos+1:].strip()
        else:
            body = line

    if new_topic:
        topic_name = new_topic
        new_topic = ""
        if topic_name in dictionary["serbian"]["dicts"]["bibliotek"]:
            dictionary["serbian"]["dicts"]["bibliotek"][topic_name]["text"] += "\n\n\n"
            print (f"Duplicate Entry : {topic_name}")
        else:
            # print (topic_name)
            dictionary["serbian"]["dicts"]["bibliotek"][topic_name] = {}
            dictionary["serbian"]["dicts"]["bibliotek"][topic_name]["text"] = ""
            dictionary["serbian"]["dicts"]["bibliotek"][topic_name]["links"] = []
            new_topic = ""
            # continue

    if dictionary["serbian"]["dicts"]["bibliotek"][topic_name]["text"]:
        dictionary["serbian"]["dicts"]["bibliotek"][topic_name]["text"] += f" {body}"
    else:
        dictionary["serbian"]["dicts"]["bibliotek"][topic_name]["text"] += body

for i in dictionary["serbian"]["dicts"]["bibliotek"]:
    dictionary["serbian"]["dicts"]["bibliotek"][i]["text"] = dictionary["serbian"]["dicts"]["bibliotek"][i]["text"].replace("\t", " ")
    while True:
        dictionary["serbian"]["dicts"]["bibliotek"][i]["text"] = dictionary["serbian"]["dicts"]["bibliotek"][i]["text"].replace("  ", " ")
        if dictionary["serbian"]["dicts"]["bibliotek"][i]["text"].find("  ") == -1:
            break

print ()
print (f'Leksikon Bibliotekarstva: {len(dictionary["serbian"]["dicts"]["bibliotek"])} items.')




# Leksikon Hriscanstva, Judaizma i Islama


doc = Document("leksikon_hriscanstva_judaizma_i_islama_aleksandar_272akovac.docx")

dictionary["serbian"]["dicts"]["leksikon_hji"] = {}


start = False
body = ""
topic_name = ""
header = False
new_topic = ""
keep_new_topic = False
counter = 0
for paragraph in doc.paragraphs:
    par_text: str = paragraph.text.strip()

    if par_text.find("Episkop“ i „Evharistija“ posebno obrađeni") != -1:
        start = True
        continue

    if par_text  == "DODATAK I":
        start = False

    if not start:
        continue

    if len(par_text) < 2:
        continue

    combine_topic = ""
    if keep_new_topic:
        combine_topic = new_topic + " "
    
    keep_new_topic = False
    new_topic = ""
    body = ""
    if paragraph.style.style_id != "BodyText":
        header = True
    else:
        header = False

    header = True
    for run in paragraph.runs:
        if header:
            if run.bold:
                check_topic: str = new_topic + run.text
                if par_text.startswith(check_topic) and check_topic == check_topic.upper():
                    new_topic += run.text
                else:
                    body += run.text
            else:
                body += run.text
        else:
            body += run.text

    if new_topic == par_text:
        keep_new_topic = True
        new_topic = combine_topic + new_topic
        continue
    
    new_topic = combine_topic + new_topic
    
    new_topic = new_topic.strip(" ,():-")
    body = body.strip()

    if new_topic:
        topic_name = new_topic
        new_topic = ""
        if topic_name in dictionary["serbian"]["dicts"]["leksikon_hji"]:
            dictionary["serbian"]["dicts"]["leksikon_hji"][topic_name]["text"] += "\n\n\n"
            print (f"Duplicate Entry : {topic_name}")
        else:
            # print (topic_name)
            dictionary["serbian"]["dicts"]["leksikon_hji"][topic_name] = {}
            dictionary["serbian"]["dicts"]["leksikon_hji"][topic_name]["text"] = ""
            dictionary["serbian"]["dicts"]["leksikon_hji"][topic_name]["links"] = []
            new_topic = ""
            # continue

    if dictionary["serbian"]["dicts"]["leksikon_hji"][topic_name]["text"]:
        dictionary["serbian"]["dicts"]["leksikon_hji"][topic_name]["text"] += f" {body}"
    else:
        dictionary["serbian"]["dicts"]["leksikon_hji"][topic_name]["text"] += body

count = 0
for i in dictionary["serbian"]["dicts"]["leksikon_hji"]:
    txt: str = dictionary["serbian"]["dicts"]["leksikon_hji"][i]["text"]

    txt = txt.replace("\t", " ")
    while True:
        txt = txt.replace("  ", " ")
        if txt.find("  ") == -1:
            break
    
    if txt != dictionary["serbian"]["dicts"]["leksikon_hji"][i]["text"]:
        count += 1

    dictionary["serbian"]["dicts"]["leksikon_hji"][i]["text"] = txt
print ()
print (f"Corrrected {count} items.")
    
print ()
print (f'Leksikon Hriscanstva, Judaizma i Islama: {len(dictionary["serbian"]["dicts"]["leksikon_hji"])} items.')





# Leksikon Lovackog Naoruzanja

with open("leksikon_lovackog_oruzja.txt", "r", encoding="utf-8") as file:
    txt = file.read()

txt_list = txt.split("\n")

dictionary["serbian"]["dicts"]["lov"] = {}

start = False
body = ""
topic_name = ""
header = False
new_topic = ""
counter = 0
for line in txt_list:
    line = line.strip().replace("\t", " ")
    while True:
        line = line.replace("  ", " ")
        if line.find("  ") == -1:
            break
    
    line = line.replace("@@@", "\n")

    if not line:
        continue

    pos = line.find(",")
    new_topic = line[:pos].strip()
    body = line[pos+1:].strip()

    if new_topic:
        topic_name = new_topic
        new_topic = ""
        if topic_name in dictionary["serbian"]["dicts"]["lov"]:
            dictionary["serbian"]["dicts"]["lov"][topic_name]["text"] += "\n\n\n"
            print (f"Duplicate Entry : {topic_name}")
        else:
            # print (topic_name)
            dictionary["serbian"]["dicts"]["lov"][topic_name] = {}
            dictionary["serbian"]["dicts"]["lov"][topic_name]["text"] = ""
            dictionary["serbian"]["dicts"]["lov"][topic_name]["links"] = []
            new_topic = ""
            # continue

    if dictionary["serbian"]["dicts"]["lov"][topic_name]["text"]:
        dictionary["serbian"]["dicts"]["lov"][topic_name]["text"] += f" {body}"
    else:
        dictionary["serbian"]["dicts"]["lov"][topic_name]["text"] += body


print ()
print (f'Leksikon Lovackog Naoruzanja: {len(dictionary["serbian"]["dicts"]["lov"])} items.')




# Spisak pojmova Polemologije

with open("leksikon_pojmova_polemologije.txt", "r", encoding="utf-8") as file:
    txt = file.read()

txt_list = txt.split("\n")

dictionary["serbian"]["dicts"]["polemologija"] = {}

start = False
body = ""
topic_name = ""
header = False
new_topic = ""
counter = 0
for line in txt_list:
    line = line.strip().replace("\t", " ")
    while True:
        line = line.replace("  ", " ")
        if line.find("  ") == -1:
            break
    
    line = line.replace("@@@", "\n")

    if not line:
        continue

    pos = line.find("-")
    new_topic = line[:pos].strip()
    body = line[pos+1:].strip()

    if new_topic:
        topic_name = new_topic
        new_topic = ""
        if topic_name in dictionary["serbian"]["dicts"]["polemologija"]:
            dictionary["serbian"]["dicts"]["polemologija"][topic_name]["text"] += "\n\n\n"
            print (f"Duplicate Entry : {topic_name}")
        else:
            # print (topic_name)
            dictionary["serbian"]["dicts"]["polemologija"][topic_name] = {}
            dictionary["serbian"]["dicts"]["polemologija"][topic_name]["text"] = ""
            dictionary["serbian"]["dicts"]["polemologija"][topic_name]["links"] = []
            new_topic = ""
            # continue

    if dictionary["serbian"]["dicts"]["polemologija"][topic_name]["text"]:
        dictionary["serbian"]["dicts"]["polemologija"][topic_name]["text"] += f" {body}"
    else:
        dictionary["serbian"]["dicts"]["polemologija"][topic_name]["text"] += body


print ()
print (f'Spisak pojmova Polemologije: {len(dictionary["serbian"]["dicts"]["polemologija"])} items.')



# Crven Ban

with open("crven_ban.txt", "r", encoding="utf-8") as file:
    txt = file.read()

txt_list = txt.split("\n")

dictionary["serbian"]["dicts"]["crven_ban"] = {}

start = False
body = ""
topic_name = ""
header = False
new_topic = ""
count = 0
counter = 0
vers_num = None
for line in txt_list:
    line = line.strip().replace("\t", " ")
    while True:
        line = line.replace("  ", " ")
        if line.find("  ") == -1:
            break
    
    line = line.replace("@@@", "\n")

    if not line:
        continue

    line_split = line.split(".")

    new_topic = ""
    body = ""

    if len(line_split) == 2:
        try:
            vers_num = int(line_split[0])
            count += 1
            tmp = line_split[1].strip(" '.,:-")
            new_topic = f'{count}. {tmp}'
            body = ""
        except:
            body = line
    else:
        body = line


    if new_topic:
        topic_name = new_topic
        new_topic = ""
        if topic_name in dictionary["serbian"]["dicts"]["crven_ban"]:
            dictionary["serbian"]["dicts"]["crven_ban"][topic_name]["text"] += "\n\n\n"
            print (f"Duplicate Entry : {topic_name}")
        else:
            # print (topic_name)
            dictionary["serbian"]["dicts"]["crven_ban"][topic_name] = {}
            dictionary["serbian"]["dicts"]["crven_ban"][topic_name]["text"] = ""
            dictionary["serbian"]["dicts"]["crven_ban"][topic_name]["links"] = []
            new_topic = ""
            # continue

    if dictionary["serbian"]["dicts"]["crven_ban"][topic_name]["text"]:
        dictionary["serbian"]["dicts"]["crven_ban"][topic_name]["text"] += f" {body}"
    else:
        dictionary["serbian"]["dicts"]["crven_ban"][topic_name]["text"] += body


print ()
print (f'Crven Ban: {len(dictionary["serbian"]["dicts"]["crven_ban"])} items.')




# Medicinski Leksikon

doc = Document("medicinski_leksikon.docx")

dictionary["serbian"]["dicts"]["medicina"] = {}

repl = [
    ["ñ", "đ"],
    ["manarski ", "mađarski "],
    ["prosvetitdjske ", "prosvetiteljske "],
    ["podmlanivanja ", "podmlađivanja "],
    [" Iek ", " lek "],
    ["ad 1.", "ad l."],
    [" tečenju ", " lečenju "],
    ["^", ""],
    ["V. ", "v. "],
    ["izogenskL", "izogenski"],
    ["f Ilija", "filija"],
    ["", ""],
    ["", ""],
    ["", ""],
    ["", ""],
    ["", ""],
    ["", ""],
    ["", ""],
    ["", ""]
]

start = False
body = ""
topic_name = "a"
header = False
new_topic = ""
counter = 0
eponim_mode = False
keep_topic = False
prepare_new_topic = False
lett_map = "abcčćdđefghijklmnoprsštuvwxyzž  "
let_pos = 0

for paragraph in doc.paragraphs:
    par_text: str = paragraph.text.strip().replace("\t", " ")

    if par_text == "BEOGRAD - ZAGREB":
        start = True
        continue

    if par_text.startswith("Abbott, Edville Gerhard (1870-1938)"):
        start = True
        eponim_mode = True

    if par_text.startswith("EPONIMNI I"):
        start = False

    if not start:
        continue

    if len(par_text) < 4 or par_text.find("Medicinski leksikon") != -1:
        # if par_text.find("Medicinski leksikon") != -1:
        #     print (f"Canceled entry: {par_text}")
        continue

    
    if eponim_mode:
        new_topic = ""
        body = ""
        for run in paragraph.runs:
            if run.bold:
                if par_text.startswith(new_topic + run.text):
                    new_topic += run.text
                else:
                    body += run.text
            else:
                body += run.text
        
        new_topic = new_topic.strip(" ,")
        body = body.strip()
        
        for j in repl:
            new_topic = new_topic.replace(j[0], j[1])
            body = body.replace(j[0], j[1])


        if new_topic:
            topic_name = new_topic
            new_topic = ""
            if topic_name in dictionary["serbian"]["dicts"]["medicina"]:
                dictionary["serbian"]["dicts"]["medicina"][topic_name]["text"] += "\n\n\n"
                print (f"Duplicate Entry : {topic_name}")
            else:
                # print (topic_name)
                dictionary["serbian"]["dicts"]["medicina"][topic_name] = {}
                dictionary["serbian"]["dicts"]["medicina"][topic_name]["text"] = ""
                dictionary["serbian"]["dicts"]["medicina"][topic_name]["links"] = []
                new_topic = ""
                # continue

        if dictionary["serbian"]["dicts"]["medicina"][topic_name]["text"]:
            dictionary["serbian"]["dicts"]["medicina"][topic_name]["text"] += f" {body}"
        else:
            dictionary["serbian"]["dicts"]["medicina"][topic_name]["text"] += body

        continue

    if keep_topic:
        new_topic = ""
        body = par_text
    else:
        pos = par_text.find(",")
        if pos > 0 and pos < 5:
            for idx, i in enumerate(lett_map):
                if i == par_text[0].lower():
                    if let_pos == idx:
                        prepare_new_topic = False
                    if prepare_new_topic:
                        prepare_new_topic = False
                        let_pos = idx
                        break
                    if let_pos != idx:
                        prepare_new_topic = True
        else:
            prepare_new_topic = False
        
        if par_text[0].lower() == lett_map[let_pos] or par_text[0].lower() == lett_map[let_pos+1]:
            new_topic = par_text[:pos].strip()
            body = par_text[pos+1:].strip()
        else:
            new_topic = ""
            body = par_text

    new_topic = new_topic.strip(" ,")
    body = body.strip()

    if body.endswith("-"):
        keep_topic = True
    else:
        keep_topic = False

    for j in repl:
        new_topic = new_topic.replace(j[0], j[1])
        body = body.replace(j[0], j[1])
    body = body.replace(";", ";\n")

    if new_topic:
        topic_name = new_topic
        new_topic = ""
        if topic_name in dictionary["serbian"]["dicts"]["medicina"]:
            dictionary["serbian"]["dicts"]["medicina"][topic_name]["text"] += "\n\n\n"
            print (f"Duplicate Entry : {topic_name}")
        else:
            # print (topic_name)
            dictionary["serbian"]["dicts"]["medicina"][topic_name] = {}
            dictionary["serbian"]["dicts"]["medicina"][topic_name]["text"] = ""
            dictionary["serbian"]["dicts"]["medicina"][topic_name]["links"] = []
            new_topic = ""
            # continue

    if dictionary["serbian"]["dicts"]["medicina"][topic_name]["text"]:
        dictionary["serbian"]["dicts"]["medicina"][topic_name]["text"] += f" {body}"
    else:
        dictionary["serbian"]["dicts"]["medicina"][topic_name]["text"] += body


repl1 = [
    [" - ", "@@@"],
    ["-  ", ""],
    ["- ", ""],
    ["@@@", " - "],
    ["", ""],
    ["", ""],
    ["", ""],
    ["", ""],
    ["", ""],
    ["", ""],
    ["", ""],
    ["", ""],
    ["", ""],
    ["", ""],
    ["", ""],
    ["", ""],
    ["", ""],
    ["", ""],
    ["", ""]
]

for i in dictionary["serbian"]["dicts"]["medicina"]:
    txt = dictionary["serbian"]["dicts"]["medicina"][i]["text"]
    for j in repl1:
        txt = txt.replace(j[0], j[1])
    while True:
        txt = txt.replace("  ", " ")
        if txt.find("  ") == -1:
            break
    dictionary["serbian"]["dicts"]["medicina"][i]["text"] = txt


print ()
print (f'Medicinski recnik: {len(dictionary["serbian"]["dicts"]["medicina"])} items.')






# Medicinski recnik ROGIC

doc = Document("medicinski_leksikon_Rogic_Dragan.docx")

dictionary["serbian"]["dicts"]["medicina_rogic"] = {}

repl = [
    ["ñ", "đ"],
    ["manarski ", "mađarski "],
    ["prosvetitdjske ", "prosvetiteljske "],
    ["podmlanivanja ", "podmlađivanja "],
    [" Iek ", " lek "],
    ["ad 1.", "ad l."],
    [" tečenju ", " lečenju "],
    ["^", ""],
    ["V. ", "v. "],
    ["izogenskL", "izogenski"],
    ["f Ilija", "filija"],
    ["", ""],
    ["", ""],
    ["", ""],
    ["", ""],
    ["", ""],
    ["", ""],
    ["", ""],
    ["", ""]
]

start = False
body = ""
topic_name = ""
header = False
new_topic = ""
counter = 0
for paragraph in doc.paragraphs:
    par_text: str = paragraph.text.strip().replace("\t", " ")

    if par_text == "MEDICINSKI LEKSIKON":
        if start is None:
            start = True
            continue
        if start is False:
            start = None

    if par_text == "PRILOZI":
        start = False

    if not start:
        continue

    if len(par_text) < 2:
        continue

    new_topic = ""
    body = ""

    if paragraph.style.font.size == 215900:
        new_topic = par_text
    else:
        body = par_text

    new_topic = new_topic.strip()
    body = body.strip()
    
    if not new_topic:
        if body.strip() == dictionary["serbian"]["dicts"]["medicina_rogic"][topic_name]["text"].strip():
            print (f"Canceled writig to dict, content already writen.  Topic: {topic_name}")
            continue

    body = body.replace(" - ", "@@@")
    body = body.replace("-", "")
    body = body.replace("@@@", " - ")

    for j in repl:
        new_topic = new_topic.replace(j[0], j[1])
        body = body.replace(j[0], j[1])
    body = body.replace(";", ";\n")

    if new_topic:
        topic_name = new_topic
        new_topic = ""
        if topic_name in dictionary["serbian"]["dicts"]["medicina_rogic"]:
            dictionary["serbian"]["dicts"]["medicina_rogic"][topic_name]["text"] += "\n\n\n"
            print (f"Duplicate Entry : {topic_name}")
        else:
            # print (topic_name)
            dictionary["serbian"]["dicts"]["medicina_rogic"][topic_name] = {}
            dictionary["serbian"]["dicts"]["medicina_rogic"][topic_name]["text"] = ""
            dictionary["serbian"]["dicts"]["medicina_rogic"][topic_name]["links"] = []
            new_topic = ""
            # continue

    if dictionary["serbian"]["dicts"]["medicina_rogic"][topic_name]["text"]:
        dictionary["serbian"]["dicts"]["medicina_rogic"][topic_name]["text"] += f" {body}"
    else:
        dictionary["serbian"]["dicts"]["medicina_rogic"][topic_name]["text"] += body

repl1 = [
    [" - ", "@@@"],
    ["-  ", ""],
    ["- ", ""],
    ["@@@", " - "],
    ["", ""],
    ["", ""],
    ["", ""],
    ["", ""],
    ["", ""],
    ["", ""],
    ["", ""],
    ["", ""],
    ["", ""],
    ["", ""],
    ["", ""],
    ["", ""],
    ["", ""],
    ["", ""],
    ["", ""]
]

for i in dictionary["serbian"]["dicts"]["medicina_rogic"]:
    txt = dictionary["serbian"]["dicts"]["medicina_rogic"][i]["text"]
    for j in repl1:
        txt = txt.replace(j[0], j[1])
    while True:
        txt = txt.replace("  ", " ")
        if txt.find("  ") == -1:
            break
    dictionary["serbian"]["dicts"]["medicina_rogic"][i]["text"] = txt


print ()
print (f'Medicinski recnik ROGIC: {len(dictionary["serbian"]["dicts"]["medicina_rogic"])} items.')




# Naratoloski recnik


doc = Document("naratoloski_recnik_dzerald_prins.docx")

dictionary["serbian"]["dicts"]["narat"] = {}

start = False
body = ""
topic_name = ""
header = False
new_topic = ""
counter = 0
repl = [
    ["^", ""],
    [" - ", "@@@"],
    ["-   ", ""],
    ["-  ", ""],
    ["- ", ""],
    ["@@@", " - "],
    ["", ""]
]
for paragraph in doc.paragraphs:
    par_text: str = paragraph.text.strip()

    if par_text.startswith("@@@rečnik"):
        start = True
        continue

    if par_text == "BIBLIOGRAFIJA":
        start = False

    if not start:
        continue

    if len(par_text) < 2:
        continue

    if paragraph.style.font.size == 95250:
        continue

    new_topic = ""
    body = ""

    if paragraph.style.font.bold:
        new_topic = par_text
        body = ""
    else:
        new_topic = ""
        body = par_text

    # if par_text.startswith("čitljivi"):
    #     print ()

    for i in repl:
        new_topic = new_topic.replace(i[0], i[1])
        body = body.replace(i[0], i[1])
    
    new_topic = new_topic.strip(" ,.:-=")
    body = body.strip()

    if new_topic:
        topic_name = new_topic
        new_topic = ""
        if topic_name in dictionary["serbian"]["dicts"]["narat"]:
            dictionary["serbian"]["dicts"]["narat"][topic_name]["text"] += "\n\n\n"
            print (f"Duplicate Entry : {topic_name}")
        else:
            # print (topic_name)
            dictionary["serbian"]["dicts"]["narat"][topic_name] = {}
            dictionary["serbian"]["dicts"]["narat"][topic_name]["text"] = ""
            dictionary["serbian"]["dicts"]["narat"][topic_name]["links"] = []
            new_topic = ""
            # continue

    if dictionary["serbian"]["dicts"]["narat"][topic_name]["text"]:
        dictionary["serbian"]["dicts"]["narat"][topic_name]["text"] += f" {body}"
    else:
        dictionary["serbian"]["dicts"]["narat"][topic_name]["text"] += body


print ()
print (f'Naratoloski recnik: {len(dictionary["serbian"]["dicts"]["narat"])} items.')




# Recnik Latinskih Imena


doc = Document("latinski_recnik_imena.docx")

dictionary["serbian"]["dicts"]["latin"] = {}

start = False
body = ""
topic_name = ""
header = False
new_topic = ""
counter = 0
last_body = ""
for paragraph in doc.paragraphs:
    par_text: str = paragraph.text.strip()

    if par_text == "@@@A":
        start = True
        continue

    if par_text == "@@@END":
        start = False
    
    if not start:
        continue

    if len(par_text) < 4:
        continue

    new_topic = ""
    body = ""

    if last_body and last_body[-1] == "-":
        body = par_text
    else:
        for run in paragraph.runs:
            if run.bold:
                if par_text.startswith(new_topic + run.text):
                    new_topic += run.text
                else:
                    body += run.text
            else:
                body += run.text

    new_topic = new_topic.strip(" ,.:=-")
    body = body.strip(" ,.:=").replace(";", ";\n")
    last_body = body

    if new_topic:
        topic_name = new_topic
        new_topic = ""
        if topic_name in dictionary["serbian"]["dicts"]["latin"]:
            dictionary["serbian"]["dicts"]["latin"][topic_name]["text"] += "\n\n\n"
            print (f"Duplicate Entry : {topic_name}")
        else:
            # print (topic_name)
            dictionary["serbian"]["dicts"]["latin"][topic_name] = {}
            dictionary["serbian"]["dicts"]["latin"][topic_name]["text"] = ""
            dictionary["serbian"]["dicts"]["latin"][topic_name]["links"] = []
            new_topic = ""
            # continue

    if dictionary["serbian"]["dicts"]["latin"][topic_name]["text"]:
        dictionary["serbian"]["dicts"]["latin"][topic_name]["text"] += f" {body}"
    else:
        dictionary["serbian"]["dicts"]["latin"][topic_name]["text"] += body



with open("latinska_vlastita_imena.txt", "r", encoding="utf-8") as file:
    txt = file.read()

txt_list = [x.strip() for x in txt.split("\n") if x.strip() != ""]

spliters = [" f. ", " m. ", " n. ", "@@@"]
for line in txt_list:
    line = line.replace("\t", " ")
    line = line.replace(" f ", " f. ")
    line = line.replace(" m ", " m. ")
    line = line.replace(" n ", " n. ")

    count_f = line.count(" f. ")
    count_m = line.count(" m. ")
    count_n = line.count(" n. ")
    count_s = line.count("@@@")
    count_t = count_f + count_m + count_n + count_s

    if count_t != 1:
        print (f"Warning. Too many elements : {line}")
        continue
    
    line_split = []
    for i in spliters:
        if line.find(i) != -1:
            line_split = line.split(i)
            break
    
    if i == "@@@":
        new_topic = line_split[0]
    else:
        new_topic = line_split[0] + i
    
    new_topic = new_topic.strip()
    body = line_split[1].strip()

    if new_topic:
        topic_name = new_topic
        new_topic = ""
        if topic_name in dictionary["serbian"]["dicts"]["latin"]:
            dictionary["serbian"]["dicts"]["latin"][topic_name]["text"] += "\n\n\n"
            print (f"Duplicate Entry : {topic_name}")
        else:
            # print (topic_name)
            dictionary["serbian"]["dicts"]["latin"][topic_name] = {}
            dictionary["serbian"]["dicts"]["latin"][topic_name]["text"] = ""
            dictionary["serbian"]["dicts"]["latin"][topic_name]["links"] = []
            new_topic = ""
            # continue

    if dictionary["serbian"]["dicts"]["latin"][topic_name]["text"]:
        dictionary["serbian"]["dicts"]["latin"][topic_name]["text"] += f" {body}"
    else:
        dictionary["serbian"]["dicts"]["latin"][topic_name]["text"] += body



print ()
print (f'Recnik Latinskih Imena: {len(dictionary["serbian"]["dicts"]["latin"])} items.')





# Recnik novijih anglicizama

def anglicizmi_fix_text_case(txt: str, title: bool = False) -> str:
    original_txt = txt
    original_txt_list = [x for x in original_txt.split("\t")]
    if len(original_txt_list) > 1:
        try:
            _ = int(original_txt_list[0])
            original_txt_list.pop(0)
            original_txt = " ".join(original_txt_list)
        except ValueError:
            _ = 0

        try:
            _ = int(original_txt_list[-1])
            original_txt_list.pop(len(original_txt_list)-1)
            original_txt = " ".join(original_txt_list)
        except ValueError:
            _ = 0
        
    repl = "!@#$%^&*()_+=-[]}{';\":,./?><\\\t\n0123456789"
    for i in repl:
        txt = txt.replace(i, " ")
    
    pos = 0
    while True:
        if pos >= len(txt):
            break
        if txt[pos] == " ":
            pos += 1
            continue
        end = txt.find(" ", pos)
        if end == -1:
            break
        word = txt[pos:end]
        if len(word) < 2:
            pos = end
            continue

        if word[1:].upper() == word[1:] or word[1:].lower() == word[1:]:
            pos = end
            continue
        
        count_up = 0
        count_low = 0
        for i in word:
            if i in "QWERTYUIOPLKJHGFDSAZXCVBNMŠĐČĆŽ":
                count_up += 1
            else:
                count_low += 1

        if count_up >= count_low:
            word = word.upper()
        else:
            word = word.lower()
        
        original_txt = original_txt[:pos] + word + original_txt[end:]
        pos = end
    
    pos = 0
    cr_map = []
    while True:
        pos = original_txt.find("/", pos)
        if pos == -1:
            break
        end = original_txt.find("/", pos+1)
        if end == -1:
            break
        
        cr_map.append(pos)
        cr_map.append(end + 1)
        pos = end + 1

    count = 0
    for i in cr_map:
        original_txt = original_txt[0:i + count] + "\n" + original_txt[i + count:]
        count += 1
        
    if title:
        original_txt = original_txt.replace("\n", "")

    return original_txt

def anglicizmi_is_valid_title(txt: str) -> bool:
    is_valid = True

    letter_pool = "QWERTYUIOPLKJHGFDSAZXCVBNMŠĐČĆŽ" + "QWERTYUIOPLKJHGFDSAZXCVBNMŠĐČĆŽ".lower()

    for i in letter_pool:
        if i in txt:
            return True
    
    return False

doc = Document("Srpski_recnik_novijih_anglicizama_Prvo_e2.docx")

dictionary["serbian"]["dicts"]["anglicizmi"] = {}

start = False
body = ""
topic_name = ""
header = False
new_topic = ""
counter = 0
protect_topic = False
cancel_writing = False
subtitle_group_mode = False
subtitle_mode = False
for paragraph in doc.paragraphs:
    par_text: str = anglicizmi_fix_text_case(paragraph.text.strip())

    if par_text == "@@@START":
        start = True
        continue

    if par_text == "@@@END":
        start = False

    if not start:
        continue

    if not par_text:
        if topic_name:
            if dictionary["serbian"]["dicts"]["anglicizmi"][topic_name]["text"][-1] != "\n":
                dictionary["serbian"]["dicts"]["anglicizmi"][topic_name]["text"] += "\n"

    if len(par_text) < 5:
        continue

    new_topic = ""
    body = ""
    subtitle_group_mode = False
    subtitle_mode = False
    if par_text.find("") != -1:
        par_text_split = par_text.split("")
        if len(par_text_split) != 2:
            print ("Error. Too many splits: ", par_text)
        
        if par_text_split[0].find("\t") == -1:
            new_topic = par_text_split[0].strip(" ")
        else:
            new_topic = par_text_split[0][par_text_split[0].find("\t")+1:].strip(" ")
        
        if par_text_split[1].find("\t") == -1:
            body = "v." + par_text_split[1]
        else:
            body = par_text_split[1][:par_text_split[1].find("\t")].strip(" ")

    else:
        for run in paragraph.runs:
            if run.font.size == 120650:
                # print ("Cancelled............................: ", run.text)
                continue
            run_txt = anglicizmi_fix_text_case(run.text)

            if run.font.subscript:
                run_txt = f"({run_txt})"

            if run.bold:
                if run.font.color.rgb[2] < 100:
                    if subtitle_group_mode:
                        body += f"@@@01{run_txt.upper()}@@@01"
                    else:
                        body += f"\n@@@01{run_txt.upper()}@@@01"
                    subtitle_group_mode = True
                    protect_topic = True
                    continue
                else:
                    subtitle_group_mode = False

                if run.font.size == 152400:
                    protect_topic = False

                if protect_topic:
                    if subtitle_mode:
                        body += f"@@@00{run_txt}@@@00"
                    else:
                        body += f"\n@@@00{run_txt}@@@00"
                    subtitle_mode = True
                    continue
                else:
                    subtitle_mode = False
                
                compare_1 = new_topic + run_txt
                compare_2 = new_topic.replace("(", "").replace(")", "") + run_txt.replace("(", "").replace(")", "")
                
                if par_text.lower().startswith(compare_1.lower()) or par_text.lower().startswith(compare_2.lower()):
                    cancel_writing = False
                    new_topic += run_txt
                else:
                    body += run_txt
            else:
                if paragraph.style.font.size == 127000:
                    body += f"@@@10{run_txt}@@@10"
                else:
                    body += run_txt

    if body.upper().startswith("SKRAĆENICE I SIMBOLI"):
        cancel_writing = True
    if body.endswith("dug nenaglašen samoglasnik"):
        cancel_writing = False
        continue
    if cancel_writing:
        continue

    new_topic = new_topic.strip(" ,/-=:;")
    new_topic = anglicizmi_fix_text_case(new_topic, title=True)
    
    body = anglicizmi_fix_text_case(body)

    if not anglicizmi_is_valid_title(new_topic):
        body = f"\n{new_topic} {body}"
        new_topic = ""

    body = body.strip(" ,")
    body = body.replace("\t", "\n")
    body = body.replace("[", "\n[")
    body = body.replace("]", "]\n")
    body = body.replace("\n\n", "\n")
    body = body.replace("\n ", "\n")
    for i in range(20):
        list_number = f"{i} "
        if body.startswith(list_number):
            body = "\n" + body
    
    if new_topic:
        topic_name = new_topic
        new_topic = ""
        if topic_name in dictionary["serbian"]["dicts"]["anglicizmi"]:
            dictionary["serbian"]["dicts"]["anglicizmi"][topic_name]["text"] += "\n\n\n"
            print (f"Duplicate Entry : {topic_name}")
        else:
            # print (topic_name)
            dictionary["serbian"]["dicts"]["anglicizmi"][topic_name] = {}
            dictionary["serbian"]["dicts"]["anglicizmi"][topic_name]["text"] = ""
            dictionary["serbian"]["dicts"]["anglicizmi"][topic_name]["links"] = []
            new_topic = ""
            # continue

    if dictionary["serbian"]["dicts"]["anglicizmi"][topic_name]["text"] and dictionary["serbian"]["dicts"]["anglicizmi"][topic_name]["text"][-1] != "\n":
        dictionary["serbian"]["dicts"]["anglicizmi"][topic_name]["text"] += f" {body}"
    else:
        dictionary["serbian"]["dicts"]["anglicizmi"][topic_name]["text"] += body


print ()
print (f'Recnik novijih anglicizama: {len(dictionary["serbian"]["dicts"]["anglicizmi"])} items.')




# Onkoloski recnik


doc = Document("onkoloski-recnik.docx")

dictionary["serbian"]["dicts"]["onkoloski"] = {}

start = False
body = ""
topic_name = ""
header = False
new_topic = ""
counter = 0
for paragraph in doc.paragraphs:
    par_text: str = paragraph.text.strip()

    if par_text == "@@@START":
        start = True
        continue

    if par_text == "@@@END":
        start = False

    if not start:
        continue

    if len(par_text) < 3:
        continue

    new_topic = ""
    body = ""
    for run in paragraph.runs:
        if run.bold:
            if par_text.startswith(new_topic + run.text):
                new_topic += run.text
            else:
                print (run.text)
                body += run.text
        else:
            body += run.text

    new_topic = new_topic.strip(" :")
    body = body.lstrip(":").strip()

    if new_topic:
        topic_name = new_topic
        new_topic = ""
        if topic_name in dictionary["serbian"]["dicts"]["onkoloski"]:
            dictionary["serbian"]["dicts"]["onkoloski"][topic_name]["text"] += "\n\n\n"
            print (f"Duplicate Entry : {topic_name}")
        else:
            # print (topic_name)
            dictionary["serbian"]["dicts"]["onkoloski"][topic_name] = {}
            dictionary["serbian"]["dicts"]["onkoloski"][topic_name]["text"] = ""
            dictionary["serbian"]["dicts"]["onkoloski"][topic_name]["links"] = []
            new_topic = ""
            # continue

    if dictionary["serbian"]["dicts"]["onkoloski"][topic_name]["text"]:
        dictionary["serbian"]["dicts"]["onkoloski"][topic_name]["text"] += f" {body}"
    else:
        dictionary["serbian"]["dicts"]["onkoloski"][topic_name]["text"] += body


print ()
print (f'Onkoloski recnik: {len(dictionary["serbian"]["dicts"]["onkoloski"])} items.')






# Pravoslavni pojmovnik

with open("Pravoslavni_pojmovnik.txt", "r", encoding="utf-8") as file:
    txt = file.read()

txt_list = txt.split("\n")

dictionary["serbian"]["dicts"]["pravoslavni_pojmovnik"] = {}

start = False
body = ""
topic_name = ""
header = False
new_topic = ""
counter = 0
for line in txt_list:
    line = cirilica_u_latinicu(line.strip())

    if not line:
        continue
    if line.startswith("Slovo "):
        continue

    pos = line.find("-")
    if pos == -1:
        print ("Error, incorect line: ", line)
        continue

    new_topic = line[:pos].strip()
    body = line[pos+1:].strip()

    if new_topic:
        topic_name = new_topic
        new_topic = ""
        if topic_name in dictionary["serbian"]["dicts"]["pravoslavni_pojmovnik"]:
            dictionary["serbian"]["dicts"]["pravoslavni_pojmovnik"][topic_name]["text"] += "\n\n\n"
            print (f"Duplicate Entry : {topic_name}")
        else:
            # print (topic_name)
            dictionary["serbian"]["dicts"]["pravoslavni_pojmovnik"][topic_name] = {}
            dictionary["serbian"]["dicts"]["pravoslavni_pojmovnik"][topic_name]["text"] = ""
            dictionary["serbian"]["dicts"]["pravoslavni_pojmovnik"][topic_name]["links"] = []
            new_topic = ""
            # continue

    if dictionary["serbian"]["dicts"]["pravoslavni_pojmovnik"][topic_name]["text"]:
        dictionary["serbian"]["dicts"]["pravoslavni_pojmovnik"][topic_name]["text"] += f" {body}"
    else:
        dictionary["serbian"]["dicts"]["pravoslavni_pojmovnik"][topic_name]["text"] += body


print ()
print (f'Pravoslavni pojmovnik: {len(dictionary["serbian"]["dicts"]["pravoslavni_pojmovnik"])} items.')





# Recnik svetske mitologije Artur Koterel


with open("kuran.txt", "r", encoding="utf-8") as file:
    txt = file.read()

txt_list = [x for x in txt.split("\n")]

dictionary["serbian"]["dicts"]["kuran"] = {}

start = False
body = ""
topic_name = ""
header = False
new_topic = ""
counter = 0
sura_mode = False
sura_title = ""
hatma_dova_mode = False
ajet_num = None
for par_text in txt_list:
    par_text = par_text.replace("\t", " ").strip()
    while True:
        par_text = par_text.replace("  ", " ")
        if par_text.find("  ") == -1:
            break

    if par_text == "@@@START":
        start = True
        continue

    if par_text == "@@@END":
        hatma_dova_mode = True
        continue

    if not start:
        continue

    if hatma_dova_mode:
        topic_name = "HATMA-DOVA"
        if not par_text:
            body = "\n"
        else:
            body = par_text + "\n"

        if topic_name not in dictionary["serbian"]["dicts"]["kuran"]:
            dictionary["serbian"]["dicts"]["kuran"][topic_name] = {}
            dictionary["serbian"]["dicts"]["kuran"][topic_name]["text"] = ""
            dictionary["serbian"]["dicts"]["kuran"][topic_name]["links"] = []

        dictionary["serbian"]["dicts"]["kuran"][topic_name]["text"] += f" {body}"
        continue
        
    if not par_text:
        continue

    new_topic = ""
    body = ""

    par_text_split = [x for x in par_text.split(" ") if x != ""]
    if par_text_split[0] == "Sura":
        sura_mode = True
        sura_title = par_text + " - "
        ajet_num = None
        new_topic = par_text
        body = ""
    
    if par_text.startswith("1. "):
        sura_mode = False

    if sura_mode and not new_topic:
        body = par_text + "\n"
    
    if not sura_mode:
        if ajet_num is None:
            ajet_num = 0

        is_new_ajet = True
        try:
            _ = float(par_text_split[0])
        except:
            is_new_ajet = False

        if is_new_ajet:
            ajet_num += 1
            new_topic = f"{sura_title}Ajet {ajet_num}."
            body = par_text + "\n"
        else:
            new_topic = ""
            body = par_text + "\n"



    if new_topic:
        topic_name = new_topic
        new_topic = ""
        if topic_name in dictionary["serbian"]["dicts"]["kuran"]:
            dictionary["serbian"]["dicts"]["kuran"][topic_name]["text"] += "\n\n\n"
            print (f"Duplicate Entry : {topic_name}")
        else:
            # print (topic_name)
            dictionary["serbian"]["dicts"]["kuran"][topic_name] = {}
            dictionary["serbian"]["dicts"]["kuran"][topic_name]["text"] = ""
            dictionary["serbian"]["dicts"]["kuran"][topic_name]["links"] = []
            new_topic = ""
            # continue

    if dictionary["serbian"]["dicts"]["kuran"][topic_name]["text"]:
        dictionary["serbian"]["dicts"]["kuran"][topic_name]["text"] += f" {body}"
    else:
        dictionary["serbian"]["dicts"]["kuran"][topic_name]["text"] += body


print ()
print (f'Kuran: {len(dictionary["serbian"]["dicts"]["kuran"])} items.')




# Recnik arhitektonskog projektovanja


doc = Document("recnik_arhitektonskog_projektovanja.docx")

dictionary["serbian"]["dicts"]["arhitekt"] = {}

start = False
body = ""
topic_name = ""
header = False
new_topic = ""
counter = 0
for paragraph in doc.paragraphs:
    par_text: str = paragraph.text.strip()

    if par_text == "@@@START":
        start = True
        continue

    if par_text == "@@@END":
        start = False

    if not start:
        continue

    if len(par_text) < 2:
        continue

    new_topic = ""
    body = ""
    new_topic_list = []
    can_new = False
    for run in paragraph.runs:
        if run.bold:
            if can_new:
                new_topic_list.append([new_topic, body])
                new_topic = ""
                body = ""
            can_new = False
            new_topic += run.text
        else:
            can_new = True
            body += run.text
    if can_new:
        new_topic_list.append([new_topic, body])


    for i in new_topic_list:
        new_topic = i[0]
        for j in new_topic.split(" "):
            if new_topic != new_topic.upper():
                print ("Warrning. Title too long: ", new_topic)
        
        if new_topic.find("INTERSEKCIJA ARHITEKTONSKIH ELEMENATA") != -1:
            print ()

        body = i[1]

        new_topic = new_topic.strip()
        body = body.strip()

        if new_topic:
            topic_name = new_topic
            new_topic = ""
            if topic_name in dictionary["serbian"]["dicts"]["arhitekt"]:
                dictionary["serbian"]["dicts"]["arhitekt"][topic_name]["text"] += "\n\n\n"
                print (f"Duplicate Entry : {topic_name}")
            else:
                # print (topic_name)
                dictionary["serbian"]["dicts"]["arhitekt"][topic_name] = {}
                dictionary["serbian"]["dicts"]["arhitekt"][topic_name]["text"] = ""
                dictionary["serbian"]["dicts"]["arhitekt"][topic_name]["links"] = []
                new_topic = ""
                # continue

        if dictionary["serbian"]["dicts"]["arhitekt"][topic_name]["text"]:
            dictionary["serbian"]["dicts"]["arhitekt"][topic_name]["text"] += f"\n{body}"
        else:
            dictionary["serbian"]["dicts"]["arhitekt"][topic_name]["text"] += body


print ()
print (f'Recnik arhitektonskog projektovanja: {len(dictionary["serbian"]["dicts"]["arhitekt"])} items.')





# Recnik latinskog jezika


doc = Document("recnik_latinskog_jezika.docx")

dictionary["serbian"]["dicts"]["latin2"] = {}

start = True
body = ""
topic_name = ""
header = False
new_topic = ""
counter = 0
for paragraph in doc.paragraphs:
    par_text: str = paragraph.text.strip()

    if par_text == "@@@START":
        start = True
        continue

    if par_text == "@@@END":
        start = False

    if not start:
        continue

    if len(par_text) < 2:
        continue

    
    if par_text.find("Scīpĭo") != -1:
        header = True
        print ()

    if header:
        bool1 = par_text.find('\t')!=-1
        bool2 = par_text.find('\n')!=-1
        # print (f"Has tab: {bool1},   Has ENTER: {bool2},   {par_text}")

    par_text_split = par_text.split("\t")

    
    new_topic = par_text_split[0]
    body = par_text_split[1]

    # for run in paragraph.runs:
    #     if run.bold:
    #         new_topic = run.text
    #     else:
    #         body += run.text

    new_topic = new_topic.strip()
    body = body.strip()

    if new_topic:
        topic_name = new_topic
        new_topic = ""
        if topic_name in dictionary["serbian"]["dicts"]["latin2"]:
            dictionary["serbian"]["dicts"]["latin2"][topic_name]["text"] += "\n\n\n"
            print (f"Duplicate Entry : {topic_name}")
        else:
            # print (topic_name)
            dictionary["serbian"]["dicts"]["latin2"][topic_name] = {}
            dictionary["serbian"]["dicts"]["latin2"][topic_name]["text"] = ""
            dictionary["serbian"]["dicts"]["latin2"][topic_name]["links"] = []
            new_topic = ""
            # continue

    if dictionary["serbian"]["dicts"]["latin2"][topic_name]["text"]:
        dictionary["serbian"]["dicts"]["latin2"][topic_name]["text"] += f" {body}"
    else:
        dictionary["serbian"]["dicts"]["latin2"][topic_name]["text"] += body


print ()
print (f'Recnik latinskog jezika: {len(dictionary["serbian"]["dicts"]["latin2"])} items.')






# Recnik Pirotskog Govora


doc = Document("Recnik_pirotskog_govora.docx")

dictionary["serbian"]["dicts"]["pirot"] = {}

start = False
body = ""
topic_name = ""
header = False
new_topic = ""
counter = 0
for paragraph in doc.paragraphs:
    par_text: str = cirilica_u_latinicu(paragraph.text.replace("\t", "\n").strip())

    if par_text == "@@@START":
        start = True
        continue

    if par_text == "@@@END":
        start = False

    if not start:
        continue

    if len(par_text) < 4:
        continue

    new_topic = ""
    body = ""
    for run in paragraph.runs:
        run_txt = cirilica_u_latinicu(run.text)
        if run.bold:
            if par_text.startswith(new_topic + run_txt):
                new_topic += run_txt
            else:
                body += run_txt
        else:
            body += run_txt

    new_topic = new_topic.strip(" ,:-.;")
    body = body.strip(" ,:;=-.\t")
    for i in range(20):
        body = body.replace(f" {i}. ", f"\n{i}. ")
    body = body.replace("(", "\n     (").replace("\t", "\n")

    if new_topic:
        topic_name = new_topic
        new_topic = ""
        if topic_name in dictionary["serbian"]["dicts"]["pirot"]:
            dictionary["serbian"]["dicts"]["pirot"][topic_name]["text"] += "\n\n\n"
            print (f"Duplicate Entry : {topic_name}")
        else:
            # print (topic_name)
            dictionary["serbian"]["dicts"]["pirot"][topic_name] = {}
            dictionary["serbian"]["dicts"]["pirot"][topic_name]["text"] = ""
            dictionary["serbian"]["dicts"]["pirot"][topic_name]["links"] = []
            new_topic = ""
            # continue

    if dictionary["serbian"]["dicts"]["pirot"][topic_name]["text"]:
        dictionary["serbian"]["dicts"]["pirot"][topic_name]["text"] += f" {body}"
    else:
        dictionary["serbian"]["dicts"]["pirot"][topic_name]["text"] += body


txt_list = txt_pirot.split("\n")
for line in txt_list:
    line_split = line.split("\t")
    if len(line_split) != 2:
        print ("Error. Line: ", line)

    new_topic = f"SKRAĆENICA: {line_split[0]}"
    body = line_split[1]
    if new_topic:
        topic_name = new_topic
        new_topic = ""
        if topic_name in dictionary["serbian"]["dicts"]["pirot"]:
            dictionary["serbian"]["dicts"]["pirot"][topic_name]["text"] += "\n\n\n"
            print (f"Duplicate Entry : {topic_name}")
        else:
            # print (topic_name)
            dictionary["serbian"]["dicts"]["pirot"][topic_name] = {}
            dictionary["serbian"]["dicts"]["pirot"][topic_name]["text"] = ""
            dictionary["serbian"]["dicts"]["pirot"][topic_name]["links"] = []
            new_topic = ""
            # continue

    if dictionary["serbian"]["dicts"]["pirot"][topic_name]["text"]:
        dictionary["serbian"]["dicts"]["pirot"][topic_name]["text"] += f" {body}"
    else:
        dictionary["serbian"]["dicts"]["pirot"][topic_name]["text"] += body

print ()
print (f'Recnik Pirotskog Govora: {len(dictionary["serbian"]["dicts"]["pirot"])} items.')





# Recnik pravnih termina za novinare

with open("recnik_prvanih_termina_za_novinare.txt", "r", encoding="utf-8") as file:
    txt = file.read()

txt_list = txt.split("\n")

dictionary["serbian"]["dicts"]["pravni_novinar"] = {}

start = False
body = ""
topic_name = ""
header = False
new_topic = ""
counter = 0
for line in txt_list:
    line = line.strip().replace("\t", " ")

    if not line:
        continue

    pos = line.find("-")
    if pos == -1:
        print ("Error. No delimiter. LINE: ", line)
        continue
    # if line.count("-") > 1:
    #     print ("Error. Too many delimiters. LINE: ", line)

    new_topic = line[:pos].strip()
    body = line[pos+1:].strip()

    if new_topic:
        topic_name = new_topic
        new_topic = ""
        if topic_name in dictionary["serbian"]["dicts"]["pravni_novinar"]:
            dictionary["serbian"]["dicts"]["pravni_novinar"][topic_name]["text"] += "\n\n\n"
            print (f"Duplicate Entry : {topic_name}")
        else:
            print (topic_name)
            dictionary["serbian"]["dicts"]["pravni_novinar"][topic_name] = {}
            dictionary["serbian"]["dicts"]["pravni_novinar"][topic_name]["text"] = ""
            dictionary["serbian"]["dicts"]["pravni_novinar"][topic_name]["links"] = []
            new_topic = ""
            # continue

    if dictionary["serbian"]["dicts"]["pravni_novinar"][topic_name]["text"]:
        dictionary["serbian"]["dicts"]["pravni_novinar"][topic_name]["text"] += f" {body}"
    else:
        dictionary["serbian"]["dicts"]["pravni_novinar"][topic_name]["text"] += body


print ()
print (f'Recnik pravnih termina za novinare: {len(dictionary["serbian"]["dicts"]["pravni_novinar"])} items.')




# Recnik Poslovica i Antiposlovica


doc = Document("recnik_srpskih_antiposlovica.docx")

dictionary["serbian"]["dicts"]["poslovice"] = {}

start = False
body = ""
topic_name = ""
header = False
new_topic = ""
counter = 0
poslovica_mode = False
repl = [
    ["-   ", ""],
    ["-  ", ""],
    ["- ", ""],
    ["\t", " "]
]
for paragraph in doc.paragraphs:
    par_text: str = paragraph.text.strip()

    if par_text == "@@@START":
        start = True
        continue

    if par_text == "@@@END":
        poslovica_mode = True
        start = False

    if not start:
        continue

    if not par_text:
        continue

    for i in repl:
        par_text = par_text.replace(i[0], i[1])

    new_topic = ""
    body = ""


    if poslovica_mode:
        if paragraph.style.name.startswith("Heading 4"):
            new_topic = par_text
            body = ""
        else:
            new_topic = ""
            body = par_text + "\n"

        if new_topic:
            topic_name = new_topic
            new_topic = ""
            if topic_name in dictionary["serbian"]["dicts"]["poslovice"]:
                dictionary["serbian"]["dicts"]["poslovice"][topic_name]["text"] += "\n\n\n"
                print (f"Duplicate Entry : {topic_name}")
            else:
                # print (topic_name)
                dictionary["serbian"]["dicts"]["poslovice"][topic_name] = {}
                dictionary["serbian"]["dicts"]["poslovice"][topic_name]["text"] = ""
                dictionary["serbian"]["dicts"]["poslovice"][topic_name]["links"] = []
                new_topic = ""
                # continue

        if dictionary["serbian"]["dicts"]["poslovice"][topic_name]["text"]:
            dictionary["serbian"]["dicts"]["poslovice"][topic_name]["text"] += f" {body}"
        else:
            dictionary["serbian"]["dicts"]["poslovice"][topic_name]["text"] += body
        continue


    if paragraph.style.name.startswith("Heading 2"):
        new_topic = par_text
        body = ""
    else:
        new_topic = ""
        body = par_text + "\n"

    if paragraph.style.name.startswith("List"):
        body = f"@@@00{body}@@@00\n"
    elif paragraph.style.name.startswith("Body") and paragraph.style.font.size == 127000:
        body = f"@@@01{body}@@@01\n"
    
    if new_topic:
        topic_name = new_topic
        new_topic = ""
        if topic_name in dictionary["serbian"]["dicts"]["poslovice"]:
            dictionary["serbian"]["dicts"]["poslovice"][topic_name]["text"] += "\n\n\n"
            print (f"Duplicate Entry : {topic_name}")
        else:
            # print (topic_name)
            dictionary["serbian"]["dicts"]["poslovice"][topic_name] = {}
            dictionary["serbian"]["dicts"]["poslovice"][topic_name]["text"] = ""
            dictionary["serbian"]["dicts"]["poslovice"][topic_name]["links"] = []
            new_topic = ""
            # continue

    if dictionary["serbian"]["dicts"]["poslovice"][topic_name]["text"]:
        dictionary["serbian"]["dicts"]["poslovice"][topic_name]["text"] += f" {body}"
    else:
        dictionary["serbian"]["dicts"]["poslovice"][topic_name]["text"] += body


print ()
print (f'Recnik Poslovica i Antiposlovica: {len(dictionary["serbian"]["dicts"]["poslovice"])} items.')





# Recnik Turcizama u Srpskokm Jeziku

with open("turcizmi_u_srpskom_jeziku.txt", "r", encoding="utf-8") as file:
    txt = file.read()

txt_list = txt.split("\n")

dictionary["serbian"]["dicts"]["turcizmi"] = {}

start = False
body = ""
topic_name = ""
header = False
new_topic = ""
counter = 0
for line in txt_list:
    line = cirilica_u_latinicu(line[line.find("\t")+1:].strip().replace("\t", " "))

    if not line:
        continue

    pos = line.find("--")

    if pos == -1 or line.count("--") > 1:
        print ()

    new_topic = line[:pos].strip().upper()
    body = line[pos+2:].strip()


    if new_topic:
        topic_name = new_topic
        new_topic = ""
        if topic_name in dictionary["serbian"]["dicts"]["turcizmi"]:
            dictionary["serbian"]["dicts"]["turcizmi"][topic_name]["text"] += "\n\n\n"
            print (f"Duplicate Entry : {topic_name}")
        else:
            # print (topic_name)
            dictionary["serbian"]["dicts"]["turcizmi"][topic_name] = {}
            dictionary["serbian"]["dicts"]["turcizmi"][topic_name]["text"] = ""
            dictionary["serbian"]["dicts"]["turcizmi"][topic_name]["links"] = []
            new_topic = ""
            # continue

    if dictionary["serbian"]["dicts"]["turcizmi"][topic_name]["text"]:
        dictionary["serbian"]["dicts"]["turcizmi"][topic_name]["text"] += f" {body}"
    else:
        dictionary["serbian"]["dicts"]["turcizmi"][topic_name]["text"] += body


print ()
print (f'Recnik Turcizama u Srpskokm Jeziku: {len(dictionary["serbian"]["dicts"]["turcizmi"])} items.')




# Recnik Urbane Svakodnevnice

doc = Document("recnik_urbane_svakodnevice.docx")

dictionary["serbian"]["dicts"]["urbani"] = {}

start = False
body = ""
topic_name = ""
header = False
new_topic = ""
counter = 0
italic = False
for paragraph in doc.paragraphs:
    par_text: str = cirilica_u_latinicu(paragraph.text.strip())

    if par_text == "@@@START":
        start = True
        continue

    if par_text == "@@@END":
        start = False

    if not start:
        continue

    if len(par_text) < 2:
        continue

    new_topic = ""
    body = ""
    header = False
    unused_bold = ""
    
    par_struc = []
    for run in paragraph.runs:
        run_txt = cirilica_u_latinicu(run.text)
        if run.bold:
            if run.font.italic:
                # print ("Canceled topic, reason: italic font:  ", run_txt)
                body += f"@@@01{run_txt}@@@01"
                continue
            if body.strip() and body.strip()[-1] not in ".!?":
                # print ("Canceled topic, reason: body not ended:  ", run_txt)
                body += f"@@@01{run_txt}@@@01"
                continue

            if body:
                par_struc.append([body, False])
                body = ""
            new_topic += run_txt
        else:
            if new_topic:
                par_struc.append([new_topic, True])
                new_topic = ""
            if run.font.italic:
                body += f"@@@01{run_txt}@@@01"
            else:
                body += run_txt
    if new_topic:
        par_struc.append([new_topic, True])
    if body:
        par_struc.append([body, False])

    topic_list = []
    # if par_struc:
    #     if not par_struc[0][1]:
    #         print (f"Warnning bad entry : {par_text}")

    header = False
    new_topic = ""
    body = ""
    for idx, part in enumerate(par_struc):
        txt = part[0]
        header = part[1]

        if header:
            if idx == 0:
                new_topic = txt
                continue
            else:
                if len(txt) > 2:
                    topic_list.append([new_topic, body])
                    new_topic = txt
                    body = ""
                else:
                    print ("Canceled New Topic: ", txt)
                    body += txt
        else:
            body += txt
    
    if new_topic or body:
        topic_list.append([new_topic, body])

    # if len(topic_list) > 1:
    #     print ()

    for i in topic_list:
        new_topic = i[0]
        body = i[1]

        new_topic = new_topic.strip(" ,")
        
        new_body = ""
        protected = False
        for j in body:
            if j == "(":
                protected = True
            elif j == ")":
                protected = False
            if protected:
                new_body += j
            else:
                if j == ".":
                    new_body += ".\n"
                else:
                    new_body += j
        body = new_body
        
        body = body.strip(" ,") + " "

        pos = 0
        new_body = ""
        old_pos = 0
        while True:
            pos = body.find("*", pos)
            if pos == -1:
                break
            end = body.find(" ", pos)
            if end == -1:
                break

            new_body += body[old_pos:pos+1]
            new_body += "@@@00"
            new_body += body[pos+1:end]
            new_body += "@@@00"

            pos = end
            old_pos = pos
        new_body += body[old_pos:]

        body = new_body.strip()

        if new_topic:
            topic_name = new_topic
            new_topic = ""
            if topic_name in dictionary["serbian"]["dicts"]["urbani"]:
                dictionary["serbian"]["dicts"]["urbani"][topic_name]["text"] += "\n\n\n"
                print (f"Duplicate Entry : {topic_name}")
            else:
                # print (topic_name)
                dictionary["serbian"]["dicts"]["urbani"][topic_name] = {}
                dictionary["serbian"]["dicts"]["urbani"][topic_name]["text"] = ""
                dictionary["serbian"]["dicts"]["urbani"][topic_name]["links"] = []
                new_topic = ""
                # continue

        if dictionary["serbian"]["dicts"]["urbani"][topic_name]["text"]:
            dictionary["serbian"]["dicts"]["urbani"][topic_name]["text"] += f" {body}"
        else:
            dictionary["serbian"]["dicts"]["urbani"][topic_name]["text"] += body


print ()
print (f'Recnik Urbane Svakodnevnice: {len(dictionary["serbian"]["dicts"]["urbani"])} items.')





# Skolski Geografski Leksikon


doc = Document("geografski_leksikon.docx")

dictionary["serbian"]["dicts"]["geografija"] = {}

start = False
body = ""
topic_name = ""
header = False
new_topic = ""
counter = 0
protected_new_topic = ""
for paragraph in doc.paragraphs:
    par_text: str = paragraph.text.strip()

    if par_text == "@@@START":
        start = True
        continue

    if par_text == "@@@END":
        start = False

    if not start:
        continue

    if len(par_text) < 2:
        continue

    new_topic = ""
    body = ""
    if protected_new_topic:
        new_topic = protected_new_topic + " "
        par_text = protected_new_topic + " " + par_text
        
    if paragraph.style.font.bold:
        if par_text.lower().find("v.") == -1:
            # print ("Found STRONG Topic name: ", par_text)
            protected_new_topic = par_text
            continue
        else:
            new_topic = par_text
    else:
        protected_new_topic = ""
        for run in paragraph.runs:
            run_txt = run.text
            if run.bold:
                if par_text.startswith(new_topic + run_txt):
                    new_topic += run_txt
                else:
                    body += run_txt
            else:
                body += run_txt

    if new_topic.lower().find("v.") != -1:
        # print (f"Found topic without body: {par_text}")
        pos = new_topic.lower().find("v.")
        body = new_topic[pos+2:]
        new_topic = new_topic[:pos]

    new_topic = new_topic.strip(" -,")
    body = body.strip(" -,")

    new_topic = new_topic.replace("\t", " ")
    body = body.replace("\t", " ")
    for i in range(15):
        body = body.replace(f" {i}. ", f"\n{i}. ")

    if body.find("; ") != -1:
        body = body.replace("; ", ";\n")
    else:
        body = body.replace(";", ";\n")
    
    while True:
        new_topic = new_topic.replace("  ", " ")
        body = body.replace("  ", " ")
        if new_topic.find("  ") == -1 and body.find("  ") == -1:
            break

    if new_topic:
        topic_name = new_topic
        new_topic = ""
        if topic_name in dictionary["serbian"]["dicts"]["geografija"]:
            dictionary["serbian"]["dicts"]["geografija"][topic_name]["text"] += "\n\n\n"
            print (f"Duplicate Entry : {topic_name}")
        else:
            # print (topic_name)
            dictionary["serbian"]["dicts"]["geografija"][topic_name] = {}
            dictionary["serbian"]["dicts"]["geografija"][topic_name]["text"] = ""
            dictionary["serbian"]["dicts"]["geografija"][topic_name]["links"] = []
            new_topic = ""
            # continue

    if dictionary["serbian"]["dicts"]["geografija"][topic_name]["text"]:
        dictionary["serbian"]["dicts"]["geografija"][topic_name]["text"] += f" {body}"
    else:
        dictionary["serbian"]["dicts"]["geografija"][topic_name]["text"] += body


print ()
print (f'Skolski Geografski Leksikon: {len(dictionary["serbian"]["dicts"]["geografija"])} items.')




# Skolski leksikon biologije


doc = Document("skolski_leksikon_biologije_s_pitanjima_za_maturu_i_prijemne.docx")

dictionary["serbian"]["dicts"]["biologija"] = {}

start = False
body = ""
topic_name = ""
header = False
new_topic = ""
counter = 0
protected_topic = ""
for paragraph in doc.paragraphs:
    par_text: str = paragraph.text.strip()

    if par_text == "@@@START":
        start = True
        continue

    if par_text == "@@@END":
        start = False

    if not start:
        continue

    if len(par_text) < 2:
        continue

    new_topic = ""
    body = ""
    if protected_topic:
        new_topic = protected_topic + " "
        par_text = protected_topic + " " + par_text
        protected_topic = False

    for run in paragraph.runs:
        run_txt = run.text
        if run.bold:
            if par_text.startswith(new_topic + run_txt):
                new_topic += run_txt
            else:
                if body and not body.strip(" ,"):
                    if par_text.startswith(new_topic + body + run_txt):
                        new_topic += body
                        new_topic += run_txt
                        body = ""
                    else:
                        # print (F"Warning, Found topic inside body:== {run_txt}       >>> {par_text}")
                        # print ()
                        body += f"@@@00{run_txt}@@@00"
                else:    
                    # print (F"Warning, Found topic inside body:== {run_txt}       >>> {par_text}")
                    # print ()
                    body += f"@@@00{run_txt}@@@00"
        else:
            body += run_txt

    if new_topic == par_text and not body:
        protected_topic = new_topic
        continue

    new_topic = new_topic.strip(" ,")
    body = body.lstrip(" ,.")

    new_topic = new_topic.replace("-   ", "")
    new_topic = new_topic.replace("-  ", "")
    new_topic = new_topic.replace("- ", "")
    new_topic = new_topic.replace("\t", " ")
    body = body.replace("-   ", "")
    body = body.replace("-  ", "")
    body = body.replace("- ", "")
    body = body.replace("\t", " ")
    while True:
        new_topic = new_topic.replace("  ", " ")
        body = body.replace("  ", " ")
        if new_topic.find("  ") == -1 and body.find("  ") == -1:
            break

    if new_topic:
        topic_name = new_topic
        new_topic = ""
        if topic_name in dictionary["serbian"]["dicts"]["biologija"]:
            dictionary["serbian"]["dicts"]["biologija"][topic_name]["text"] += "\n\n\n"
            print (f"Duplicate Entry : {topic_name}")
        else:
            # print (topic_name)
            dictionary["serbian"]["dicts"]["biologija"][topic_name] = {}
            dictionary["serbian"]["dicts"]["biologija"][topic_name]["text"] = ""
            dictionary["serbian"]["dicts"]["biologija"][topic_name]["links"] = []
            new_topic = ""
            # continue

    if dictionary["serbian"]["dicts"]["biologija"][topic_name]["text"]:
        dictionary["serbian"]["dicts"]["biologija"][topic_name]["text"] += f" {body}"
    else:
        dictionary["serbian"]["dicts"]["biologija"][topic_name]["text"] += body


print ()
print (f'Skolski leksikon biologije: {len(dictionary["serbian"]["dicts"]["biologija"])} items.')




# Slovenska Mitoloska Enciklopedijski recink

def slo_mit_encikl_get_topic(txt: str, silent: bool = False) -> str:
    pos = txt.find(" -")
    if pos == -1:
        pos = txt.find(" ->")
        if pos == -1:
            return ""
    
    txt = txt[:pos].strip()

    if not txt:
        return ""

    txt_split = [x for x in txt.split(" ") if x != ""]

    if txt_split[0] != txt_split[0].upper():
        return ""
    
    valid_letters = "QWERTYUOPLKJHGFDSAZXCBNMŠĐČĆŽ"
    has_letters = False
    for i in txt_split[0]:
        if i in valid_letters:
            has_letters = True
            break
    
    if not has_letters:
        if not silent:
            print ("Canceled, reason FIRST WORD : ", txt)
        return ""

    count = 0
    switch_to_low = False
    for i in txt_split:
        if i != i.upper() and i.find("(") == -1:
            count += 1
            switch_to_low = True
        else:
            if switch_to_low == True:
                return ""
    
    if len(txt_split[0]) < 5 and  count > 1:
        return ""
    
    if len(txt_split[0]) > 4 and  count > 3:
        return

    valid_letters = "QWERTYUOPLKJHGFDSAZXCBNMŠĐČĆŽIV)(,.-;[]{} "
    ignore_this = False
    for l in txt:
        if l in "([{":
            ignore_this = True
        elif l in ")]}":
            ignore_this = False

        if ignore_this:
            continue

        if l.upper() not in valid_letters:
            if not silent:
                print ("Canceled title name : ", txt)
            return ""

    return txt

def slo_mit_encikl_is_page_info(txt: str, silent: bool = False) -> str:
    if txt.count("\t") < 2:
        return False
    
    txt_split = txt.split("\t")

    if not silent:
        print ("May be page info .... ", end="")

    result  = True
    for i in txt_split:
        if i != i.upper():
            result = False
            break
    
    if not silent:
        print (f"{result}  :  {txt}")
    return result





doc = Document("slovenska_mitologija_enciklopedijski_recnik.docx")

dictionary["serbian"]["dicts"]["slo_mit_encikl"] = {}

start = False
body = ""
topic_name = ""
header = False
new_topic = ""
counter = 0

italic_txt = "@@@00"
literatura_txt = "\n@@@01"
autor_txt = "\n@@@02"

italic = False
literatura = False
autor = False

topic_delim = " -"
doc_map = []
for paragraph in doc.paragraphs:
    par_text: str = cirilica_u_latinicu(paragraph.text.strip())

    if par_text == "@@@START":
        start = True
        continue

    if par_text == "@@@END":
        start = False

    if not start:
        continue

    if len(par_text) < 2:
        continue

    par_text = par_text.replace("Nj", "NJ").replace("Lj", "LJ").replace("Dž", "DŽ")


    # if par_text.find("AŽDAJA") != -1:
    #     print ()


    if slo_mit_encikl_is_page_info(par_text, silent=True):
        continue

    # # Check is it page info
    # page_info = True
    # par_text_split = par_text.split("-")
    # if par_text.count("-") == 2:
    #     if par_text_split[0] != par_text_split[0].upper():
    #         page_info = False
    #     if par_text_split[2] != par_text_split[2].upper():
    #         page_info = False
    #     if page_info:
    #         try:
    #             _ = int(par_text_split[1])
    #         except:
    #             print (f"Is this Page Info ????????????   : {par_text}")
    #             page_info = False
    # else:
    #     if par_text.count("-") > 2:
    #         for i in par_text_split:
    #             if i != i.upper():
    #                 page_info = False
    #     else:
    #         page_info = False

    # if page_info:
    #     print (f"Found Page Info: {par_text}")
    #     continue

    new_topic = ""
    body = ""

    new_topic = slo_mit_encikl_get_topic(par_text, silent=True)

    if new_topic:
        if literatura:
            doc_map.append([literatura_txt, "b"])
            literatura = False
        if autor:
            doc_map.append([autor_txt, "b"])
            autor = False
        if italic:
            doc_map.append([italic_txt, "b"])
            italic = False
        # print (new_topic)
        doc_map.append([new_topic, "t"])

    # # Find New Topic
    # pos = par_text.find(topic_delim)
    # if pos != -1:
    #     pot_topic = par_text[:pos].strip()

    #     if pot_topic == pot_topic.upper() and (len(pot_topic) > 2 or par_text[pos:pos+3] == " ->"):
    #         if literatura:
    #             doc_map.append([literatura_txt, "b"])
    #             literatura = False
    #         if autor:
    #             doc_map.append([autor_txt, "b"])
    #             autor = False
    #         if italic:
    #             doc_map.append([italic_txt, "b"])
    #             italic = False
    #         new_topic = pot_topic
    #         print (new_topic)
    #         doc_map.append([new_topic, "t"])

    # Check is it literatura
    if par_text.startswith("Lit.:") and not new_topic:
        par_text = "\n\n" + literatura_txt + par_text
        doc_map.append([par_text, "b"])
        literatura = True
        continue
    
    if literatura:
        if new_topic:
            literatura = False
        else:
            doc_map.append([par_text, "b"])
            continue

    # Find Body
    run_text_total = ""
    if new_topic:
        found_new_topic = False
    else:
        found_new_topic = True
    run_text_total = ""

    for run in paragraph.runs:
        # if run.font.size == 95250:
        #     print (f"Image Found: {run.text}")
        #     continue

        run_text_total += cirilica_u_latinicu(run.text)

        if not found_new_topic:
            if run_text_total.find(topic_delim) == -1:
                continue
            run_txt = run_text_total[run_text_total.find(topic_delim)+len(topic_delim):]
            found_new_topic = True
        else:
            run_txt = cirilica_u_latinicu(run.text)
        
        if run.font.italic or paragraph.style.font.italic:
            run_txt = f"{italic_txt}{run_txt}{italic_txt}"
        
        doc_map.append([run_txt, "b"])

    doc_map.append([" ", "b"])


start = False
body = ""
topic_name = ""
header = False
new_topic = ""
counter = 0
for part in doc_map:
    if part[1] == "t":
        new_topic = part[0].strip(" >,-")
        body = ""
    elif part[1] == "b":
        new_topic = ""
        body = part[0]
    else:
        print ("Error, Document Map corrupted.", part[0])
        continue

    if new_topic:
        topic_name = new_topic
        new_topic = ""
        if topic_name in dictionary["serbian"]["dicts"]["slo_mit_encikl"]:
            topic_name = topic_name + " (1)"
            print (f"Duplicate Entry : {topic_name}")
        # print (topic_name)
        dictionary["serbian"]["dicts"]["slo_mit_encikl"][topic_name] = {}
        dictionary["serbian"]["dicts"]["slo_mit_encikl"][topic_name]["text"] = ""
        dictionary["serbian"]["dicts"]["slo_mit_encikl"][topic_name]["links"] = []
        new_topic = ""
        # continue

    if dictionary["serbian"]["dicts"]["slo_mit_encikl"][topic_name]["text"]:
        dictionary["serbian"]["dicts"]["slo_mit_encikl"][topic_name]["text"] += body
    else:
        dictionary["serbian"]["dicts"]["slo_mit_encikl"][topic_name]["text"] += body

for i in dictionary["serbian"]["dicts"]["slo_mit_encikl"]:
    txt: str = " " + dictionary["serbian"]["dicts"]["slo_mit_encikl"][i]["text"] + " "

    txt_t = i
    for j in "~!@#$%^&*()+=_-{}[];'|:\"\\,./?><":
        txt_t = txt_t.replace(j, " ")
    
    txt_split = [x for x in txt_t.split(" ") if x != ""]
    repl = []
    forbiden_items = [f" {x} " for x in "auisoAUISO"]
    
    for j in txt_split:
        if j == j.upper():
            for j_pref in [" ", "0", "\n"]:
                for j_suff in [". ", " ", ".,", ".@", "@"]:
                    repl_item = f"{j_pref}{j[0]}{j_suff}"
                    repl_with = f"{j_pref}{j.lower().capitalize()}{j_suff.replace('.', '')}"
                    if repl_item not in forbiden_items:
                        repl.append([repl_item, repl_with])

                    repl_item = f"{j_pref}{j[0].lower()}{j_suff}"
                    repl_with = f"{j_pref}{j.lower()}{j_suff.replace('.', '')}"
                    if repl_item not in forbiden_items:
                        repl.append([repl_item, repl_with])
        else:
            break

    txt = txt.replace("@@@00@@@00", "")
    replaced = []
    for j in repl:
        if j[0] not in replaced:
            txt = txt.replace(j[0], j[1])
        replaced.append(j[0])

    txt = txt.strip(" >,-")
    txt = txt.replace("@@@00@@@00", "")
    txt = txt.replace("@@@00\n@@@00", "\n")
    txt = txt.replace("-    ", "")
    txt = txt.replace("-   ", "")
    txt = txt.replace("-  ", "")
    txt = txt.replace("- ", "")
    txt = txt.replace("-\n", "")
    txt = txt.replace("\n-", "")
    txt = txt.replace("->", "")

    count = 0
    new_txt = ""
    old_j = ""
    protected = False
    for j in txt:
        if j in "([":
            protected = True
        if j in ")]":
            protected = False
        
        if j != "@":
            count += 1
        
        new_txt += j
        if not protected and j in ".!?;" and count > 50 and old_j in "qwertyuioplkjhgfdsazxcvbnmšđčćž":
            new_txt += "\n"
            count = 0
        old_j = j
    
    dictionary["serbian"]["dicts"]["slo_mit_encikl"][i]["text"] = new_txt

print ()
print (f'Slovenska Mitoloska Enciklopedijski recink: {len(dictionary["serbian"]["dicts"]["slo_mit_encikl"])} items.')






# Tehnicki Recnik

with open("tehnicki_recnik.txt", "r", encoding="utf-8") as file:
    txt = file.read()

txt_list = txt.split("\n")

dictionary["serbian"]["dicts"]["tehnicki"] = {}

start = False
body = ""
topic_name = ""
header = False
new_topic = ""
counter = 0
for line in txt_list:
    line = line.strip()

    if len(line) < 3:
        continue

    if line.find("@") != -1:
        pos = line.find("@")
    elif line.find("-") != -1:
        if line.count("-") != 1:
            print (f"Incorect number of delimiters : {line}")
            continue
        pos = line.find("-")
    else:
        print (f"Bad entry:  {line}")
        continue
    
    new_topic = line[:pos].strip()
    body = line[pos+1:].strip()

    if new_topic:
        topic_name = new_topic
        new_topic = ""
        if topic_name in dictionary["serbian"]["dicts"]["tehnicki"]:
            dictionary["serbian"]["dicts"]["tehnicki"][topic_name]["text"] += "\n\n\n"
            print (f"Duplicate Entry : {topic_name}")
        else:
            # print (topic_name)
            dictionary["serbian"]["dicts"]["tehnicki"][topic_name] = {}
            dictionary["serbian"]["dicts"]["tehnicki"][topic_name]["text"] = ""
            dictionary["serbian"]["dicts"]["tehnicki"][topic_name]["links"] = []
            new_topic = ""
            # continue

    if dictionary["serbian"]["dicts"]["tehnicki"][topic_name]["text"]:
        dictionary["serbian"]["dicts"]["tehnicki"][topic_name]["text"] += f" {body}"
    else:
        dictionary["serbian"]["dicts"]["tehnicki"][topic_name]["text"] += body


print ()
print (f'Tehnicki Recnik: {len(dictionary["serbian"]["dicts"]["tehnicki"])} items.')




# Tolkinov recnik


doc = Document("Tolkinov_recnik.docx")

dictionary["serbian"]["dicts"]["tolkin"] = {}

start = False
body = ""
topic_name = ""
header = False
new_topic = ""
counter = 0
topics_pool = ["start   "]
body_mode = False
for paragraph in doc.paragraphs:
    par_text: str = paragraph.text.strip()

    if par_text == "Angmar":
        start = True

    if not start:
        continue

    # if not par_text:
    #     continue

    new_topic = ""
    body = ""

    if body_mode and paragraph.style.name.startswith("Body") and not par_text:
        body = ""
    else:
        if not par_text:
            continue

    if paragraph.style.name.startswith("Normal"):
        if par_text:
            topics_pool.append(par_text)
        continue
    elif paragraph.style.name.startswith("Body"):
        body += par_text
    else:
        print (f"Unknown paragraph style:  {par_text}")

    if body:
        if body[0] == "1":
            body = body[1:]
            topics_pool.pop(0)
            new_topic = topics_pool[0]
    else:
        body = "\n\n"

    if body.endswith(".") or body.endswith("!") or body.endswith("?"):
        body_mode = True
    else:
        body_mode = False

    new_topic = new_topic.strip()
    

    if new_topic:
        topic_name = new_topic
        new_topic = ""
        if topic_name in dictionary["serbian"]["dicts"]["tolkin"]:
            dictionary["serbian"]["dicts"]["tolkin"][topic_name]["text"] += "\n\n\n"
            print (f"Duplicate Entry : {topic_name}")
        else:
            # print (topic_name)
            dictionary["serbian"]["dicts"]["tolkin"][topic_name] = {}
            dictionary["serbian"]["dicts"]["tolkin"][topic_name]["text"] = ""
            dictionary["serbian"]["dicts"]["tolkin"][topic_name]["links"] = []
            new_topic = ""
            # continue

    if dictionary["serbian"]["dicts"]["tolkin"][topic_name]["text"]:
        if body == "\n\n":
            if not dictionary["serbian"]["dicts"]["tolkin"][topic_name]["text"].endswith("\n"):
                dictionary["serbian"]["dicts"]["tolkin"][topic_name]["text"] += body
        else:
            dictionary["serbian"]["dicts"]["tolkin"][topic_name]["text"] += f" {body}"
    else:
        dictionary["serbian"]["dicts"]["tolkin"][topic_name]["text"] += body


print ()
print (f'Tolkinov recnik: {len(dictionary["serbian"]["dicts"]["tolkin"])} items.')




# Istorijski Leksikon


doc = Document("vladimir_corovic_istorijski_leksikon.docx")

dictionary["serbian"]["dicts"]["istorijski"] = {}

start = False
body = ""
topic_name = ""
header = False
new_topic = ""
counter = 0
for paragraph in doc.paragraphs:
    par_text: str = cirilica_u_latinicu(paragraph.text.strip())

    if par_text == "@@@START":
        start = True
        continue

    if par_text == "@@@END":
        start = False

    if not start:
        continue

    if len(par_text) < 3:
        continue

    new_topic = ""
    body = ""

    if paragraph.style.name.startswith("Heading 1"):
        print (f"Found letter: {par_text}")
        continue
    elif paragraph.style.name.startswith("Heading 3"):
        # print ("New Topic: ", par_text)
        new_topic = par_text
    elif paragraph.style.name.startswith("Body"):
        body = par_text
        if body[0] == "\t":
            body = "\n" + body
        
        if body.lower().startswith("l i t e r a t u r a"):
            body = f"\n\n@@@00{body}@@@00"
    else:
        print ("Unrecognized paragraph style:  ", par_text)
        continue

    new_topic = new_topic.replace("\t", " ").strip()
    body = body.replace("\t", " ")

    if new_topic:
        topic_name = new_topic
        new_topic = ""
        if topic_name in dictionary["serbian"]["dicts"]["istorijski"]:
            dictionary["serbian"]["dicts"]["istorijski"][topic_name]["text"] += "\n\n\n"
            print (f"Duplicate Entry : {topic_name}")
        else:
            # print (topic_name)
            dictionary["serbian"]["dicts"]["istorijski"][topic_name] = {}
            dictionary["serbian"]["dicts"]["istorijski"][topic_name]["text"] = ""
            dictionary["serbian"]["dicts"]["istorijski"][topic_name]["links"] = []
            new_topic = ""
            # continue

    if dictionary["serbian"]["dicts"]["istorijski"][topic_name]["text"]:
        dictionary["serbian"]["dicts"]["istorijski"][topic_name]["text"] += f" {body}"
    else:
        dictionary["serbian"]["dicts"]["istorijski"][topic_name]["text"] += body

for i in dictionary["serbian"]["dicts"]["istorijski"]:
    txt: str = dictionary["serbian"]["dicts"]["istorijski"][i]["text"]
    txt = txt.replace("-    ", "")
    txt = txt.replace("-   ", "")
    txt = txt.replace("-  ", "")
    txt = txt.replace("- ", "")

    txt = " " + txt + " "
    txt_literatura = ""
    pos = txt.lower().find("l i t e r a t u r a")
    if pos >= 0:
        txt_literatura = txt[pos:]
        txt = txt[:pos]

    txt_t = i
    for j in "~!@#$%^&*()+=_-{}[];'|:\"\\,./?><":
        txt_t = txt_t.replace(j, " ")
    txt_t = txt_t.replace("Nj", "NJ")
    txt_t = txt_t.replace("Lj", "LJ")
    txt_t = txt_t.replace("Dž", "DŽ")
    
    txt_split = [x for x in txt_t.split(" ") if x != ""]
    repl = []
    forbiden_items = [f" {x} " for x in "auisoAUISO"]
    
    for j in txt_split:
        if j == j.upper():
            for j_pref in [" ", "0", "\n"]:
                for j_suff in [". ", " ", ".,", ".@", "@"]:
                    repl_item = f"{j_pref}{j[0]}{j_suff}"
                    repl_with = f"{j_pref}{j.lower().capitalize()}{j_suff.replace('.', '')}"
                    if repl_item not in forbiden_items:
                        repl.append([repl_item, repl_with])

                    repl_item = f"{j_pref}{j[0].lower()}{j_suff}"
                    repl_with = f"{j_pref}{j.lower()}{j_suff.replace('.', '')}"
                    if repl_item not in forbiden_items:
                        repl.append([repl_item, repl_with])
        else:
            break

    txt = txt.replace("@@@00@@@00", "")
    replaced = []
    for j in repl:
        if j[0] not in replaced:
            txt = txt.replace(j[0], j[1])
        replaced.append(j[0])

    txt = txt + txt_literatura

    dictionary["serbian"]["dicts"]["istorijski"][i]["text"] = txt

print ()
print (f'Istorijski Leksikon: {len(dictionary["serbian"]["dicts"]["istorijski"])} items.')



# Vlaski Recnik


doc = Document("vlaski_recnik_cirilica.docx")

dictionary["serbian"]["dicts"]["vlaski"] = {}

start = False
body = ""
topic_name = ""
header = False
new_topic = ""
counter = 0
for paragraph in doc.paragraphs:
    par_text: str = cirilica_u_latinicu(paragraph.text.strip())

    if par_text == "@@@START":
        start = True
        continue

    if par_text == "@@@END":
        start = False

    if not start:
        continue

    if len(par_text) < 3:
        continue

    new_topic = ""
    body = ""
    one_letter_body = ""
    watch_for_one_letter_body = False
    topics_list = []
    for run in paragraph.runs:
        run_txt = cirilica_u_latinicu(run.text)
        if run.bold:
            if body and new_topic:
                topics_list.append([new_topic, body])
                new_topic = ""
                body = ""

            if one_letter_body:
                new_topic += one_letter_body
                one_letter_body = ""
            new_topic += run_txt
            watch_for_one_letter_body = True

        else:
            if one_letter_body and len(run_txt) <= 2:
                one_letter_body += run_txt
                continue

            if one_letter_body:
                body += one_letter_body
                one_letter_body = ""

            if len(run_txt) <= 2 and watch_for_one_letter_body:
                one_letter_body = run_txt
                watch_for_one_letter_body = False
                continue

            watch_for_one_letter_body = False
            body += run_txt
    if one_letter_body:
        body += one_letter_body
        one_letter_body = ""

    if body and new_topic:
        topics_list.append([new_topic, body])
        new_topic = ""
        body = ""

    for i in topics_list:
        new_topic = i[0].replace("\t", " ")
        body = i[1].replace("\t", " ")

        new_topic = "Vlaški: " + new_topic.strip(" -").lower()
        body = "Srpski:\n" + body.strip(" -").lower()

        if new_topic:
            topic_name = new_topic
            new_topic = ""
            if topic_name in dictionary["serbian"]["dicts"]["vlaski"]:
                dictionary["serbian"]["dicts"]["vlaski"][topic_name]["text"] += "\n\n\n"
                print (f"Duplicate Entry : {topic_name}")
            else:
                # print (topic_name)
                dictionary["serbian"]["dicts"]["vlaski"][topic_name] = {}
                dictionary["serbian"]["dicts"]["vlaski"][topic_name]["text"] = ""
                dictionary["serbian"]["dicts"]["vlaski"][topic_name]["links"] = []
                new_topic = ""
                # continue

        if dictionary["serbian"]["dicts"]["vlaski"][topic_name]["text"]:
            dictionary["serbian"]["dicts"]["vlaski"][topic_name]["text"] += f" {body}"
        else:
            dictionary["serbian"]["dicts"]["vlaski"][topic_name]["text"] += body

print (f'Vlasko -> Srpski: {len(dictionary["serbian"]["dicts"]["vlaski"])} items.')




def vlaski_get_topics(txt: str) -> list:
    txt = txt.replace("\t", " ")
    txt_original = txt
    topics = []
    if txt.count("-") == 1 and txt.count("\t") == 1:
        txt_split = txt.split("\t")
        topics.append([txt_split[0], txt_split[1]])
        return topics
    
    if txt.count("\t") > 1:
        print (f"Error. Too many tab delimiters: {txt}")
        txt = txt.replace("\t", " ")
    
    if txt.count("-") == 1:
        txt_split = txt.split("-")
        topics.append([txt_split[0], txt_split[1]])
        return topics

    if txt.count("-") == 0 and txt.count("\t") > 0:
        txt_t = txt.replace('\t', ' TAB ')
        print (f"Warnning,  Is This New Topic Item ?  :  {txt_t}")
        txt_split = txt.split("\t")
        topics.append([txt_split[0], txt_split[1]])
        return topics
        
    while True:
        txt = txt.replace("  ", " ")
        if txt.find("  ") == -1:
            break

    txt_words = []
    delim = " "
    pos = 0
    while True:
        pos = txt.find(delim)
        if pos == -1:
            if txt:
                if txt.strip():
                    txt_words.insert(0, txt.strip())
            break

        word = txt[:pos+1]
        if word.strip():
            txt_words.insert(0, word.strip())
        txt = txt[pos+1:]
        if txt:
            if txt[0] == "(":
                delim = ")"
            else:
                delim = " "
        
    new_topic = ""
    body = ""
    topic_name_mode = False
    for idx, word in enumerate(txt_words):
        if word == "-":
            if new_topic:
                if new_topic.find(";") != -1:
                    pos = new_topic.rfind(";")
                    body_t = new_topic[:pos+1]
                    new_topic = new_topic[pos+1:]
                    topics.append([new_topic, body])
                    body = body_t
                    new_topic = ""
                else:
                    print (f"Error, cannot resolve string : {txt_original}")
                    return None
            topic_name_mode = True
            continue

        if topic_name_mode:
            new_topic = f"{word} {new_topic}"
            if idx == len(txt_words) - 1:
                topics.append([new_topic, body])
                break

            topic_name_mode = False

            if word[0] in "(," or txt_words[idx+1][-1] in ",;":
                topic_name_mode = True
                continue

            topics.append([new_topic, body])
            new_topic = ""
            body = ""
        else:
            body = f"{word} {body}"
        
    return topics

def vlaski_get_topic_names(txt: str) -> list:
    txt = txt.replace("\t", " ")
    while True:
        txt = txt.replace("  ", " ")
        if txt.find("  ") == -1:
            break

    txt_words = []
    delim = " "
    pos = 0
    while True:
        pos = txt.find(delim)
        if pos == -1:
            break

        word = txt[:pos+1]
        if word.strip():
            txt_words.append(word.strip())
        txt = txt[pos+1:]
        if txt:
            if txt[0] == "(":
                delim = ")"
            else:
                delim = " "

    if txt.strip():
        txt_words.append(txt.strip())

    topics = []
    new_topic = ""
    for idx, word in enumerate(txt_words):
        if word[0] in ",(" or (new_topic and new_topic[-1] in ","):
            new_topic = f"{new_topic} {word}"
        else:
            if new_topic:
                topics.append(new_topic)
            new_topic = word
    
    if new_topic:
        topics.append(new_topic)

    return topics

start = False
body = ""
topic_name = ""
header = False
new_topic = ""
counter = 0
topics_buffer =  []
expected_list_paragraph = False
for paragraph in doc.paragraphs:
    par_text: str = paragraph.text.strip()

    if par_text == "@@@LATIN":
        start = True
        continue

    if par_text == "@@@END":
        start = False

    if not start:
        continue

    if not par_text:
        continue

    if par_text in ["еtu", "еćilji"]:
        continue

    new_par_text = ""
    protected = False
    for i in par_text:
        if i == "(":
            protected = True
        if i == ")":
            protected = False
        if protected:
            if i in "-\t":
                i = " "
        new_par_text += i
    par_text = new_par_text


    new_topic = ""
    body = ""
    topics_list = []

    # if par_text.startswith("аrhеоlоgiја"):
    #     print ()

    if paragraph.style.name.startswith("List") or par_text.startswith("-"):
        expected_list_paragraph = False
        if topics_buffer:
            new_topic = topics_buffer[0]
        else:
            # print ("Error. Tpoic Buffer exhausted.... : ", par_text)
            continue
        body = par_text
        topics_buffer.pop(0)
        topics_list.append([new_topic, body])
    elif paragraph.style.name.startswith("Body") or paragraph.style.name.startswith("Normal"):
        if par_text.find("-") != -1:            
            # if expected_list_paragraph:
            #     print ("Warning, List Paragaph expected !  : ", par_text)
            if topics_buffer:
                # print (f"Warning, Topic Buffer not exhausted, but is should be. Buffer items: {len(topics_buffer)}     {par_text}")
                topics_buffer = []
            topics = vlaski_get_topics(par_text)
            for i in topics:
                topics_list.append(i)
        else:
            if paragraph.style.name.startswith("Normal"):
                continue
            topics = vlaski_get_topic_names(par_text)
            for i in topics:
                topics_buffer.append(i)
            expected_list_paragraph = True
            continue
    elif paragraph.style.name.startswith("Heading 2"):
        # print ("Found New Letter start : ", par_text)
        continue
    else:
        print (f"Warning, Unknown paragraph header !  : {par_text}")
        continue


    for i in topics_list:
        new_topic = i[0].replace("\t", " ").replace("_", " ")
        body = i[1].replace("\t", " ").replace("_", " ")

        new_topic = "Srpski: " + new_topic.strip(" -").lower()
        body = "Vlaški:\n" + body.strip(" -").lower()

        if new_topic:
            topic_name = new_topic
            new_topic = ""
            if topic_name in dictionary["serbian"]["dicts"]["vlaski"]:
                dictionary["serbian"]["dicts"]["vlaski"][topic_name]["text"] += "\n\n\n"
                print (f"Duplicate Entry : {topic_name}")
            else:
                # print (f"{topic_name}       {body}")
                dictionary["serbian"]["dicts"]["vlaski"][topic_name] = {}
                dictionary["serbian"]["dicts"]["vlaski"][topic_name]["text"] = ""
                dictionary["serbian"]["dicts"]["vlaski"][topic_name]["links"] = []
                new_topic = ""
                # continue

        if dictionary["serbian"]["dicts"]["vlaski"][topic_name]["text"]:
            dictionary["serbian"]["dicts"]["vlaski"][topic_name]["text"] += f" {body}"
        else:
            dictionary["serbian"]["dicts"]["vlaski"][topic_name]["text"] += body


print ()
print (f'Vlaski Recnik PART I: {len(dictionary["serbian"]["dicts"]["vlaski"])} items.')


doc = Document("vlaski_recnik_latinica.docx")

start = False
body = ""
topic_name = ""
header = False
new_topic = ""
counter = 0
for paragraph in doc.paragraphs:
    par_text: str = paragraph.text.strip()

    if par_text == "@@@START":
        start = True
        continue

    if par_text == "@@@END":
        start = False

    if not start:
        continue

    if len(par_text) < 2:
        continue

    new_topic = ""
    body = ""
    one_letter_body = ""
    watch_for_one_letter_body = False
    topics_list = []
    for run in paragraph.runs:
        run_txt = cirilica_u_latinicu(run.text)
        if run.bold:
            if body:
                topics_list.append([new_topic, body])
                new_topic = ""
                body = ""

            if one_letter_body:
                new_topic += one_letter_body
                one_letter_body = ""
            new_topic += run_txt
            watch_for_one_letter_body = True

        else:
            if one_letter_body and len(run_txt) <= 2:
                one_letter_body += run_txt
                continue

            if one_letter_body:
                body += one_letter_body
                one_letter_body = ""

            if len(run_txt) <= 2 and watch_for_one_letter_body:
                one_letter_body = run_txt
                watch_for_one_letter_body = False
                continue

            watch_for_one_letter_body = False
            body += run_txt
    if one_letter_body:
        body += one_letter_body
        one_letter_body = ""

    if body:
        topics_list.append([new_topic, body])
        new_topic = ""
        body = ""

    for i in topics_list:
        new_topic = i[0].replace("\t", " ")
        body = i[1].replace("\t", " ")

        new_topic = "Srpski: " + new_topic.strip(" -")
        body = "Vlaški:\n" + body.strip(" -")
        while True:
            new_topic = new_topic.replace("  ", " ")
            body = body.replace("  ", " ")
            if new_topic.find("  ") == -1 and body.find("  ") == -1:
                break

        if new_topic:
            topic_name = new_topic
            new_topic = ""
            if topic_name in dictionary["serbian"]["dicts"]["vlaski"]:
                dictionary["serbian"]["dicts"]["vlaski"][topic_name]["text"] += "\n\n\n"
                print (f"Duplicate Entry : {topic_name}")
            else:
                # print (topic_name)
                dictionary["serbian"]["dicts"]["vlaski"][topic_name] = {}
                dictionary["serbian"]["dicts"]["vlaski"][topic_name]["text"] = ""
                dictionary["serbian"]["dicts"]["vlaski"][topic_name]["links"] = []
                new_topic = ""
                # continue

        if dictionary["serbian"]["dicts"]["vlaski"][topic_name]["text"]:
            dictionary["serbian"]["dicts"]["vlaski"][topic_name]["text"] += f" {body}"
        else:
            dictionary["serbian"]["dicts"]["vlaski"][topic_name]["text"] += body



skracenice_list = vlaski_skracenice.split("\n")

for i in skracenice_list:
    pos = i.find("=")
    if pos < 0:
        print ("Error, no delimiter. ", i)
        continue
    new_topic = i[:pos].strip()
    body = i[pos+1:].strip()

    if new_topic:
        topic_name = new_topic
        new_topic = ""
        if topic_name in dictionary["serbian"]["dicts"]["vlaski"]:
            dictionary["serbian"]["dicts"]["vlaski"][topic_name]["text"] += "\n\n\n"
            print (f"Duplicate Entry : {topic_name}")
        else:
            print (topic_name)
            dictionary["serbian"]["dicts"]["vlaski"][topic_name] = {}
            dictionary["serbian"]["dicts"]["vlaski"][topic_name]["text"] = ""
            dictionary["serbian"]["dicts"]["vlaski"][topic_name]["links"] = []
            new_topic = ""
            # continue

    if dictionary["serbian"]["dicts"]["vlaski"][topic_name]["text"]:
        dictionary["serbian"]["dicts"]["vlaski"][topic_name]["text"] += f" {body}"
    else:
        dictionary["serbian"]["dicts"]["vlaski"][topic_name]["text"] += body


print ()
print (f'Vlaski Recnik: {len(dictionary["serbian"]["dicts"]["vlaski"])} items.')




# Krivicni zakonik

def zakonik_clear_head_item(txt: str) -> str:
    for i in range(10):
        txt = txt.replace(f"@@@0{i}", "")
    txt = txt.replace("@@@10", "")
    return txt

doc = Document("krivicni_zakonik.docx")

dictionary["serbian"]["dicts"]["zakon_krivicni_zakonik"] = {}

start = False
body = ""
topic_name = ""
new_topic = ""
counter = 0
header = []
header_mode = False
sub_title = ""
clan = ""
save_topic = False
clan_list_num = 0
for paragraph in doc.paragraphs:
    par_text: str = paragraph.text.strip()

    if par_text == "@@@START":
        start = True
        continue

    if par_text == "@@@END":
        start = False

    if not start:
        continue

    if len(par_text) < 2:
        continue

    new_topic = ""
    body = ""

    if paragraph.style.name.startswith("Heading 1"):
        clan_list_num = 0
        header = []
        header.append(par_text)
        header_mode = True
        continue
    
    if header_mode:
        if paragraph.alignment == 1 or paragraph.style.name.startswith("Heading"):
            if not par_text.startswith("Član"):
                header.append(par_text)
                continue

            new_topic = par_text + " - " + zakonik_clear_head_item(header[len(header)-1])
            par_text = f"@@@10{par_text}@@@10"
            header.append(par_text)

            header_mode = False
            for idx, i in enumerate(header):
                format_num = idx
                if format_num > 2:
                    format_num = 2
                if idx == len(header) - 1:
                    body += i + "\n"
                else:
                    body += f"@@@0{format_num}{i}@@@0{format_num}\n"
                    header[idx] = f"@@@0{format_num}{i}@@@0{format_num}"

            header.pop(len(header)-1)
            header.pop(len(header)-1)
            save_topic = True
        else:
            print ("Error, unexpected end of title: ", par_text)
            continue
    
    if not save_topic:
        if paragraph.alignment == 1 or paragraph.style.name.startswith("Heading"):
            if par_text.startswith("Član"):
                clan_list_num = 0
                new_topic = par_text + " - " + zakonik_clear_head_item(header[len(header)-1])
                for i in header:
                    body += i + "\n"
                body += f"@@@10{par_text}@@@10\n"
                while len(header) > 2:
                    header.pop(len(header)-1)
            else:
                header.append(f"@@@02{par_text}@@@02")
                continue
        else:
            if paragraph.style.name.startswith("List"):
                clan_list_num += 1
                body = f"\n@@@03({clan_list_num}) @@@03"
            body += par_text
    else:
        save_topic = False

    new_topic = new_topic.strip()

    if new_topic:
        topic_name = new_topic
        new_topic = ""
        if topic_name in dictionary["serbian"]["dicts"]["zakon_krivicni_zakonik"]:
            dictionary["serbian"]["dicts"]["zakon_krivicni_zakonik"][topic_name]["text"] += "\n\n\n"
            print (f"Duplicate Entry : {topic_name}")
        else:
            # print (topic_name)
            dictionary["serbian"]["dicts"]["zakon_krivicni_zakonik"][topic_name] = {}
            dictionary["serbian"]["dicts"]["zakon_krivicni_zakonik"][topic_name]["text"] = ""
            dictionary["serbian"]["dicts"]["zakon_krivicni_zakonik"][topic_name]["links"] = []
            new_topic = ""
            # continue

    if dictionary["serbian"]["dicts"]["zakon_krivicni_zakonik"][topic_name]["text"]:
        dictionary["serbian"]["dicts"]["zakon_krivicni_zakonik"][topic_name]["text"] += f" {body}"
    else:
        dictionary["serbian"]["dicts"]["zakon_krivicni_zakonik"][topic_name]["text"] += body


for i in dictionary["serbian"]["dicts"]["zakon_krivicni_zakonik"]:
    txt = dictionary["serbian"]["dicts"]["zakon_krivicni_zakonik"][i]["text"]

    txt = txt.replace("@@@10", "\n@@@10")
    
    dictionary["serbian"]["dicts"]["zakon_krivicni_zakonik"][i]["text"] = txt

print ()
print (f'Krivicni zakonik: {len(dictionary["serbian"]["dicts"]["zakon_krivicni_zakonik"])} items.')





# Zakon o krivicnom postupku

def zakonik_postupak_clear_head_item(txt: str) -> str:
    for i in range(10):
        txt = txt.replace(f"@@@0{i}", "")
    txt = txt.replace("@@@10", "")
    return txt

doc = Document("Zakonik_o_krivicnom_postupku.docx")

dictionary["serbian"]["dicts"]["zakon_krivicni_postupak"] = {}

start = False
body = ""
topic_name = ""
new_topic = ""
counter = 0
header = []
header_mode = False
sub_title = ""
clan = ""
save_topic = False
clan_list_num = 0
zakon_part = ""
for paragraph in doc.paragraphs:
    par_text: str = paragraph.text.strip()

    if par_text == "@@@START":
        start = True
        continue

    if par_text == "@@@END":
        start = False

    if not start:
        continue

    if len(par_text) < 2:
        continue

    new_topic = ""
    body = ""

    if paragraph.style.name.startswith("Title"):
        zakon_part = par_text
        continue

    if paragraph.style.name.startswith("Heading 1"):
        clan_list_num = 0
        header = []
        header.append(par_text)
        header_mode = True
        continue
    
    if header_mode:
        if paragraph.alignment == 1 or paragraph.style.name.startswith("Heading"):
            if not par_text.startswith("Član"):
                header.append(par_text)
                continue

            new_topic = par_text + " - " + zakonik_postupak_clear_head_item(header[len(header)-1])
            par_text = f"@@@10{par_text}@@@10"
            header.append(par_text)

            header_mode = False
            if zakon_part:
                body += f"@@@04{zakon_part}@@@04\n"
            for idx, i in enumerate(header):
                format_num = idx
                if format_num > 2:
                    format_num = 2
                if idx == len(header) - 1:
                    body += i + "\n"
                else:
                    body += f"@@@0{format_num}{i}@@@0{format_num}\n"
                    header[idx] = f"@@@0{format_num}{i}@@@0{format_num}"

            while len(header) > 2:
                header.pop(len(header)-1)
            
            save_topic = True
        else:
            print ("Error, unexpected end of title: ", par_text)
            continue
    
    if not save_topic:
        if paragraph.alignment == 1 or paragraph.style.name.startswith("Heading"):
            if par_text.startswith("Član"):
                clan_list_num = 0
                new_topic = par_text + " - " + zakonik_postupak_clear_head_item(header[len(header)-1])

                if zakon_part:
                    body += f"@@@04{zakon_part}@@@04\n"

                for i in header:
                    body += i + "\n"
                body += f"@@@10{par_text}@@@10\n"
                while len(header) > 2:
                    header.pop(len(header)-1)
            else:
                header.append(f"@@@02{par_text}@@@02")
                continue
        else:
            if paragraph.style.name.startswith("List"):
                clan_list_num += 1
                body = f"\n@@@03({clan_list_num}) @@@03"
            body += par_text
    else:
        save_topic = False

    new_topic = new_topic.strip()

    if new_topic:
        topic_name = new_topic
        new_topic = ""
        if topic_name in dictionary["serbian"]["dicts"]["zakon_krivicni_postupak"]:
            dictionary["serbian"]["dicts"]["zakon_krivicni_postupak"][topic_name]["text"] += "\n\n\n"
            print (f"Duplicate Entry : {topic_name}")
        else:
            # print (topic_name)
            dictionary["serbian"]["dicts"]["zakon_krivicni_postupak"][topic_name] = {}
            dictionary["serbian"]["dicts"]["zakon_krivicni_postupak"][topic_name]["text"] = ""
            dictionary["serbian"]["dicts"]["zakon_krivicni_postupak"][topic_name]["links"] = []
            new_topic = ""
            # continue

    if dictionary["serbian"]["dicts"]["zakon_krivicni_postupak"][topic_name]["text"]:
        dictionary["serbian"]["dicts"]["zakon_krivicni_postupak"][topic_name]["text"] += f" {body}"
    else:
        dictionary["serbian"]["dicts"]["zakon_krivicni_postupak"][topic_name]["text"] += body


for i in dictionary["serbian"]["dicts"]["zakon_krivicni_postupak"]:
    txt = dictionary["serbian"]["dicts"]["zakon_krivicni_postupak"][i]["text"]

    txt = txt.replace("@@@10", "\n@@@10")
    
    dictionary["serbian"]["dicts"]["zakon_krivicni_postupak"][i]["text"] = txt

print ()
print (f'Krivicni zakonik: {len(dictionary["serbian"]["dicts"]["zakon_krivicni_postupak"])} items.')




# Zakon o Radu

def zakon_rad_clear_head_item(txt: str) -> str:
    for i in range(10):
        txt = txt.replace(f"@@@0{i}", "")
    txt = txt.replace("@@@10", "")
    txt = txt.replace("\n", "")
    return txt

doc = Document("zakon_o_radu.docx")

dictionary["serbian"]["dicts"]["zakon_o_radu"] = {}

start = False
body = ""
topic_name = ""
header = []
new_topic = ""
counter = 0
heading1 = ""
heading2 = ""
add_topic = False
for paragraph in doc.paragraphs:
    par_text: str = cirilica_u_latinicu(paragraph.text.strip())

    if par_text == "@@@START":
        start = True
        continue

    if par_text == "@@@END":
        start = False

    if not start:
        continue

    if len(par_text) < 2:
        continue

    if (par_text.startswith("Član 46")):
        print ()

    new_topic = ""
    body = ""

    if paragraph.style.name.startswith("Heading 1"):
        heading1 = f"@@@04{par_text}@@@04\n"
        continue
    elif paragraph.style.name.startswith("Heading 2"):
        heading2 = f"@@@03{par_text}@@@03\n"
        continue
    elif paragraph.alignment == 1:
        add_topic = True
        if not (par_text.startswith("Član")):
            header.append(f"@@@00{par_text}@@@00\n")
        else:
            header.append(f"@@@10{par_text}@@@10\n")
        continue
    
    if add_topic:
        new_topic = zakon_rad_clear_head_item(header[-1])
        if len(header) > 1:
            add_txt = zakon_rad_clear_head_item(header[len(header)-2])
            new_topic += f" - {add_txt}"
        else:
            add_txt = zakon_rad_clear_head_item(heading2)
            if heading2:
                new_topic += f" - {add_txt}"

        body += heading1
        body += heading2
        for i in header:
            body += i
        body += "\n"
        header = []
        add_topic = False

    body_flag = ""
    pos = par_text.find(")")
    if pos > 0:
        try:
            _ = int(par_text[:pos])
            par_text = f"@@@01{par_text[:pos+1]}@@@01{par_text[pos+1:]}"
        except:
            _ = 0
    
    is_italic = True
    for run in paragraph.runs:
        if not run.italic:
            is_italic = False
            break

    if par_text.startswith("*") or is_italic:
        body_flag = "@@@02"

    body += f"{body_flag}{par_text}{body_flag}\n"

    new_topic = new_topic.strip()

    if new_topic:
        topic_name = new_topic
        new_topic = ""
        if topic_name in dictionary["serbian"]["dicts"]["zakon_o_radu"]:
            dictionary["serbian"]["dicts"]["zakon_o_radu"][topic_name]["text"] += "\n\n\n"
            print (f"Duplicate Entry : {topic_name}")
        else:
            # print (topic_name)
            dictionary["serbian"]["dicts"]["zakon_o_radu"][topic_name] = {}
            dictionary["serbian"]["dicts"]["zakon_o_radu"][topic_name]["text"] = ""
            dictionary["serbian"]["dicts"]["zakon_o_radu"][topic_name]["links"] = []
            new_topic = ""
            # continue

    if dictionary["serbian"]["dicts"]["zakon_o_radu"][topic_name]["text"]:
        dictionary["serbian"]["dicts"]["zakon_o_radu"][topic_name]["text"] += f" {body}"
    else:
        dictionary["serbian"]["dicts"]["zakon_o_radu"][topic_name]["text"] += body


print ()
print (f'Zakon o Radu: {len(dictionary["serbian"]["dicts"]["zakon_o_radu"])} items.')




# Dusanov zakonik


doc = Document("Dusanov_zakonik.docx")

dictionary["serbian"]["dicts"]["dusan"] = {}

start = False
body = ""
topic_name = ""
header = True
new_topic = ""
counter = 0
recnik_mode = False
clan = 1
has_body = False
for paragraph in doc.paragraphs:
    par_text: str = paragraph.text.strip().replace("\t", "")
    while True:
        par_text = par_text.replace("  ", " ")
        if par_text.find("  ") == -1:
            break

    if par_text == "@@@START":
        start = True
        continue

    if par_text == "@@@END":
        start = False

    if par_text == "@@@RECNIK":
        start = True
        recnik_mode = True
        continue

    if not start:
        continue


    new_topic = ""
    body = ""

    if recnik_mode:
        pos = par_text.find("-")
        if pos == -1:
            print ("Error, delimiter not found  : ", par_text)
            continue
        new_topic = par_text[:pos]
        body = par_text[pos+1:].strip()
    else:
        if not par_text and not paragraph.style.name.startswith("List"):
            if not has_body:
                continue
            clan += 1
            header = True
            has_body = False
            continue
        if header:
            if par_text:
                new_topic = f"Član ({clan}) - {par_text}"
                body = f"@@@01{par_text}@@@01\n\n"
            else:
                new_topic = f"Član ({clan})"
            header = False
        else:
            body += par_text + "\n"
            has_body = True


    new_topic = new_topic.strip()

    if new_topic:
        topic_name = new_topic
        new_topic = ""
        if topic_name in dictionary["serbian"]["dicts"]["dusan"]:
            dictionary["serbian"]["dicts"]["dusan"][topic_name]["text"] += "\n\n\n"
            print (f"Duplicate Entry : {topic_name}")
        else:
            # print (topic_name)
            dictionary["serbian"]["dicts"]["dusan"][topic_name] = {}
            dictionary["serbian"]["dicts"]["dusan"][topic_name]["text"] = ""
            dictionary["serbian"]["dicts"]["dusan"][topic_name]["links"] = []
            new_topic = ""
            # continue

    if dictionary["serbian"]["dicts"]["dusan"][topic_name]["text"]:
        dictionary["serbian"]["dicts"]["dusan"][topic_name]["text"] += f" {body}"
    else:
        dictionary["serbian"]["dicts"]["dusan"][topic_name]["text"] += body


print ()
print (f'Dusanov zakonik: {len(dictionary["serbian"]["dicts"]["dusan"])} items.')




# Zakon o opstem upravnom postupku

def zakon_upravni_clear_head_item(txt: str) -> str:
    for i in range(10):
        txt = txt.replace(f"@@@0{i}", "")
    txt = txt.replace("@@@10", "")
    txt = txt.replace("\n", "")
    return txt

doc = Document("Zakon_o_opstem_upravnom_postupku.docx")

dictionary["serbian"]["dicts"]["zakon_upravni"] = {}

start = False
body = ""
topic_name = ""
header = []
new_topic = ""
counter = 0
title = ""
heading1 = ""
heading2 = ""
add_topic = False
for paragraph in doc.paragraphs:
    par_text: str = cirilica_u_latinicu(paragraph.text.strip())

    if par_text == "@@@START":
        start = True
        continue

    if par_text == "@@@END":
        start = False

    if not start:
        continue

    if len(par_text) < 2:
        continue

    # if (par_text.startswith("Član 46")):
    #     print ()

    new_topic = ""
    body = ""

    if paragraph.style.name.startswith("Title"):
        title = f"@@@10{par_text}@@@10\n"
        continue
    if paragraph.style.name.startswith("Heading 1"):
        heading1 = f"@@@04{par_text}@@@04\n"
        continue
    elif paragraph.style.name.startswith("Heading 2"):
        heading2 = f"@@@03{par_text}@@@03\n"
        continue
    elif paragraph.alignment == 1:
        add_topic = True
        if not (par_text.startswith("Član")):
            header.append(f"@@@00{par_text}@@@00\n")
        else:
            header.append(f"@@@10{par_text}@@@10\n")
        continue
    
    if add_topic:
        new_topic = zakon_upravni_clear_head_item(header[-1])
        if len(header) > 1:
            add_txt = zakon_upravni_clear_head_item(header[len(header)-2])
            new_topic += f" - {add_txt}"
        else:
            add_txt = zakon_upravni_clear_head_item(heading2)
            if heading2:
                new_topic += f" - {add_txt}"

        body += title
        body += heading1
        body += heading2
        for i in header:
            body += i
        body += "\n"
        header = []
        add_topic = False

    body_flag = ""
    pos = par_text.find(")")
    if pos > 0:
        try:
            _ = int(par_text[:pos].replace("(", ""))
            par_text = f"@@@01{par_text[:pos+1]}@@@01{par_text[pos+1:]}"
        except:
            _ = 0
    
    is_italic = True
    for run in paragraph.runs:
        if not run.italic:
            is_italic = False
            break

    if par_text.startswith("*") or is_italic:
        body_flag = "@@@02"

    body += f"{body_flag}{par_text}{body_flag}\n"

    new_topic = new_topic.strip()

    if new_topic:
        topic_name = new_topic
        new_topic = ""
        if topic_name in dictionary["serbian"]["dicts"]["zakon_upravni"]:
            dictionary["serbian"]["dicts"]["zakon_upravni"][topic_name]["text"] += "\n\n\n"
            print (f"Duplicate Entry : {topic_name}")
        else:
            # print (topic_name)
            dictionary["serbian"]["dicts"]["zakon_upravni"][topic_name] = {}
            dictionary["serbian"]["dicts"]["zakon_upravni"][topic_name]["text"] = ""
            dictionary["serbian"]["dicts"]["zakon_upravni"][topic_name]["links"] = []
            new_topic = ""
            # continue

    if dictionary["serbian"]["dicts"]["zakon_upravni"][topic_name]["text"]:
        dictionary["serbian"]["dicts"]["zakon_upravni"][topic_name]["text"] += f"{body}"
    else:
        dictionary["serbian"]["dicts"]["zakon_upravni"][topic_name]["text"] += body


print ()
print (f'Zakon o opstem upravnom postupku: {len(dictionary["serbian"]["dicts"]["zakon_upravni"])} items.')




# Zakoni razni

def zakon_ostali_clear_head_item(txt: str) -> str:
    for i in range(10):
        txt = txt.replace(f"@@@0{i}", "")
    txt = txt.replace("@@@10", "")
    txt = txt.replace("\n", "")
    return txt

def zakon_ostali_is_clan(txt: str) -> bool:
    if txt.count(" ") > 3:
        return False
    
    txt = txt.lower()
    if not txt.startswith("član") and not txt.startswith("čla n"):
        return False
    
    return True


doc = Document("ostali_zakoni.docx")

dictionary["serbian"]["dicts"]["zakon_razni"] = {}

# zakoni_map : [search_string, topic_prefix, cyrilic]
zakoni_map = [
    ["@@@START ZOMRRS", "ZOMRRS", True],
    ["@@@START ZOSZNR", "ZOSZNR", True],
    ["@@@START ZOV", "ZOV", True],
    ["@@@START ZOAZ", "ZOAZ", True],
    ["@@@START ZOPR", "ZOPR", True],
    ["@@@START ZOPII", "ZOPII", False],
    ["@@@START ZOVO", "ZOVO", False],
    ["@@@START ZOS", "ZOS", False],
    ["@@@START ZOSBiDO", "ZOSBiDO", True],
    ["@@@START ZOSB", "ZOSB", False],
    ["@@@START ZOPPSL", "ZOPPSL", False],
    ["@@@START ZOSIOZ", "ZOSIOZ", True],
    ["@@@START ZOLN", "ZOLN", True],
    ["@@@START ZOUO", "ZOUO", False],
    ["@@@START ZOZŽS", "ZOZŽS", False],
    ["@@@START ZOH", "ZOH", False],
    ["@@@START ZOTOT", "ZOTOT", False],
    ["@@@START ZOAIAO", "ZOAIAO", False],
    ["@@@START POKIIKO", "POKIIKO", False],
    ["@@@START ZOZPOL", "ZOZPOL", False],
    ["@@@START ZOPVN", "ZOPVN", True],
    ["@@@START ZOPDV", "ZOPDV", True],
    ["@@@START ZOJN", "ZOJN", True],
    ["@@@START ZODŽ", "ZODŽ", True],
    ["@@@START ZORR", "ZORR", False],
    ["@@@START ZOizb", "ZOizb", False],
    ["@@@START ZODO", "ZODO", False],
    ["@@@START ZOEK", "ZOEK", False],
    ["@@@START ZOOO", "ZOOO", False],
    ["@@@START ZOOOIV", "ZOOOIV", False],
    ["@@@START ZOLS", "ZOLS", False],
    ["@@@START ZOSUJIP", "ZOSUJIP", False],
    ["@@@START ZOU", "ZOU", False]

]
for zakon in zakoni_map:
    zakon_started = False
    print (f"Working on {zakon[1]} ... searching ... ", end="")
    start = False
    body = ""
    topic_name = ""
    header = []
    new_topic = ""
    title = ""
    heading1 = ""
    heading2 = ""
    add_topic = False
    topic_prefix = zakon[1] + ": "
    clan_num = 0

    for paragraph in doc.paragraphs:
        if zakon[2]:
            par_text: str = cirilica_u_latinicu(paragraph.text.strip())
        else:
            par_text: str = paragraph.text.strip()

        par_text = par_text.replace("ñ", "đ")

        if par_text == zakon[0]:
            print ("found ... processing ... ", end="")
            zakon_started = True
            start = True
            continue

        if par_text == "@@@END":
            if zakon_started:
                print ("Done.")
            zakon_started = False
            start = False

        if not start:
            continue

        if not par_text:
            continue

        new_topic = ""
        body = ""

        if paragraph.alignment == 1 and (par_text.lower().find("brisan") >= 0 or par_text.lower().find("(prestalo da važi)") >= 0):
            print (f"Warning, alignment not set properly : {par_text}", end="")

        if paragraph.style.name.startswith("Title"):
            title = f"@@@10{par_text}@@@10\n"
            continue
        if paragraph.style.name.startswith("Heading 1"):
            heading1 = f"@@@04{par_text}@@@04\n"
            continue
        elif paragraph.style.name.startswith("Heading 2"):
            heading2 = f"@@@03{par_text}@@@03\n"
            continue
        elif paragraph.alignment in [1] or paragraph.style.name.startswith("Heading 3") or paragraph.style.name.startswith("Heading 4") or zakon_ostali_is_clan(par_text):
            add_topic = True
            if not (par_text.startswith("Član")):
                header.append(f"@@@00{par_text}@@@00\n")
            else:
                clan_num = 0
                header.append(f"@@@10{par_text}@@@10\n")
            continue
        
        if add_topic:
            new_topic = topic_prefix + zakon_ostali_clear_head_item(header[-1])
            if len(header) > 1:
                add_txt = zakon_ostali_clear_head_item(header[len(header)-2])
                new_topic += f" - {add_txt}"
            else:
                if heading2.replace("-", "").strip():
                    add_txt = zakon_ostali_clear_head_item(heading2)
                    new_topic += f" - {add_txt}"
                else:
                    add_txt = zakon_ostali_clear_head_item(heading1)
                    new_topic += f" - {add_txt}"

            body += title
            body += heading1
            body += heading2
            for i in header:
                body += i
            body += "\n"
            header = []
            add_topic = False

        if paragraph.style.name.startswith("List"):
            clan_num += 1
            par_text = f"@@@01({clan_num}) @@@01{par_text}"
        else:
            body_flag = ""
            pos = par_text.find(")")
            if pos > 0:
                try:
                    _ = int(par_text[:pos].replace("(", ""))
                    par_text = f"@@@01{par_text[:pos+1]}@@@01{par_text[pos+1:]}"
                except:
                    _ = 0
        
        is_italic = True
        for run in paragraph.runs:
            if not run.italic:
                is_italic = False
                break

        if par_text.startswith("*") or is_italic:
            body_flag = "@@@02"

        body += f"{body_flag}{par_text}{body_flag}\n"

        new_topic = new_topic.strip(" -")

        if new_topic:
            topic_name = new_topic
            new_topic = ""
            if topic_name in dictionary["serbian"]["dicts"]["zakon_razni"]:
                counter = 1
                while True:
                    topic_name = f"{topic_name} ({counter})"
                    if topic_name not in dictionary["serbian"]["dicts"]["zakon_razni"]:
                        break
                    counter += 1
                print (f"\nDuplicate Entry Corrected to : {topic_name}")
            
            # print (topic_name)
            dictionary["serbian"]["dicts"]["zakon_razni"][topic_name] = {}
            dictionary["serbian"]["dicts"]["zakon_razni"][topic_name]["text"] = ""
            dictionary["serbian"]["dicts"]["zakon_razni"][topic_name]["links"] = []
            new_topic = ""
            # continue

        if dictionary["serbian"]["dicts"]["zakon_razni"][topic_name]["text"]:
            dictionary["serbian"]["dicts"]["zakon_razni"][topic_name]["text"] += f"{body}"
        else:
            dictionary["serbian"]["dicts"]["zakon_razni"][topic_name]["text"] += body


print ()
print (f'Zakon razni, TOTAL : {len(zakoni_map)} docs.  Processed: {len(dictionary["serbian"]["dicts"]["zakon_razni"])} items.')



# Ustav

def ustav_clear_head_item(txt: str) -> str:
    for i in range(10):
        txt = txt.replace(f"@@@0{i}", "")
    txt = txt.replace("@@@10", "")
    txt = txt.replace("\n", "")
    return txt

doc = Document("ustav.docx")

dictionary["serbian"]["dicts"]["ustav"] = {}

start = False
body = ""
topic_name = ""
header = []
new_topic = ""
title = ""
heading1 = ""
heading2 = ""
add_topic = False
topic_prefix = ""
clan_num = 0

for paragraph in doc.paragraphs:
    par_text: str = cirilica_u_latinicu(paragraph.text.strip())

    par_text = par_text.replace("ñ", "đ")

    if par_text == "@@@START":
        start = True
        continue

    if par_text == "@@@END":
        start = False

    if not start:
        continue

    if not par_text:
        continue

    # if (par_text.startswith("Član 46")):
    #     print ()

    new_topic = ""
    body = ""

    if paragraph.alignment == 1 and (par_text.lower().find("brisan") or par_text.lower().find("(prestalo da važi)")):
        print (f"Warning, alignment not set properly : {par_text}")

    if paragraph.style.name.startswith("Title"):
        title = f"@@@10{par_text}@@@10\n"
        continue
    if paragraph.style.name.startswith("Heading 1"):
        heading1 = f"@@@04{par_text}@@@04\n"
        continue
    elif paragraph.style.name.startswith("Heading 2"):
        heading2 = f"@@@03{par_text}@@@03\n"
        continue
    elif paragraph.alignment in [1] or paragraph.style.name.startswith("Heading 3") or paragraph.style.name.startswith("Heading 4"):
        add_topic = True
        if not (par_text.startswith("Član")):
            header.append(f"@@@00{par_text}@@@00\n")
        else:
            clan_num = 0
            header.append(f"@@@10{par_text}@@@10\n")
        continue
    
    if add_topic:
        new_topic = topic_prefix + ustav_clear_head_item(header[-1])
        if len(header) > 1:
            add_txt = ustav_clear_head_item(header[len(header)-2])
            new_topic += f" - {add_txt}"
        else:
            if heading2.replace("-", "").strip():
                add_txt = ustav_clear_head_item(heading2)
                new_topic += f" - {add_txt}"
            else:
                add_txt = ustav_clear_head_item(heading1)
                new_topic += f" - {add_txt}"

        body += title
        body += heading1
        body += heading2
        for i in header:
            body += i
        body += "\n"
        header = []
        add_topic = False

    if paragraph.style.name.startswith("List"):
        clan_num += 1
        par_text = f"@@@01({clan_num}) @@@01{par_text}"
    else:
        body_flag = ""
        pos = par_text.find(")")
        if pos > 0:
            try:
                _ = int(par_text[:pos].replace("(", ""))
                par_text = f"@@@01{par_text[:pos+1]}@@@01{par_text[pos+1:]}"
            except:
                _ = 0
    
    is_italic = True
    for run in paragraph.runs:
        if not run.italic:
            is_italic = False
            break

    if par_text.startswith("*") or is_italic:
        body_flag = "@@@02"

    body += f"{body_flag}{par_text}{body_flag}\n"

    new_topic = new_topic.strip()

    if new_topic:
        topic_name = new_topic
        new_topic = ""
        if topic_name in dictionary["serbian"]["dicts"]["ustav"]:
            counter = 1
            while True:
                topic_name = f"{topic_name} ({counter})"
                if topic_name not in dictionary["serbian"]["dicts"]["ustav"]:
                    break
                counter += 1
            print (f"Duplicate Entry Corrected to : {topic_name}")
        
        # print (topic_name)
        dictionary["serbian"]["dicts"]["ustav"][topic_name] = {}
        dictionary["serbian"]["dicts"]["ustav"][topic_name]["text"] = ""
        dictionary["serbian"]["dicts"]["ustav"][topic_name]["links"] = []
        new_topic = ""
        # continue

    if dictionary["serbian"]["dicts"]["ustav"][topic_name]["text"]:
        dictionary["serbian"]["dicts"]["ustav"][topic_name]["text"] += f"{body}"
    else:
        dictionary["serbian"]["dicts"]["ustav"][topic_name]["text"] += body


print ()
print (f'Ustav: {len(dictionary["serbian"]["dicts"]["ustav"])} items.')








# FIX SPECIAL CHARACTER IN ALL DICTIONARIES
# with open("dictionary.json", "r", encoding="utf-8") as file:
#     dictionary = json.load(file)

from unidecode import unidecode

SER_MAP = [
    ["ć", "|c0"], 
    ["č", "|c1"],
    ["Ć", "|C0"], 
    ["Č", "|C1"],
    ["đ", "|d0"],
    ["Đ", "|D0"],
    ["ž", "|z0"],
    ["Ž", "|Z0"],
    ["š", "|s0"],
    ["Š", "|S0"],
]

# TEXT_WRAPPING  { dict } = [ character(s), occurrences ]
# occurrences = 1 first occurrence, -1 all occurrences
TEXT_WRAPPING = {
    "bos": [";", -1],
    "vujaklija": [";", -1],
    "google&ms": [":;", -1]
}

TEXT_REPLACE = {
    "filoz2": [
        ["-   ", "", None],
        ["-  ", "", None],
        ["- ", "", None],
        [" Pet", "", "last"]
        ],
    "svet_mit": [
        ["-   ", "", None],
        ["-  ", "", None],
        ["- ", "", None]
        ],
    "hrt": [
        ["-   ", "", None],
        ["-  ", "", None],
        ["- ", "", None]
        ],
    "jez_nedoum": [
        ["-   ", "", None],
        ["-  ", "", None],
        ["- ", "", None]
        ],
    "bibliotek": [
        [" - ", "~~~", None],
        ["\t", " ", None],
        ["-   ", "", None],
        ["-  ", "", None],
        ["- ", "", None],
        ["~~~", " - ", None]
        ],
        "lov": [
            ["@@@", "\n", None]
        ],
    "latin": [
        ["-   ", "", None],
        ["-  ", "", None],
        ["- ", "", None]
        ],
    "onkoloski": [
        ["-   ", "", None],
        ["-  ", "", None],
        ["- ", "", None]
        ],
    "biologija": [
        [" - ", "~~~", None],
        ["-   ", "", None],
        ["-  ", "", None],
        ["- ", "", None],
        ["~~~", " - ", None]
        ]

}


def encode_ser_char(text: str) -> str:
    for i in SER_MAP:
        text = text.replace(i[0], i[1])
    return text

def decode_ser_char(text: str, fix_keyword: bool = False) -> str:
    for i in SER_MAP:
        text = text.replace(i[1], i[0])
    if fix_keyword:
        map_for_k = [">>", "<<", "(tm)", "..", ",,"]
        for i in map_for_k:
            text = text.replace(i, "")
    return text


print ()
print ("Fixing keywords...", end="")
count = 0
key_to_change = []
for i in dictionary["serbian"]["dicts"]:
    for j in dictionary["serbian"]["dicts"][i]:
        new_key = encode_ser_char(j)
        new_key = unidecode(new_key)
        new_key = decode_ser_char(new_key, fix_keyword=True)
        if new_key != j:
            count += 1
            if new_key in dictionary["serbian"]["dicts"][i]:
                print (f"ERROR.  Key already exist: {new_key} -- changed to :{new_key} (1)")
                new_key += " (1)"
            key_to_change.append([i, new_key, j])
            
for i in key_to_change:
    dictionary["serbian"]["dicts"][i[0]][i[1]] = dictionary["serbian"]["dicts"][i[0]].pop(i[2])

print (f" {count} fixed.")

print ("Fixing text...", end="")
count = 0
for i in dictionary["serbian"]["dicts"]:
    for j in dictionary["serbian"]["dicts"][i]:
        if j == "plug-in":
            print ("Found")
        new_text = encode_ser_char(dictionary["serbian"]["dicts"][i][j]["text"])
        new_text = unidecode(new_text)
        new_text = decode_ser_char(new_text)
        if i in TEXT_WRAPPING:
            for k in TEXT_WRAPPING[i][0]:
                new_text = new_text.replace(k, f"{k}\n", TEXT_WRAPPING[i][1])

        if i in TEXT_REPLACE:
            for k in TEXT_REPLACE[i]:
                if k[2]:
                    if k[2] == "last" and len(new_text) >= len(k[0]):
                        if new_text[-len(k[0]):] == k[0]:
                            new_text = new_text[:-len(k[0])] + k[1]
                else:
                    new_text = new_text.replace(k[0], k[1])

        if new_text != dictionary["serbian"]["dicts"][i][j]["text"]:
            count += 1
            dictionary["serbian"]["dicts"][i][j]["text"] = new_text
print (f" {count} fixed.")



# FILL LINKS IN ALL DICTIONARIES

def _links_add(link: str, text: str, links: list):
    original_link = link
    original_text = text
    original_text_to_compare = text
    text = f" {text.lower()} "

    for i in END_OF_TOPIC:
        text = text.replace(i, " ")

    link_repl = ["=", "(", ")", "[", "]"]
    for i in link_repl:
        link = link.replace(i, ",")
    link_list = [f' {x.lower().strip()} ' for x in link.split(",") if len(x.strip()) > 2]
    link_list.append(f" {original_link.lower()} ")

    text = f" {text.strip()} "
    for i in link_list:
        for j in ["():;[]{}.,!?\"'\\/><=*~"]:
            if original_text.find(f"{j}{i.strip()}{j}") >= 0:
                original_text = original_text.replace(f"{j}{i.strip()}{j}", f"{j} {i.strip()} {j}")
                text = text.replace(f"{j}{i.lower().strip()}{j}", f"{j} {i.lower().strip()} {j}")
            if original_text.find(f"{j}{i.strip()}") >= 0:
                original_text = original_text.replace(f"{j}{i.strip()}", f"{j} {i.strip()}")
                text = text.replace(f"{j}{i.lower().strip()}", f"{j} {i.lower().strip()}")
            if original_text.find(f"{i.strip()}{j}") >= 0:
                original_text = original_text.replace(f"{i.strip()}{j}", f"{i.strip()} {j}")
                text = text.replace(f"{i.lower().strip()}{j}", f"{i.lower().strip()} {j}")
                
        if text.find(i) >= 0:
            if i.strip() not in [x[0] for x in links]:
                links.append([i.strip(), original_link])
    
    if original_text != original_text_to_compare:
        return original_text
    else:
        return None

def _links_fix(links: list):
    remove_list = []
    for idx, i in enumerate(links):
        for j in links:
            if i != j:
                if f" {j[0]} ".find(f" {i[0]} ") >= 0:
                    if idx not in remove_list:
                        remove_list.append(idx)
    remove_list.sort(reverse=True)
    for i in remove_list:
        links.pop(i)


print ("Searching for links ... ")
count = 0
for i in dictionary["serbian"]["dicts"]:
    dict_t = time.perf_counter()
    print ()
    print (f"Start: {i} ", end="")
    # if i == "vujaklija":
    #     print ("Skip VUJAKLIJA")
    #     continue
    if i == "en-sr":
        print ("Skip EN-SR")
        continue
    percent = int(len(dictionary["serbian"]["dicts"][i]) / 20)
    count1 = 0
    for j in dictionary["serbian"]["dicts"][i]:
        count1 += 1
        if count1 % percent == 0:
            print (f'{int(count1/len(dictionary["serbian"]["dicts"][i])*100)}% ', end="")

        dictionary["serbian"]["dicts"][i][j]["links"] = []
        for k in dictionary["serbian"]["dicts"][i]:
            if k != j:
                result = _links_add(k, dictionary["serbian"]["dicts"][i][j]["text"], dictionary["serbian"]["dicts"][i][j]["links"])
                if result:
                    dictionary["serbian"]["dicts"][i][j]["text"] = result
        count += len(dictionary["serbian"]["dicts"][i][j]["links"])
    print (f"  100% Done.  TIME: {(time.perf_counter() - dict_t)/60: ,.2f} min")

print ()
print (f'{count} links in total found.')

print ("Fixing links ... ", end="")
count = 0
for i in dictionary["serbian"]["dicts"]:
    for j in dictionary["serbian"]["dicts"][i]:
        prev = len(dictionary["serbian"]["dicts"][i][j]["links"])
        _links_fix(dictionary["serbian"]["dicts"][i][j]["links"])
        count += prev - len(dictionary["serbian"]["dicts"][i][j]["links"])

print (f'{count} links fixed.')


# Add names and description for all dictionaries

names_map = [
    [
        """mit""",
        ":mitologija",
        """Srpski Mitološki Rečnik (Kulišić)""",
        """Srpski Mitoloski Recnik (Kulisic),
Godina izdanja: 1970.
Tip: Jednojezični
Jezik: Srpski
Š. KULIŠIĆ / P.Ž. PETROVIĆ / N. PANTELIĆ
SRPSKI MITOLOŠKI REČNIK
NOLIT
BEOGRAD
1970

BIBLIOTEKA SINTEZE

REČNIK OBUHVATA MITOLOGIJU SRBA I CRNOGORACA
O RELIGIJI STARIH SRBA I JUŽNIH SLOVENA""" 
    ],
    [
        """vujaklija""",
        "",
        """Rečnik Stranih Reči (Vujaklija)""",
        """Recnik Stranih Reci i Izraza (Milan Vujaklija),
Godina: 1980.
REDAKTORI 
Dr SVETOMIR RISTIĆ i dr RADOMIR ALEKSIĆ 
DOPUNILI, PROŠIRILI I REDIGOVALI ŠTAMPANO IZDANjE 
Dr RADOMIR ALEKSIĆ i REDAKCIJA »P R O S V E T E «
Prvo izdanje Leksikona stranih reči i izraza Milana Vujaklije objavljeno je 1937. godine.
Drugo izdanje, u redakciji dr Svetomira Ristića i dr Radomira Aleksića, izašlo je 1954. godine i štampane, do danas, u tiražu od preko 200.000 primeraka.
Dobar prijem na koji je Leksikon naišao u najširim krugovima korisnika podstakao je Redakciju Prosvete da pripremi ovo novo, popunjeno i ažurirano izdanje Leksikona Milana Vujaklije.""" 
    ],
    [
        """san""",
        ":sanjati",
        """Tumačenje Snova (Sanovnik)""",
        """Tumacenje Snova (Sanovnik),
Oniromantija - značenje i tumačenje snova, pored astrologije, ovom tehnikom ljudi se bave još od najstarijih vremena.
Tumačenje snova je spominjano čak i u starom zavetu, gde Josip tumači faraonove snove.
Nauka danas smatra da sadržaji snova odražavaju slučajnu situaciju neuralnih krugova mozga,
bilo radi spontanog izbijanja ili zbog delovanja okoline,
te tendencije mozga da unese red u nepravilnoj aktivnosti.
Psiholozi su bili dugo pod uticajem Frojda koji je smatrao da su snovi odraz naših potisnutih i podsvesnih želja i misli,
te naročito naših seksualnih konflikata.
Prema Frojdu svi naši snovi su bez izuzetka ispunjenje neke želje.
Fizički stimulans sam po sebi nije odgovoran za stvaranje snova.
Naše želje su samo impuls za snove. Mehanizam snova je veoma suptilan.""" 
    ],
    [
        """zargon""",
        "",
        """Rečnik Žargona""",
        """Recnik Žargona,
Žargon je specifična vrsta reči koja pripada jednoj grupi, sredini ili neformalnoj skupini.
Upotreba žargona može biti i zatvorena komunikacija unutar grupe radi zaštite njenih interesa.
Žargon se upotrebljava i kao element književnog jezika.

Leksemi koji se koriste u žargonu nazivaju se žargonizmima.""" 
    ],
    [
        """bos""",
        "",
        """Rečnik Bosanskih Izraza""",
        """Recnik Bosanskih Izraza,
Izrazi koji se koriste na podrucuju Bosne i Hercegovine.""" 
    ],
    [
        """en-sr""",
        "",
        """Englesko-Srpski Recnik""",
        """Englesko-Srpski Recnik""" 
    ],
    [
        """psiho""",
        "psihologija",
        """Rečnik Psihologije""",
        """Recnik Psihologije,
Izrazi koji se upotrebljavaju u Psihologiji.

Psihologija je nauka koja proučava psihu, psihičke procese i psihičke osobine u njihovom nastanku, razvoju i objektivnim pojavljivanjima kod čoveka i životinja.
Saznanja psihologije počivaju na dve vrste podataka. Na podacima dobijenim posmatranjem ponašanja (tzv. objektivni podaci) i podacima neposrednog iskustva (tzv. subjektivni podaci) dobijenim samoposmatranjem (introspekcijom).

Psihologija proučava individualna i grupna ponašanja. Sama reč psihologija potiče od grčkih reči psyha i logos i doslovno znači „nauka o duši”.
Tokom istorije, još od antičkog doba, psihologija je bila grana filozofije, da bi u XVIII i 19. veku, započela svoj „vlastiti put” kao nezavisna naučna disciplina.
Psihologija se kao naučna disciplina zvanično počela računati 1879. kada je nemački filozof i psiholog Vilhelm Vunt u Lajpcigu osnovao Prvu psihološku laboratoriju, tzv. Prva laboratorija ljudske duše.
Ipak, danas je psihologija vrlo vezana za biologiju i društvene nauke.

Psihologija proučava dve tačke psihičkog, odnosno duševnog života ljudi, koje čine dve vrste psihičkih pojava, a to su:
Psihički procesi
Psihičke osobine""" 
    ],
    [
        """stari_izrazi""",
        ":stara rec",
        """Rečnik Starih Srpskih Izraza""",
        """Recnik Starih Srpskih Izraza,
Stari, cesto zastareli, srpski izrazi koji se retko ili uopste vise ne koriste.""" 
    ],
    [
        """filoz""",
        ":filozofija",
        """Rečnik Filozofije""",
        """Recnik Filozofije,
Recnik termina koji se koriste u Filozofiji.

Filozofija (grč. φιλοσοφία) jeste nauka koja se bavi opštim i temeljnim problemima u vezi sa stvarnošću, postojanjem, znanjem, vrednostima, razumom, umom i jezikom.
Leksikon stranih reči i izraza definiše filozofiju kao „mudrost, naučni rad na izgrađivanju opšteg pogleda na svet i sam taj pogled na svet“.

Reč filozofija je nastala u starogrčkom jeziku u obliku φιλοσοφία, a posle je preneta u lat. philosophia i u bukvalnom prevodu znači ljubav prema mudrosti.
Pitagori se pripisuje da je prvi upotrebio reči filozof i filozofija.
Reč je uvedena kao suprotnost sofistima — mudracima važnim u antičkoj Grčkoj koji su prodavali svoje znanje kao učitelji — dok se filozofi, kao ljudi koji vole mudrost, nisu bavili mudrošću zbog novca.

Istorijski, filozofija je obuhvatala sve oblike znanja.
Od vremena antičkog grčkog filozofa Aristotela do 19. veka, „prirodna filozofija“ je obuhvatala astronomiju, medicinu i fiziku.
Na primer, Njutnova knjiga iz 1687. godine se zvala Matematički principi prirodne filozofije, a kasnije je klasifikovana kao knjiga fizike.
U 19. veku je razvoj modernih istraživačkih univerziteta doveo do profesionalizacije i specijalizacije akademske filozofije i drugih disciplina.
U moderno doba, neka istraživanja koja su tradicionalno bila deo filozofije su postala zasebne akademske discipline, uključujući psihologiju, sociologiju, lingvistiku i ekonomiju.""" 
    ],
    [
        """emo""",
        ":emocija",
        """Rečnik Emocija""",
        """Recnik Izraza Emocija,
Spisak termina koji se koriste za opis raznih emotivnih stanja.

Emocija je uzbuđenost izazvana situacijom ili stimulusom koji je osobi važan.
Sastoji se iz tri komponente: fiziološke, izražajne i subjektivne.
Fiziološka komponenta priprema organizam za adekvatno reagovanje, izražavanjem se obavlja neverbalna komunikacija, a preko doživljaja spoznajemo u kakvom smo stanju.
Emocije imaju adaptivnu i komunikativnu funkciju. Dele se na osnovne, složene, prijatne, neprijatne.
Osnovne emocije su urođene reakcije na draži, i izražavaju se kod svih ljudi na sličan način. To su sreća, tuga, ljutnja, strah, gađenje i iznenađenje.
Afekt, raspoloženje i sentiment su emocionalna stanja različitog intenziteta i trajanja, a afektivni ton je doživljaj prijatnosti i neprijatnosti i sastavni je deo svakog osećanja.""" 
    ],
    [
        """biljke""",
        "biljka",
        """Rečnik Verovanja o Biljkama""",
        """Recnik Verovanja o Biljkama,
Veselin Čajkanović
REČNIK SRPSKIH NARODNIH VEROVANjA O BILjKAMA
""" 
    ],
    [
        """it""",
        ":informatika",
        """Rečnik IT Izraza""",
        """Recnik IT Izraza,
Recnik termina koji se koriste u IT oblasti.

Informacionu tehnologiju (IT) (engl. information technology (IT)) Američka asocijacija za informacione tehnologije definiše kao
"izučavanje, projektovanje, razvoj, implementacija (sprovođenje) i podrška ili upravljanje računarskim informacionim sistemima (IS), softverskim aplikacijama i hardverom".
IT koriste računare i računarske programe da pretvore, uskladište (smeste), štite, obrade, bezbedno šalju i primaju informacije.

Termin (pojam) „informaciona tehnologija” često obuhvata i znatno šire polje oblasti tehnologije.
Sve one aktivnosti kojima se IT profesionalci bave, od instalacija aplikativnih programa do projektovanja složenih računarskih mreža i informacionih sistema.
Neke od tih aktivnosti su: umrežavanje i inženjering računarskog hardvera, dizajniranje softvera i baza podataka, kao i upravljanje i administracija informacionim sistemom.
Informaciona tehnologija je opšti termin koji opisuje tehnologiju koja pomaže proizvodnji, manipulaciji, skladištenju, komunikaciji i distribuciji informacija.

Prvi koji je upotrebio termin „Informacione tehnologije” bio je Džim Domsik iz Mičigena i to novembra 1981. godine.
Termin je upotrebio kako bi modernizovao do tada korišćeni izraz „obrada podataka”.
U to vreme Domsik je radio kao računarski menadžer u automobilskoj industriji.

Standarde u ovoj oblasti definisale su organizacije ABET i ACM.""" 
    ],
    [
        "bokeljski",
        "",
        "Bokeljski rečnik",
        """Bokeljima se smatraju stanovnici Boke Kotorske (gradovi Kotor, Tivat, Risan i Herceg Novi, a u istorijskom smislu i Budva)
        koja se nalazi u jugozapadnom delu današnje Crne Gore.
        Većina ih je pravoslavne veroispovesti, dok postoji i jedan broj katolika.
        Bokelji su dominantno srpskog porekla i srpskog nacionalnog osećaja i identiteta.

Bokelji imaju i svoju narodnu nošnju i svi, bez obzira na versko opredeljenje, slave krsnu slavu, karakterističan srpski pravoslavni običaj.
Po istraživanju akademika Slavka Mijuškovića, kod Bokelja katoličke vere postojala je svest o crnogorskom poreklu, slavili su i dalje krsne slave i poštovali običaje."""
    ],
    [
        "bank",
        ":bankarstvo",
        "Glosar Bankarstva i Osiguranja",
        """ENGLESKO-HRVATSKI
GLOSAR
bankarstva, osiguranja i ostalih financijskih usluga

Nakladnik: Ministarstvo vanjskih poslova i europskih integracija
Glavni urednik: mr. sc. Jasminka Novak
Autori: dr.sc. Ante Babić, Goranka Cvijanović Vuković, Silvija Lučić, Sandra Papac
Suradnici: Saša Cimeša, Vedran Duvnjak, Lidija Jurilj, Damir Kaufman, Dubravka Kobaš, Vlatka Pirš, Blaženka Razum, Koraljka Sansović, Maja Savin, Dalibor Vrgoč, Ankica Zerec, Zrinka Živković Leksikografski savjetnik: prof. dr. sc. Maja Bratanić
Lektura: Institut za hrvatski jezik i jezikoslovlje
Redaktor: dr. sc. Branka Tafra
Grafičko oblikovanje: Bee biro
Računalna obrada jezične građe: Miljenka Prohaska Kragović"""
    ],
    [
        "google&ms",
        ":informatika",
        "Rečnik Google i Microsoft Termina",
        """Ovaj rad predstavlja korpus najčešćih i najpotrebnijih termina iz sfere informacionih tehnologija na engleskom jeziku i njihovih objašnjenja,
        kao i prevodnih adaptacija na srpskom koje se javljaju u srpskim lokalizacijama proizvoda kompanija Google i Microsoft.
        Korpus termina i njihovih prevoda sastavljen je prema najčešće korišćenim terminima na sajtovima
        Microsoft pomoći i podrške (http://support.microsoft.com/?ln=sr)
        i Google centra za podršku (https://support.google.com/?hl=sr),
        koji koriste najnoviju terminologiju koja se primenjuje i u proizvodima ove dve kompanije."""
    ],
    [
        "fraze",
        "",
        "Frazeoloski recnik Djordje Otasevic",
        """Srpski jezik bogat je frazeologizmima (idiomima, frazemima).
        Iako je ovaj rečnik okrenut prvenstveno savremenom standardnom jeziku, svoje mesto u njemu našao je i jedan broj pokrajinskih i zastarelih frazeologizama
        jer se sreću u mnogim značajnim delima srpske književnosti.
        Frazeologizmi koji sadrže opscenu leksiku nisu uneti u rečnik.
        Pri izradi su korišćeni brojni rečnici, priručnici, monografije i naučni radovi.
        Najvažniji rečnici su:
        Rečnik srpskohrvatskog književnog i narodnog jezika 1-18 (SANU, Beograd, 1959-),
        Rečnik srpskohrvatskoga književnog jezika 1-6 (Matica srpska, Novi Sad, 1967-1976)
        Frazeološki rječnik hrvatskoga ili srpskog jezika Josipa Matešića (Zagreb, 1982).
        Noviji frazeologizmi, kojih u pomenutim rečnicima, razumljivo, nije moglo da bude, obrađeni su na osnovu autorove građe za Veliki rečnik novih i nezabeleženih reči."""
    ],
    [
        "ekonom",
        "ekonomija",
        "Ekonomski Rečnik Miomir Jaksić",
        """Ovaj recnik sadzi 1100 osnovnih pojmova koji se najcesce srecu u udzbenicima i ekonomskoj literaturi.

Autori:
prof. dr Miomir JAKSIC
mr Nikola FABRIS
mr Alesandra PRASCEVIC"""
    ],
    [
        "ekonom2",
        "ekonomija",
        "Rečnik Ekonomskih Pojmova",
        """Ekonomija je naučna disciplina koja proučava osnovna pravila ponašanja i ekonomske zakonitosti u ekonomskim aktivnostima.
        U svakoj epohi razvoja, ekonomija proučava ekonomske aktivnosti, kako društvo koristi oskudne resurse radi proizvodnje dobara i usluga i vrši njihovu raspodelu među članovima društva.
        Ekonomija je povezana i sa drugim naukama: sociologijom, demografijom, politikom i drugim."""
    ],
    [
        "proces",
        "pravni izraz",
        "Glosar Procesnih Pojmova",
        """GLOSAR PROCESNIH POJMOVA
Ovo su pojmovi koji se koriste u svim ili u pojedinim sudskim postupcima.
Mnogi od njih se sreću svakodnevno, a pojedini su nastali u praksi tako da se ne mogu pronaći u procesnim zakonima"""
    ],
    [
        """filoz2""",
        ":filozofija",
        """Filozofijski Rječnik Vladimira Filipovića""",
        """Recnik Filozofije,
Recnik termina koji se koriste u Filozofiji.

Filozofija (grč. φιλοσοφία) jeste nauka koja se bavi opštim i temeljnim problemima u vezi sa stvarnošću, postojanjem, znanjem, vrednostima, razumom, umom i jezikom.
Leksikon stranih reči i izraza definiše filozofiju kao „mudrost, naučni rad na izgrađivanju opšteg pogleda na svet i sam taj pogled na svet“.

Reč filozofija je nastala u starogrčkom jeziku u obliku φιλοσοφία, a posle je preneta u lat. philosophia i u bukvalnom prevodu znači ljubav prema mudrosti.
Pitagori se pripisuje da je prvi upotrebio reči filozof i filozofija.
Reč je uvedena kao suprotnost sofistima — mudracima važnim u antičkoj Grčkoj koji su prodavali svoje znanje kao učitelji — dok se filozofi, kao ljudi koji vole mudrost, nisu bavili mudrošću zbog novca.

Istorijski, filozofija je obuhvatala sve oblike znanja.
Od vremena antičkog grčkog filozofa Aristotela do 19. veka, „prirodna filozofija“ je obuhvatala astronomiju, medicinu i fiziku.
Na primer, Njutnova knjiga iz 1687. godine se zvala Matematički principi prirodne filozofije, a kasnije je klasifikovana kao knjiga fizike.
U 19. veku je razvoj modernih istraživačkih univerziteta doveo do profesionalizacije i specijalizacije akademske filozofije i drugih disciplina.
U moderno doba, neka istraživanja koja su tradicionalno bila deo filozofije su postala zasebne akademske discipline, uključujući psihologiju, sociologiju, lingvistiku i ekonomiju.""" 
    ],
    [
        """srp_srednji_vek""",
        ":srednji vek",
        """Leksikon Srpskog Srednjeg Veka""",
        """Leksikon Srpskog Srednjeg Veka
        Sima Cirkovic
        Rade Mihaljcic
        Beograd 1999."""
    ],
    [
        "sind",
        ":sindikat",
        "Sindikalni Leksikon Dr Rajko Kosanović",
        """BELEŠKA O AUTORU
Dr Rajko Kosanović rođen je 25. avgusta 1959. godine.
Diplomirao je na Ekonomskom fakultetu Univerziteta u Beogradu, na smeru Ekonomska teorija.
Magistrirao je na Pravnom fakultetu Univerziteta u Beogradu, u decembru 1999. godine, kada je odbranio magistarski rad pod nazivom „Finansiranje zdravstvene zaštite u zemljama u tranziciji“.
Doktorirao je na Pravnom fakultetu Univerziteta u Beogradu 16. maja 2005. godine, odbranivši doktorsku disertaciju pod nazivom „Reforma zdravstvenog osiguranja u Republici Srbiji – pravnoekonomski aspekti“.
Zaposlen je u Savezu samostalnih sindikata Srbije, na poslovima savetnika predsednika sindikata, od 1. maja 2007. godine.
Do dolaska u Savez samostalnih sindikata Srbije bio je zaposlen u Sekretarijatu za zdravstvo Skupštine grada Beograda, od 1.aprila 1993. do 30. aprila 2007. godine, na poslovima pomoćnika sekretara.
Predsednik je Upravnog odbora Republičkog zavoda za zdravstveno osiguranje.
Objavio je oko 200 stručnih radova.
"""
    ],
    [
        "religije",
        ":religija",
        "Leksikon Temeljnih Religijskih Pojmova",
        """LEKSIKON TEMELJNIH RELIGIJSKIH POJMOVA 
ŽIDOVSTVO 
KRŠĆANSTVO 
ISLAM 

Preveli s njemačkog 
NEDELJKA PARAVIĆ (židovstvo) 
LJILJANA MATKOVIĆ VLAŠIĆ (kršćanstvo) 
ŽELJKO PAVIĆ (islam)

PROMETEJ
Zagreb, 2005.
"""
    ],
    [
        "svet_mit",
        ":mitologija",
        "Rečnik Svetske Mitologije Artur Koterel",
        """Zapadna Azija
Južna i srednja Azija
Istočna Azija
Evropa
Amerika
Afrika
Okeanija"""
    ],
    [
        "arvacki",
        "",
        "Arvacki recnik",
        """Ovdje imademo rječnik sastavljen od riječi Hrvatskih pisaca od Marina Držića pa do Vladimira Nazora.
Prema kvantumu riječi nepoznatih današnjemu standardnomu jeziku, oko 14 tisuća riječi i pojmova,
lako bi se dalo zaključiti da su ovih 60-ak pisaca pisali na nekom drugom jeziku."""
    ],
    [
        "frajer",
        "",
        "Beogradski Frajerski Rečnik",
        """Petrit Imami
Beogradski frajerski rečnik
Izdavač:
NNK International, Beograd
Terazije 14/111, tel: 011/361-36-81. 688-975
Zj izdavača:
Miroslav Damjanović, direktor
Urednik:
Vukosava Damjanović
Tehnički urednik:
Mile Nedeljković
Slog i konce:
Milan Bogdanović
Štampa:
”Margo-art” - Beograd"""
    ],
    [
        "astroloski",
        ":astrologija",
        "Astrološki Rečnik",
        """Astrologija (grč. αστρολογία - nauka o zvezdama) je pseudonauka koja položaje, odnose i kretanja
planeta i drugih nebeskih tela u odnosu na zodijačke znakove i sazvežđa, posmatrano za zemlje,
dovodi u vezu sa osnovnim osobinama individua (pojedinaca) i njihovom interakcijom, sa društvom, nacijama,
događajima u svetu itd. Sam pojam logos (gr. λόγος) se obično tumači kao reč, što bi svakako opravdalo verovanje
da zvezde nešto saopštavaju. Osoba koja izučava astrologiju se naziva astrolog, a njegov zadatak je da protumači
te "zvezdane poruke" i potom ih prenese zainteresovanima. Jedan od najbitnijih termina koji se vezuju za astrologiju
je horoskop koji se koristi pri svakoj astrološkoj analizi.

Astrologiju ne treba mešati sa astronomijom (grč. αστρονομία - pravila među zvezdama), naukom o zvezdama."""
    ],
    [
        "biblija_stari_zavet",
        "biblija stari zavet",
        "Biblija - Stari Zavet",
        """Stari zavet (originalno: Sveto pismo Staroga zavjeta) jeste hrišćanski naziv za zbirku svetih spisa koji prethode 
Isusu Hristu u Novom zavetu. Jevreji ga nazivaju Tanah. Sastoji se od knjiga zakona, istorije, proročanstava i pesama.
Za hrišćane, Stari zavet čini prvi deo Svetog pisma (drugi deo čini Novi zavet). Vremenom je hrišćanski Stari zavet
stekao izvesne razlike u odnosu na jevrejski Tanah.

Za judaizam, zbirka svetih knjiga Tanaha na hebrejskom jeziku je uobličena još u 5. veku p. n. e. 
Rana hrišćanska crkva koristila je grčki prevod iz 2. veka p. n. e, koji je uključivao i spise kojih nema u jevrejskom kanonu.
Jeronim je sumnjao u autentičnost tih knjiga, ali su one tek protestantskom reformacijom odvojene od Starog zaveta.
Pravoslavni i rimokatolici ih zovu devterokanonskim, a protestanti apokrifnim knjigama."""

    ],
    [
        "biblija_novi_zavet",
        "biblija novi zavet",
        "Biblija - Novi Zavet",
        """Novi zavet je hrišćanski naziv za zbirku spisa koje hrišćani smatraju svetim. To je novi savez Isusa koga hrišćani
smatraju Hristom sa ljudima. Hrišćani veruju da su zbivanja i zagonetke Starog zaveta dobili u njemu svoju
odgonetku i punoću. U svetom pismu Novog zaveta jevanđelisti su opisali događaje od rođenja Isusa Hrista,
pa sve do njegovog vaskrsenja iz mrtvih i vaznesenja. Opisan je i silazak Svetog duha na prve Hristove učenike,
apostole i njihovo propovedanje njegovog imena i nauke. U njima je ukazano i na sve ono što će se dogoditi do
drugog dolaska Hristovog i strašnog suda. Knjige Novog zaveta se sastoje iz četiri jevanđelja
(gr. evangelion - dobra vest ili blaga vest), nekolicine apostolskih pisama i drugih dela.
Novi zavet je na srpski preveo Vuk Stefanović Karadžić 1819. godine i objavio ga 1847. godine u Beču,
pod naslovom Novi zavjet Gospoda našega Isusa Hrista."""
    ],
    [
        "bibl_leksikon",
        ":biblija",
        "Biblijski Leksikon",
        """Poslije Rječnika biblijske teologije, koji je izašao kao prva knjiga u našoj biblioteci »Volumina theological godine 1969,
izdajemo — evo — još jedan biblijski rječnik: Biblijski leksikon. Premda opsegom manji, ovaj će Leksikon znatno pripomoći
čitaocima Biblije, studentima i propovjednicima, jer pored biblijskih pojmova donosi i biblijska realia.
Time će za određeno vrijeme, dok ne dobijemo velikog univerzalnog biblijskog leksikona, zadovoljiti otvorene potrebe.

Preveli:
MARIJAN GRGIĆ. JOSIP KOLANOVIĆ. MIUENKO ŽAGAR Stručna lektura:
ADALBERT REBIĆ
Jezična lektura:
MATE KRIZMAN, MARIJAN NEŽMAH
Korektura:
MARIJAN NEŽMAH
Grafički uredio:
MARIJAN NEŽMAH
Omot:
VLADIMIR ROZIJAN
Tipografi:
GRGO RUMORA, JOSIP ŠTOR, ŽELJKO PINTAR
Meteri:
ANTUN BAKAN i SLAVKO NAVRATIL
Izdaje: Kršćanska sadašnjost. Zagreb. Marulićev trg 14
Odgovara: dr Vjekoslav Bajsić. Zagreb. Kaptol 31
Izdano u Zagrebu 1972. © za hrvatski prijevod: Kršćanska sadašnjost. Zagreb Tisak: »Vjesnik«. Zagreb"""
    ],
    [
        "tis_mit",
        ":mitologija",
        "Tračanska, Ilirska i Slovenska Mitologija",
        """Dragoslav Srejović
Leksikon religija i mitova drevne
Evrope
"""
    ],
    [
        "dz_pravni",
        ":pravni izraz",
        "Džepni Pravni Rečnik Englesko-srpski",
        """Džepni Pravni Rečnik Englesko-srpski
Željko Kuvizić

Izdavač
Građevinska knjiga, d.o.o.
Željko Kuvizić
Direktor i glavni urednik
Stana Šehalić
Urednik
Milica Ceranić
Lektura i korektura
Rada Ilić
Dizajn korica
Nenad Samardžić
Komjuterska priprema
Radovan Galonja
Štampa
AMB Grafika, Novi Sad"""
    ],
    [
        "eponim",
        ":eponim",
        "Eponimski Leksikon (Dubravko Mršić)",
        """Eponimski Leksikon (Dubravko Mršić)
Glavna urednica
JELENA HEKMAN
Urednik
RANKO MATASOVIĆ
Recenzenti
VALENTIN PUTANEC
JOSIP SILIĆ
CIP - Katalogizacija u publikaciji
Nacionalna i sveučilišna knjižnica, Zagreb
UDK 81'373 (038)
MRŠIĆ, Dubravko
Eponimski leksikon / Dubravko Mršić. Zagreb : Matica hrvatska, 2000. - (Biblioteka Leksikoni / rječnici)
ISBN 953-150-600-0
"""
    ],
    [
        "jung",
        ":psihologija jung",
        "Leksikon Osnovnih Jungovskih Pojmova",
        """Leksikon osnovnih jungovskih pojmova

Preveo sa nemačkog Zoran Jovanović

Beograd, 1998 DERETA
 
Naslov originala: Helmut Hark Lexikon
Jungscher Grundbegriffe
"""
    ],
    [
        "hrt",
        ":hrt",
        "Leksikon Hrvatskog Radija i Televizije",
        """LEKSIKON RADIJA I TELEVIZIJE
Drugo dopunjeno i izmijenjeno izdanje u povodu devedesete obljetnice Hrvatskoga radija i šezdesete obljetnice Hrvatske televizije

Izdavački savjet
Mirko Galić, Bruno Kragić, Zdenko Ljevak, Marija Nemčić, Goran Radman, Vladimir Rončević, Saša Runjić, Ernest Strika, Sanda Vojković, Antun Vujić

Nakladnici
Hrvatska radiotelevizija Naklada Ljevak

Za nakladnike Siniša Kovačić Zdenko Ljevak

Leksikografska podrška
Leksikografski zavod Miroslav Krleža
"""
    ],
    [
        "imena",
        ":vlastito ime",
        "Znacenje Vlastitih Imena",
        """Čuvena poslovica kaže: “Velikim ljudima ne treba ništa više od svoga imena!” 
Svrha ovog rečnika jeste da Vam otkrije značenje, numerologiju, porijeklo i tumačenje Vašeg imena
ili imena Vama dragih osoba, da Vam olakša odabir imena za Vaše novorođenče.
Izbor imena za Vaše dijete treba biti zabavno i vođeno srećnim mislima, ali takođe može biti jako mukotrpan,
zbunjujuć i nadasve iscrpljujuć poduhvat, zato smo tu mi, da Vam pomognemo u toj avanturi.

Nadamo se da ćete se zabaviti i što-šta novo naučiti listajući stanice ove knjige.

Izvor: https://www.knjigaimena.com/"""
    ],
    [
        "kosarka",
        "basketball",
        "Rečnik Košarkaških termina",
        """Košarka je vrsta timskog sporta, koji uglavnom čine dva tima od po pet igrača, 
koji se takmiče jedan protiv drugog na terenu pravougaonog oblika. 
Glavni cilj je postizanje što više poena, što se ostvaruje ubacivanjem lopte u koš, 
pritom pokušavajući da protivnički tim postigne što manje poena.

Prečnik košarkaške lopte je 24 centimetra, a prečnik obruča skoro duplo više, 46 centimetara. 
Koš je vredan jedan poen prilikom šutiranja slobodnih bacanja, dva poena prilikom šutiranja unutar linije 
tri poena, i tri poena ukoliko se šutira izvan linije tri poena.

Pobjeđuje tim koji na kraju utakmice postigne najviše poena, 
ukoliko je na kraju regularnog dela rezultat izjednačen, igraju se produžeci."""
    ],
    [
        "jez_nedoum",
        "",
        "Rečnik Jezičkih Nedoumica",
        """Ivan Klajn
REČNIK JEZIČKIH
NEDOUMICA
Četvrto, prerađeno i dopunjeno izdanje 
"""
    ],
    [
        "bibliotek",
        "",
        "Enciklopedijski Leksikon Bibliotekarstva",
        """ENCIKLOPEDIJSKI LEKSIKON BIBLIOTEKARSTVA

SASTAVIO:
Dr KOSTA GRUBAČIĆ
"""
    ],
    [
        "leksikon_hji",
        "",
        "Leksikon Hrišćanstva, Judaizma i Islama",
        """LEKSIKON HRIŠĆANSTVA, JUDAIZMA I ISLAMA
Beograd, 2006.
Izdavač
AGENCIJA „MATIĆ“
Beograd
Bulevar kralja Aleksandra 192

Za izdavača
Goran Matić

Glavni i odgovorni urednik
magistrant pol. nauka Milka Pantić


Autor
mr Aleksandar Đakovac

Redaktor:
Protojerej-stavrofor prof. Dr Radovan Bigović

Recenzenti
Prof. Dr Zoran Krstić Prof. Dr Ksenija Končarević


Lektor i korektor
prof. Mirjana S. Petrović


Redakcija i prodaja:
(011) 2198-931, 415-301 063/277-069, 063/8081-945
Tel./ Faks (011) 108-514
e-mail: mgmatic@eunet.yu www.mgmatic.co.yu

Štampa:
Udruženje Nauka i društvo Srbije Beograd, Radoslava Grujića 11a

"""
    ],
    [
        "lov",
        "",
        "Leksikon Lovačkog Oružja",
        """Leksikon lovačkog oružja
Plemeniti učitelj lovstva, Mr Dušan Bojović"""
    ],
    [
        "polemologija",
        "",
        "Spisak Pojmova Polemologije",
        """Spisak pojmova od Prof. dr. Izet Beridan
Leksikon pojmova
Uvod u sigurnosne i odbrambene studije

"""
    ],
    [
        "crven_ban",
        "",
        "Crven Ban - Vuk Stefanović Karadžić",
        """VUK STEF. KARADŽIĆ OSOBITE PJESME I POSKOČICE
(prema: Srpske narodne pjesme iz neobjavljenih rukopisa Vuka Stef.
Karadžića, knjiga peta Osobite pjesme i poskočice, SANU, Beograd, 1974)
"""
    ],
    [
        "medicina",
        "",
        "Medicinski Leksikon",
        """Medicina (lat. ars medicina — "umetnost lečenja") bavi se dijagnostikom, preventivom i terapijom
fizičke i psihičke bolesti čoveka.
Medicina označava i nauku bolesti i praktičnu primenu. 
Reč medicina je izvedena iz latinske reči ars medicina, sa značenjem umetnost lečenja.
Medicina obuhvata raznovrsne aktivnosti zdravstvene zaštite kojima se održava i obnavlja 
zdravlje putem prevencije i tretmana bolesti."""
    ],
    [
        "medicina_rogic",
        "",
        "Medicinski Leksikon Momčilo Babić",
        """Babić Momčilo, glavni urednik

Članovi: 
    Bojić Milovan,
    Bošnjak-Petrović Vesna,
    Delić Dragan, 
    Drecun Vasilije, 
    Đorđević Miodrag, 
    Grbić Radivoje, 
    Hadži Đokić Jovan, 
    Ilić Aleksandar, 
    Kalezić Vasilije, 
    Kažić Tomislav, 
    Majkić-Singh Nada, 
    Manojlović Dragoljub, 
    Mujović Spomenka, 
    Oprić Miroslav, 
    Petrović Vlastimir, 
    Trišović Dimitrije
"""
    ],
    [
        "narat",
        "",
        "Naratološki Rečnik - Džerald Prins",
        """Džerald Prins - Naratološki rečnik
Prevela s engleskog
Brana Miladinov
"""
    ],
    [
        "latin",
        "",
        "Latinski Rečnik Imena",
        """Slađana Milonković i Pirpška Čaki

REČNIK VLASTITIH IMENA

STYLOS
"""
    ],
    [
        "anglicizmi",
        "",
        "Srpski Rečnik Novijih Anglicizama",
        """TVRTKO PRĆIĆ JASMINA DRAŽIĆ • MIRA MILIĆ
MILAN AJDŽANOVIĆ • SONJA FILIPOVIĆ KOVAČEVIĆ OLGA PANIĆ KAVGIĆ • STRAHINJA STEPANOV

SRPSKI REČNIK NOVIJIH ANGLICIZAMA
Prvo, elektronsko, izdanje
"""
    ],
    [
        "onkoloski",
        "",
        "Rečnik Onkoloških Termina",
        """Onkologija je specijalizovana disciplina medicine koja se bavi tretmanom kancera i tumora.
Onkologija je interdisciplinarna nauka o uzrocima, faktorima, prevenciji i opštoj prirodi
pojave i rasta i lečenja tumora."""
    ],
    [
        "pravoslavni_pojmovnik",
        "",
        "Pravoslavni Pojmovnik",
        """Pravoslavlje je jedno od grana kršćanstva. 
Korijeni pravoslavlja sežu u početke kršćanstva. 
Formira se u istočnom dijelu Rimskog carstva-Vizantiji gde i ostaje kao zvanična religija 
do propasti Vizantijskog carstva. 
Ono je specifičan oblik kršćanstva nastao u okrilju vizantinske(vaseljenske) crkve 
čiju doktrinu nasleđuje grčka crkva.
Oznaka grčka ne objašnjava, već iskrivljuje pojam. 
Od početaka Crkve, pravoslavlje je obuhvatalo manje-više celokupan prostor mediterana, 
uključujući i bliski istok i severnu afriku. 
Ni posle pada Vizantije nije se pravoslavlje pokazivalo kao grčka crkva, 
jer živi uporedo u slovenskim, arapskim i grčkim oblastima. 
Napomena je da je razuđenost obreda poseban kvalitet koji nudi pravoslavno iskustvo Crkve."""
    ],
    [
        "kuran",
        "",
        "KUR'AN (Prevod Kurana Duraković)",
        """Sarajevo, oktobra 2002. godine"""
    ],
    [
        "arhitekt",
        "",
        "Rečnik Arhitektonskog Projektovanja (Maldini)",
        """MALDINI: RECNIK ARHITEKTONSKOG PROJEKTOVANJA"""
    ],
    [
        "latin2",
        "",
        "Rečnik Latinskog Jezika",
        """Rečnik Latinskog Jezika"""
    ],
    [
        "pirot",
        "",
        "Rečnik Pirotskog Govora",
        """Pirotski govor spada u Timočko-lužnički dijalekat, belopalanačko-pirotskog poddijalekta.

To je govor grada Pirota i sela u njegovom okruženju u gradskim granicama. 
„U okviru navedenog područja uočavaju se veće ili manje razlike između govora sa desne strane Nišave 
(govor sela u podnožju Vidliča, u Visoku i Budžaku), 
koji su po mnogo čemu bliski timočkim govorima i govora sela sa leve strane Nišave, koji su bliski govorima Lužnice.“"""
    ],
    [
        "pravni_novinar",
        "",
        "Rečnik Pravnih Termina Za Novinare",
        """Ovo izdanje predstavlja dopunu priručnika
“Javnost u sudnici: novinarski vodič kroz sudske postupke”, priređivača Vladana Simeunovića, izdavača “Medija centar” Beograd, 2003
ISBN 86-82827-31-x COBISS-ID 104157452
CIP - Katalogizacija u publikaciji Narodna biblioteka Srbije, Beograd 347.91/95 (035)
343.11 (035)
"""
    ],
    [
        "poslovice",
        "",
        "Rečnik Srpskih Poslovica i Antiposlovica",
        """REČNIK SRPSKIH ANTIPOSLOVICA
Đorđe Otašević


Elektronsko izdanje - skraćeni prikaz
1. septembar 2011. br. 106. - Dodatak

© ETNA - Elektronski ĉasopis za satiru www.aforizmi.org/etna


Urednik Vesna Denĉić
"""
    ],
    [
        "turcizmi",
        "",
        "Rečnik Turcizama u Srpskom Jeziku",
        """Turcizmi su reči preuzete iz turskog jezika. 
Posebno su karakteristični za jezike naroda koji su bili pod vlašću Osmanskog carstva 
(grčki, bošnjački, bugarski, makedonski, srpski, jermenski)."""
    ],
    [
        "urbani",
        "",
        "Rečnik Urbane Svakodnevnice",
        """REČNIK URBANE SVAKODNEVNICE
Simo C. Ćirković"""
    ],
    [
        "geografija",
        "",
        "Školski Geografski Leksikon",
        """ŠKOLSKI GEOGRAFSKI LEKSIKON

Dr. sc. ZORAN CURIĆ BOŽICA CURIĆ, prof.

Nakladnik: HRVATSKO GEOGRAFSKO DRUŠTVO
Posebna izdanja, svezak 14
Urednik: Dr. sc. ZLATKO PEPEONIK, red. prof.
Recenzenti:
Dr. sc. MATE MATAS, docent Dr. sc. IVO NEJAŠMIĆ, izv. prof.
Dr. sc. TOMISLAV ŠEGOTA, red. prof.
Crteži: IVICA RENDULIĆ, prof. Lektura: JURE KARAKAŠ, prof. 
Korektura: BOŽICA CURIĆ, prof.
Grafički urednik: Dr. sc. ZORAN CURIĆ, docent
Slog, prijelom i tisak: HEROINA, Zagreb

Zagreb, 1999.
"""
    ],
    [
        "biologija",
        "",
        "Školski Leksikon Biologije",
        """Nakladnik
HINUS
Zagreb, Miramarska 13 b
tel.: 615 41 96, tel./fax: 611 55 18 e-mail: hinus@zg.hinet.hr

Urednik
Hrvoje Zrnčić

Recenzenti
Prof.dr.sc. Mirjana Kalafatić Mirjana Martek, prof.

"""
    ],
    [
        "slo_mit_encikl",
        "",
        "Slovenska Mitologija Enciklopedijski Rečnik",
        """SLOVENSKA MITOLOGIJA
ENCIKLOPEDIJSKI REČNIK

Redaktori

Svetlana M. Tolstoj
Ljubinko Radenković

ZEPTER BOOK WORLD BEOGRAD 2001
"""
    ],
    [
        "tehnicki",
        "",
        "Tehnički Rečnik",
        """Tehnički rečnik sadrži prevode stručnih termina iz različitih tehničkih oblasti kao što su: 
građevinarstvo, arhitektura, mašinstvo, elektrotehnika, saobraćaj, hemija itd. 
"""
    ],
    [
        "tolkin",
        "",
        "Tolkinov Rečnik",
        """Džon Ronald Ruel Tolkin (engl. John Ronald Reuel Tolkien;
Blumfontejn, 3. januar 1892 — Oksford, 2. septembar 1973) 
bio je engleski univerzitetski profesor, književnik i filolog.

Van naučnih krugova, najpoznatiji je kao autor romana „Gospodar prstenova“, zatim njegovog prethodnika, „Hobita“,
kao i velikog broja posthumno izdatih knjiga o istoriji zamišljenog sveta zvanog Arda, najviše jednog njenog kontinenta,
Srednje zemlje, gde se odigrava radnja ova njegova dva najpoznatija romana. 
Velika popularnost i uticaj ovih dela su ustoličila Tolkina kao oca žanra moderne epske fantastike."""
    ],
    [
        "istorijski",
        "",
        "Istorijski Leksikon",
        """Vladimir Ćorović
ISTORIJSKI LEKSIKON
Članci Vladimira Ćorovića
iz Narodne enciklopedije srpsko-hrvatsko-slovenačke
(1925-1929)

Podaci o štampanom izdanju:
Vladimir Ćorović: Istorijski leksikon
Srpska književna zadruga Beograd, 2006.

Priredio DRAGAN LAKIĆEVIĆ

Predgovor NENAD LjUBINKOVIĆ
"""
    ],
    [
        "vlaski",
        "",
        "Vlaški Rečnik",
        """Vlasi je naziv za vlahofono autohtono stanovništvo u istočnoj Srbiji koja živi na području između
Morave, Timoka i Dunava na teritoriji četiri okruga: Borski, Braničevski, Zaječarski i Pomoravski. 
U manjem broju naseljavaju i Podunavski, Nišavski i Rasinski. 
U pogledu konfesije, Vlasi pripadaju pravoslavnoj veroispovesti i Srpskoj pravoslavnoj crkvi, 
ali su se u njihovom verskom i obrednom životu do danas zadržali u znatnoj svežini mnogi 
prehrišćanski, indoevropski i paleobalkansi elementi."""
    ],
    [
        "zakon_krivicni_zakonik",
        "",
        "Krivični Zakonik",
        """Krivični Zakonik"""

    ],
    [
        "zakon_krivicni_postupak",
        "",
        "Zakonik o Krivičnom Postupku",
        """ZAKONIK O KRIVIČNOM POSTUPKU

("Sl. glasnik RS", br. 72/2011, 101/2011, 121/2012, 32/2013, 45/2013, 55/2014 i 35/2019)"""
    ],
    [
        "zakon_o_radu",
        "",
        "Zakon o Radu",
        """Preuzeto sa www.pravno-informacioni-sistem.rs
Redakcijski prečišćen tekst
 
ZAKON o radu
"Službeni glasnik RS" br. 24 od 15. marta 2005, 61 od 18. jula 2005, 54 od 17. jula 2009, 32 od 8. aprila 2013, 75 od 21. jula 2014, 13 od 24. februara 2017 - US, 113 od 17. decembra 2017, 95 od 8. decembra 2018 - Autentično tumačenje
 
"""
    ],
    [
        "dusan",
        "",
        "Dušanov Zakonik",
        """Dušanov zakonik
Stojan Jasić

Donošenje Dušanovog Zakonika uslovili su različiti, ekonomski i društveni momenti. 
Sredinom XIV veka privremena moć feudalne Srbije dospela je do kulminacione tačke. 
Osvajanja stranih teritorija imala su za posledicu porast etničke i političke raznolikosti društva feudalne Srbije. 
Zbog njih se i pravni poredak feudalne Srbije suočio sa novim stranim pravnim primesama. 
Vladajuća srpska vlastela našla su se otuda u složenoj i delikatnoj političkoj situaciji. 
Sem toga, odnose između vlastele i eksploatisanog sebarskog stanovništva trebalo je regulisati po 
jedinstvenim pravnim kriterijumima i jedinstvenim pravnim sredstvom. 
Na osnovama dotadašnjeg razvoja, sredinom XIV veka, stekle su se povoljne okolnosti da i vladalac feudalne Srbije 
cementira i pravno izrazi svoju supremaciju. 
U suštini ti momenti izazivali su potrebu da se stvori političko spajanje raznorodnih delova državne 
teritorije i unificiranje pravnog poretka. 
Tim ciljevima imao je da posluži Dušanov Zakonik. 
Donošenju Zakonika prethodili su značajni unutrašnjo-politički akti. 
Godine 1346. srpska arhiepiskopija proglašena je za patrijaršiju, a kralj krunisan za cara. 
Izrada Dušanovog Zakonika predstavljala je dopunsku, tj. propratnu pravno-političku aktivnost na liniji tih događaja. 
Osnovano se misli da je inicijativa za nju data 1347. g. - Materijal za to pruža Dušanovo pismo sačuvano 
uz Rakovački prepis Dušanovog Zakonika. 
U izvore Dušanovog Zakonika ubrajaju se: izvesni vizantijski propisi (osobito Vasilike), vladalačke povelje, 
ugovori sa Dubrovnikom i pravni običaji.

Dušanov Zakonik ulazi u regulisanje kako međusobnih odnosa crkvenih ustanova tako i crkve i države,
a brine i o organizacionom jačanju crkve.
Dušanov Zakonik propisuje i sudski postupak i reguliše ga u pogledu njegovog toka, upotrebe izvesnih dokaznih sredstava. 
Kako u izvesnim materijalno-pravnim tako i u procesnim odredbama i Dušanov Zakonik ozakonjuje nejednako 
postupanje s licima s obzirom na njihovu stalešku pripadnost.
"""
    ],
    [
        "zakon_upravni",
        "",
        "Zakon o Opštem Upravnom Postupku",
        """Preuzeto sa www.pravno-informacioni-sistem.rs
Redakcijski prečišćen tekst
 
ZAKON o opštem upravnom postupku
"Službeni glasnik RS", br. 18 od 1. marta 2016, 95 od 8. decembra 2018 - Autentično tumačenje, 2 od 13. januara 2023 - US
"""
    ],
    [
        "zakon_razni",
        "",
        "Kolekcija Raznih Zakona",
        """ZAKON  O MIRNOM REŠAVANjU RADNIH SPOROVA 
Sl. glasnik RS br. 125/04 , 104/09 , 50/18 )

ZAKON O SPREČAVANjU ZLOSTAVLjANjA NA RADU
("Službeni glasnik RS", broj 36/10)
Osnovni tekst na snazi od 05.06.2010. godine, u primeni od 04.09.2010.godine

ZAKON O VOLONTIRANjU
("Službeni glasnik RS", broj 36/10)
Osnovni tekst na snazi od 05.06.2010.godine , u primeni od 06.12.2010.godine

ZAKON o agencijskom zapošljavanju
"Službeni glasnik RS", broj 86 od 6. decembra 2019.

ZAKON O PLANIRANJU I IZGRADNJI
('Sl. glasnik RS', br. 72/2009, 81/2009 - ispr., 64/2010 - odluka US, 24/2011, 121/2012, 42/2013 - odluka US, 50/2013 - odluka US, 98/2013 - odluka US, 132/2014, 145/2014, 83/2018, 31/2019, 37/2019 - dr. zakon, 9/2020, 52/2021 i 62/2023)

ZAKON O VISOKOM OBRAZOVANJU
("Sl. glasnik RS", br. 88/2017, 73/2018, 27/2018 - dr. zakon, 67/2019, 6/2020 - dr. zakoni, 11/2021 - autentično tumačenje, 67/2021 i 67/2021 - dr. zakon)

ZAKON O STEČAJU
("Sl. glasnik RS", br. 104/2009, 99/2011 - dr. zakon i 71/2012 - odluka US)

ZAKON O STEČAJU I LIKVIDACIJI BANAKA I DRUŠTAVA ZA OSIGURANjE
("Sl. glasnik RS", br. 14/2015 i 44/2018 - dr. zakon)

Zakon o sanaciji, stečaju i likvidaciji banaka
(„Službeni list SFRJ“, 84/89 i 63/90 i „Službeni list SRJ“, br.37/93, 26/95, 28/96, 44/99 i 53/01)

Zakon o prinudnom poravnanju, stečaju i likvidaciji
(„Službeni list SFRJ“, br. 84/89 i „Službeni list SRJ“, br. 37/93 i 28/96)

ZAKON O STANOVANjU I ODRŽAVANjU ZGRADA
("Sl. glasnik RS", br. 104/2016 i 9/2020 - dr. zakon)
Napomena: Odredbe člana 140. Zakona o stanovanju i održavanju zgrada ("Sl. glasnik RS", br. 104/2016) prestaju da važe 12. februara 2020. godine, danom stupanja na snagu Zakona o izmenama i dopunama Zakona o planiranju i izgradnji ("Sl. glasnik RS", br. 9/2020).

ZAKON O LEČENjU NEPLODNOSTI POSTUPCIMA BIOMEDICINSKI POTPOMOGNUTOG OPLOĐENjA
"Sl.glasnik RS", br. 72/2009 od 3.9.2009. godine.

ZAKON O UPRAVLJANJU OTPADOM
("Sl. glasnik RS", br. 36/2009, 88/2010 i 14/2016)

ZAKON O ZAŠTITI ŽIVOTNE SREDINE
("Sl. glasnik RS", br. 135/2004, 36/2009, 36/2009 - dr. zakon, 72/2009 - dr. zakon, 43/2011 - odluka US i 14/2016)

ZAKON O HEMIKALIJAMA
("Sl. glasnik RS", br. 36/2009, 88/2010, 92/2011, 93/2012 i 25/2015)

ZAKON O TRANSPORTU OPASNOG TERETA
("Sl. glasnik RS", br. 88/2010)

ZAKON O AMBALAŽI I AMBALAŽNOM OTPADU

PRAVILNIK O KATEGORIJAMA, ISPITIVANJU I KLASIFIKACIJI OTPADA
("Sl. glasnik RS", br. 56/2010)

ZAKON O ZAŠTITI PODATAKA O LIČNOSTI

ZAKON O PROCENITELjIMA VREDNOSTI NEPOKRETNOSTI

ZAKON O POREZU NA DODATU VREDNOST
(„Službeni glasnik RS”, broj 84/04, 86/04  - ispravka, 61/05, 61/07, 93/12, 108/13, 68/14  - dr. zakon, 142/14, 83/15, 108/16, 113/17, 30/18, 72/19, 153/20, 138/22) 
Prečišćen tekst zaključno sa izmenama iz Sl. gl. RS br. 138/22  koje su u primeni od 01/01/2023  
(izmene u čl.: 10a, 14, 15, 17, 17a, 17b, 17v, 28, 51a, 56a).

ZAKON O JAVNIM NABAVKAMA

ZAKON O DOBROBITI ŽIVOTINjA

ZAKON O RODNOJ RAVNOPRAVNOSTI

ZAKON O IZBEGLICAMA
("Sl. glasnik RS", br. 18/92, "Sl. list SRJ", br. 42/2002 - odluka SUS i "Sl. glasnik RS", br. 30/2010)

ZAKON O DUALNOM OBRAZOVANJU
("Sl. glasnik RS", br. 101/2017)

ZAKON O ELEKTRONSKIM KOMUNIKACIJAMA
("Sl. glasnik RS", br. 44/2010, 60/2013 - odluka US i 62/2014)

ZAKON O JAVNOJ SVOJINI
("Sl. glasnik RS", br. 72/2011, 88/2013, 105/2014, 104/2016 - dr. zakon, 108/2016, 113/2017 i 95/2018)

ZAKON O OBLIGACIONIM ODNOSIMA
("Sl. list SFRJ", br. 29/78, 39/85, 45/89 - odluka USJ i 57/89, "Sl. list SRJ", br. 31/93 i "Sl. list SCG", br. 1/2003 - Ustavna povelja)

Zakon o osnovnom obrazovanju i vaspitanju
Zakon je objavljen u "Službenom listu RCG", br. 64/2002, 49/2007  i u "Službenom listu CG", br. 45/2010, 40/2011 - drugi zakon, 39/2013 (čl. 34. nije u prečišćenom tekstu) i 47/2017.

ZAKON O LOKALNOJ SAMOUPRAVI
"Sl. glasnik RS", br. 129/2007, 83/2014 - dr. zakon, 101/2016 - dr. zakon, i 47/2018

ZAKON O SLUŽBENOJ UPOTREBI JEZIKA I PISAMA
("Sl. glasnik RS", br. 45/91, 53/93, 67/93, 48/94, 101/2005 - dr. zakon, 30/2010, 47/2018 i 48/2018 - ispr.)

ZAKON O ULAGANJIMA
("Sl. glasnik RS", br. 89/2015 i 95/2018)

"""
    ],
    [
        "ustav",
        "",
        "Ustav Republike Srbije",
        """U S T A V	R E P U B L I K E   S R B I J E
Objavljen u "Službenom glasniku RS", br. 98/2006
"""
    ]

]

for i in names_map:
    if i[0] in dictionary["serbian"]["dicts"]:
        dictionary["serbian"]["dicts"][i[0]]["@search_string"] = i[1]
        dictionary["serbian"]["dicts"][i[0]]["@name"] = i[2]
        dictionary["serbian"]["dicts"][i[0]]["@desc"] = i[3]

print (f"Added names and description for {len(names_map)} dictionaries.")


with open("dictionary.json", "w", encoding="utf-8") as file:
    json.dump(dictionary, file, indent=2)

total_t_start = int((time.perf_counter() - total_t_start))
print ()
print (print (f"  All Done.  TIME: {int(total_t_start/3600):02d}:{int(total_t_start/60)%60:02d}:{total_t_start%60:02d}"))

